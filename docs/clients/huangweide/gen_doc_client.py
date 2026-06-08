from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

# font helper
def set_font(run, size=11, bold=False, color=None, east='微軟正黑體'):
    run.font.name = east
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east)
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_heading('', level=1)
    p.clear()
    run = p.add_run(text)
    set_font(run, 17, True, (0x0F,0x76,0x6E))
    return p

def h2(text):
    p = doc.add_heading('', level=2)
    p.clear()
    run = p.add_run(text)
    set_font(run, 13, True, (0x05,0x96,0x69))
    return p

def h3(text):
    p = doc.add_heading('', level=3)
    p.clear()
    run = p.add_run(text)
    set_font(run, 12, True, (0x1A,0x20,0x2C))
    return p

def body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 11)
    return p

def bold_body(segments):
    # segments: list of (text, is_bold)
    p = doc.add_paragraph()
    for txt, is_bold in segments:
        run = p.add_run(txt)
        set_font(run, 11, is_bold)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, 11)
    p.paragraph_format.left_indent = Cm(1 + level * 0.8)
    return p

def confirm_box(text):
    """A shaded confirmation box"""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    # shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF7ED')
    cell._tc.get_or_add_tcPr().append(shd)
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run('📋 請確認：' + text)
    set_font(run, 11, True, (0xC2,0x41,0x0C))
    return t

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0F766E')
        cell._tc.get_or_add_tcPr().append(shd)
        for run in cell.paragraphs[0].runs:
            set_font(run, 11, True, (0xFF,0xFF,0xFF))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        row = t.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            if ri % 2 == 0:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F0FDF4')
                cell._tc.get_or_add_tcPr().append(shd)
            for run in cell.paragraphs[0].runs:
                set_font(run, 10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

def sp(n=1):
    for _ in range(n): doc.add_paragraph()

# ══════════════════════════════════════════
# COVER
# ══════════════════════════════════════════
sp()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('普健生醫\nAI 行銷管理平台\n規格確認書')
set_font(r, 24, True, (0x0F,0x76,0x6E))

sp()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run('請黃維德藥師確認以下規格後，即可進入開發階段')
set_font(r2, 12, False, (0x64,0x74,0x8B))

sp()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = date_p.add_run('2026 年 6 月 6 日')
set_font(r3, 11, False, (0x94,0xA3,0xB8))

doc.add_page_break()

# ══════════════════════════════════════════
# 前言
# ══════════════════════════════════════════
h1('前言：這個平台是做什麼的？')
sp()
body('這份文件的目的是讓您確認：我們即將為普健生醫建置的「AI 行銷管理平台」，功能與運作方式是否符合您的期望。')
sp()
body('用一句話說明這個平台的價值：')
sp()
bold_body([
    ('行銷人員每天上來，點一個按鈕，AI 自動幫普健總站和每一間聯盟藥局各產 3 篇文章，', False),
    ('藥師審閱修改後送出，主管批准，文章就上到網站 — 全程不需要懂任何技術。', True),
])
sp()
body('這個平台解決的核心問題：')
bullet('一間藥局的藥師不擅長寫 SEO 文章，有了 AI 自動產文，不再是負擔')
bullet('100 間藥局同時更新內容，若靠人工根本不可能，AI 讓這件事變成一個按鈕')
bullet('每間藥局的文章角度不同，避免 Google 認為是重複內容')
bullet('所有聯盟藥局共享同一份藥局地圖，顧客可以找到離自己最近的合作藥局')

doc.add_page_break()

# ══════════════════════════════════════════
# 1. 誰會用這個平台
# ══════════════════════════════════════════
h1('一、誰會用這個平台？')
sp()
body('平台有三種使用者，登入方式統一用 Gmail（Google 帳號），不需要另外設密碼。')
sp()

add_table(
    ['使用者類型', '適用對象', '每天做什麼'],
    [
        ['最高管理員', '黃維德藥師', '查看所有藥局、所有文章、帳號管理，具有全部權限'],
        ['行銷主管', '行銷團隊主管', '每日一鍵產文、審稿、批准文章上線、管理藥局資料'],
        ['藥師', '各聯盟藥局藥師', '登入後只看到自己藥局今天的 3 篇草稿，修改後送出審核'],
    ],
    col_widths=[3.5, 3.5, 9]
)

sp()
confirm_box('以上三種角色是否符合您的期望？是否有需要新增或調整的角色？')
sp()

# ══════════════════════════════════════════
# 2. 每天的作業是怎麼跑的
# ══════════════════════════════════════════
h1('二、每天的作業是怎麼跑的？')
sp()
h2('行銷主管的每天流程')
sp()

add_table(
    ['步驟', '操作', '說明'],
    [
        ['1', '登入平台', '用自己的 Gmail 帳號登入，看到今日 Dashboard'],
        ['2', '選今日主題', '從四大主題中選擇（可多選）：\n過敏保養 / 葉黃素護眼 / 關節保健 / 體重管理'],
        ['3', '點「一鍵產文」', 'AI 自動幫普健總站產 3 篇 + 每間聯盟藥局各產 3 篇\n每篇角度不同，約 2 分鐘完成'],
        ['4', '審稿', '逐篇確認，可直接修改，或退回給藥師修改（附留言說明）'],
        ['5', '批准發布', '按批准後，文章即時上到各藥局網站'],
    ],
    col_widths=[1.5, 4, 11]
)

sp()
h2('藥師的每天流程')
sp()

add_table(
    ['步驟', '操作', '說明'],
    [
        ['1', '登入平台', '用自己的 Gmail 帳號登入，只看到自己藥局的內容'],
        ['2', '查看今日 3 篇草稿', 'AI 已根據藥師姓名、藥局地區自動產出，每篇角度不同'],
        ['3', '逐篇閱讀 / 修改', '可在編輯器直接修改措辭，確保符合自己的風格與法規'],
        ['4', '送出審核', '確認後按「送出」，行銷主管收到通知'],
        ['5', '等待上線通知', '主管批准後，文章自動上線，藥師收到通知'],
    ],
    col_widths=[1.5, 4, 11]
)

sp()
confirm_box('藥師送出後，需要行銷主管再次批准才上線（非藥師自行發布）— 此流程您是否同意？')
sp()

# ══════════════════════════════════════════
# 3. 文章怎麼避免重複
# ══════════════════════════════════════════
h1('三、每間藥局的文章為什麼不會一樣？')
sp()
body('Google 若判定 100 間藥局的文章內容雷同，會降低所有網站的搜尋排名。')
sp()
body('本平台在 AI 產文時，會自動為每間藥局注入不同的「背景資料」，讓每篇文章都有真正不同的角度：')
sp()

add_table(
    ['分站', '文章角度', '範例標題'],
    [
        ['普健生醫總站', '學術 / 成分解析', '入秋過敏高峰！DP2 外用專利技術的舒緩機制解析'],
        ['台北大安 長青藥局', '在地觀察（藥師視角）', '台北盆地秋季過敏特別嚴重？陳藥師的 10 年觀察'],
        ['新竹東區 長青藥局', '季節 / 氣候切入', '新竹風城乾燥氣候，過敏族群這樣保養最有效'],
        ['桃園中壢 長青藥局', '家庭故事敘事', '一個媽媽的故事：孩子過敏反覆發作怎麼辦'],
        ['台中西屯 長青藥局', 'Q&A 問答格式', 'Q&A｜過敏藥吃了想睡怎麼辦？台中張藥師解答'],
    ],
    col_widths=[4, 4, 8.5]
)

sp()
body('每篇文章的結尾，都會自動加入「加入 LINE OA 領取免費衛教資訊」的按鈕，導流到各藥局的 LINE 帳號。')
sp()
confirm_box('四大主題（過敏 / 葉黃素 / 關節 / 體重管理）是否為目前主要推廣方向？是否需要調整或新增其他主題？')
sp()

# ══════════════════════════════════════════
# 4. 新藥局怎麼加入
# ══════════════════════════════════════════
h1('四、新藥局如何加入聯盟？')
sp()
body('藥局加入聯盟後，需完成一次性的設定作業，預計 20-30 分鐘，全程有圖文步驟引導。')
sp()

add_table(
    ['步驟', '誰來做', '內容'],
    [
        ['1  發出邀請', '行銷主管', '在平台點「邀請藥局」，系統產生邀請連結，用 LINE 或 Email 傳給藥局'],
        ['2  填寫藥局資料', '藥局', '填入：藥局名稱、地址、電話、藥師姓名、執照號碼、上傳藥師大頭照、選擇主打服務'],
        ['3  設定登入帳號', '藥局', '填入藥師 Gmail，之後藥師用此 Gmail 登入平台，不需另設密碼'],
        ['4  建立藥局網站', '藥局\n（平台引導）', '① 免費申請 Vercel 帳號\n② 在 Vercel 購買自己的域名（約 NT$700/年，藥局自付）\n③ 點一下「一鍵建站」按鈕，約 2 分鐘完成\n④ 複製 2 個資料（系統 ID + 金鑰）貼回平台'],
        ['5  系統自動完成', '系統', '驗證成功後，藥局網站自動上線、出現在聯盟地圖、行銷主管收到通知'],
    ],
    col_widths=[3.5, 3, 10]
)

sp()
body('說明：Vercel 是一個免費的網站架設平台（全球知名，Next.js 官方平台）。域名費用由各藥局自行承擔，普健無需代墊。')
sp()
confirm_box('藥局在加入時需要自費購買域名（約 NT$700/年），這個費用結構您是否同意？是否需要調整加盟條件？')
sp()

# ══════════════════════════════════════════
# 5. 顧客如何找到附近藥局
# ══════════════════════════════════════════
h1('五、顧客如何找到附近的聯盟藥局？')
sp()
body('每個藥局的網站，都會在頁面底部顯示一張「聯盟藥局地圖」，標示全台所有加入普健聯盟的藥局位置。')
sp()
body('顧客動線：')
bullet('Google 搜尋「過敏 新竹藥局」或「異位性皮膚炎保養」')
bullet('落地到對應的藥局分站，閱讀專業文章')
bullet('頁面底部看到「找離你最近的聯盟藥局」地圖')
bullet('點地圖上的藥局 → 直接 Google 導航 / 加入 LINE OA 預約諮詢')
sp()
confirm_box('顧客進站後，希望引導他們採取什麼行動為優先？\n□ 加入 LINE OA   □ 直接電話聯絡   □ 預約現場諮詢   □ 其他：___________')
sp()

# ══════════════════════════════════════════
# 6. 平台功能清單
# ══════════════════════════════════════════
h1('六、平台功能完整清單')
sp()
body('以下為平台規劃的所有功能，請確認是否符合需求，或有需要移除 / 新增的項目。')
sp()

add_table(
    ['功能', '說明', '必要 / 可選'],
    [
        ['Gmail 登入', '所有使用者用 Google 帳號登入，無需另設密碼', '必要'],
        ['一鍵 AI 產文', '選主題後，AI 自動為所有藥局產出當日文章', '必要'],
        ['文章編輯器', '行銷主管 / 藥師可修改 AI 產出的草稿', '必要'],
        ['送審 / 批准流程', '藥師送審 → 主管批准 → 自動上線', '必要'],
        ['藥局管理', '新增、編輯、停用藥局資料，含藥師帳號設定', '必要'],
        ['藥局邀請 / 上線引導', '5 步驟引導新藥局完成設定', '必要'],
        ['聯盟藥局地圖', '各分站顯示全台聯盟藥局地圖，支援導航', '必要'],
        ['LINE OA 自動注入', '每篇文章結尾自動加入各藥局的 LINE 加入按鈕', '必要'],
        ['見證案例管理', '上傳使用者見證（照片 + 短文），同步到所有分站', '建議'],
        ['DP2 說明頁管理', '統一維護 DP2 專利說明，一次更新同步所有分站', '建議'],
        ['轉化數據追蹤', '各站文章瀏覽量、LINE 按鈕點擊數統計', '建議'],
        ['發布歷史記錄', '查看所有歷史文章，依日期、藥局、主題篩選', '建議'],
    ],
    col_widths=[4, 10, 2.5]
)

sp()
confirm_box('以上功能清單是否符合需求？「建議」項目是否列入開發範圍？')
sp()

# ══════════════════════════════════════════
# 7. 開發時程
# ══════════════════════════════════════════
h1('七、建議開發時程')
sp()
body('開發分三期進行，第一期完成後即可開始邀請藥局加入測試。')
sp()

add_table(
    ['期別', '預計時程', '完成後可以做的事'],
    [
        ['第一期', '第 1–2 週', '行銷主管可登入、選主題、一鍵產文、審稿、發布\n第一批文章可以上到網站'],
        ['第二期', '第 3–4 週', '藥師可登入查看自己的文章、修改、送出審核\n可以開始邀請第一批藥局加入'],
        ['第三期', '第 5–6 週', 'LINE OA 自動注入、聯盟地圖上線、見證案例管理\n完整閉環：流量 → 文章 → LINE → 諮詢'],
    ],
    col_widths=[2.5, 3, 11]
)

sp()
confirm_box('第一期完成（約 2 週後）是否有計畫邀請哪幾間長青藥局作為試辦？請提供聯絡窗口讓我們提前準備邀請連結。')
sp()

# ══════════════════════════════════════════
# 簽核欄
# ══════════════════════════════════════════
h1('八、確認簽核')
sp()
body('請您閱讀完畢後，在下方簽名確認規格，我們即可正式進入開發階段。如有任何問題或修改意見，歡迎直接在文件上標注。')
sp(2)

sign_t = doc.add_table(rows=3, cols=3)
sign_t.style = 'Table Grid'
labels = [['確認項目', '確認結果', '備註'],
          ['規格內容符合需求', '□ 同意   □ 需調整', ''],
          ['同意進入開發階段', '□ 同意   □ 待確認', '']]
for ri, row_data in enumerate(labels):
    for ci, val in enumerate(row_data):
        cell = sign_t.rows[ri].cells[ci]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            set_font(run, 11, ri == 0)
        if ri == 0:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '0F766E')
            cell._tc.get_or_add_tcPr().append(shd)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

for row in sign_t.rows:
    row.cells[0].width = Cm(6)
    row.cells[1].width = Cm(5)
    row.cells[2].width = Cm(5.5)

sp(2)

sign_line = doc.add_paragraph()
sign_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_sign = sign_line.add_run('黃維德 藥師 簽名：___________________　　日期：________')
set_font(r_sign, 11)

sp()
note_p = doc.add_paragraph()
note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_note = note_p.add_run('本文件由普健生醫行銷團隊製作　2026 年 6 月 6 日')
set_font(r_note, 9, False, (0x94,0xA3,0xB8))

# ── Save ──
out = r'C:\Users\JasonLee\claude_code_projects\CMO\docs\clients\huangweide\普健生醫平台規格確認書_20260606.docx'
doc.save(out)
print(f'SAVED: {out}')
