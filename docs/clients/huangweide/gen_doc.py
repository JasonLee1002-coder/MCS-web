from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

# ── Default font ──
style = doc.styles['Normal']
style.font.name = 'Noto Sans TC'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    p.runs[0].font.size = Pt(18)
    p.runs[0].font.name = '微軟正黑體'
    p.runs[0].bold = True
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = '微軟正黑體'
    p.runs[0].bold = True
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = '微軟正黑體'
    p.runs[0].bold = True
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.runs  # clear
    if bold_parts:
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.font.size = Pt(11)
            run.font.name = '微軟正黑體'
            if i % 2 == 1:
                run.bold = True
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = '微軟正黑體'
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '微軟正黑體'
    p.paragraph_format.left_indent = Cm(1 + level * 0.8)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.name = '微軟正黑體'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '0F766E')
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.name = '微軟正黑體'
            if ri % 2 == 0:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'F0FDF4')
                cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(width)
    return t

def spacer(n=1):
    for _ in range(n):
        doc.add_paragraph()

# ════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('普健生醫\nAI 行銷管理平台\n設計規劃書')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
run.font.name = '微軟正黑體'

spacer(1)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('藥局聯盟 × AI 產文 × 流量閉環\n整合平台設計說明')
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
run2.font.name = '微軟正黑體'

spacer(1)
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = date_p.add_run(f'版本日期：2026 年 6 月 6 日\n提案單位：普健生醫行銷團隊')
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
run3.font.name = '微軟正黑體'

doc.add_page_break()

# ════════════════════════════════════════════════
# 1. 平台定位
# ════════════════════════════════════════════════
heading1('一、平台定位')
body('本平台名稱：**普健生醫 AI 行銷管理中心**', bold_parts=True)
spacer()
body('核心目標：讓行銷人員不需碰任何技術工具，在一個網頁介面完成從「AI 產文」到「全站發布」的完整作業，並讓各藥局藥師可以審閱、修改、送審自己藥局的每日文章。')
spacer()
body('部署位置：架設於黃維德先生的 Vercel 帳號，使用 Vercel Postgres 資料庫，技術棧為 Next.js 14。')

spacer()

# ════════════════════════════════════════════════
# 2. 角色與權限
# ════════════════════════════════════════════════
heading1('二、角色與權限設計')
body('平台共分三種登入角色，所有人統一使用 Gmail 帳號登入，系統自動對應角色與權限。')
spacer()

add_table(
    ['角色', '適用對象', '看得到的資料', '可執行操作'],
    [
        ['super_admin\n（最高管理員）', '黃維德藥師', '全部藥局、全部文章、所有帳號', '所有功能 + 帳號管理 + 系統設定'],
        ['admin\n（行銷主管）', '行銷主管', '全部藥局、全部文章', 'AI 產文、審稿、批准發布、管理藥局資料'],
        ['pharmacist\n（藥師）', '各藥局藥師', '僅自己藥局今日 3 篇文章', '編輯文章、送出審核'],
    ],
    col_widths=[3, 3, 5, 5.5]
)

spacer()
body('※ 藥師送出審核後，需由 admin 或 super_admin 批准，文章才正式上線。藥師無法自行發布。')

spacer()

# ════════════════════════════════════════════════
# 3. 藥師日常流程
# ════════════════════════════════════════════════
heading1('三、藥師每日操作流程')
body('藥師的介面極度簡化，每日登入只做一件事：')
spacer()

steps_pharma = [
    ('Step 1', '登入', '用自己的 Gmail 帳號登入平台，系統自動識別身份，只顯示自己藥局的內容。'),
    ('Step 2', '查看今日 3 篇草稿', 'AI 已根據今日主題、藥師姓名、藥局地區自動產生 3 篇不同角度的文章草稿。'),
    ('Step 3', '逐篇閱讀 / 修改', '點開每篇文章，可在編輯器直接修改內容、調整語氣，確保符合藥局風格與法規。'),
    ('Step 4', '按「送出審核」', '確認無誤後送出，行銷主管收到通知進行最終審查。'),
    ('Step 5', '主管批准，自動上線', '主管批准後，文章即時發布至該藥局的網站，藥師收到上線通知。'),
]

for num, title_s, desc in steps_pharma:
    heading3(f'{num}｜{title_s}')
    body(desc)
    spacer()

# ════════════════════════════════════════════════
# 4. 行銷主管每日流程
# ════════════════════════════════════════════════
heading1('四、行銷主管每日操作流程')
spacer()

steps_admin = [
    ('Step 1', '登入 Dashboard', '看到今日全局狀況：幾篇待審、幾篇已發布、哪些藥局今日尚未動作。'),
    ('Step 2', '選擇今日主題', '從四大主題選一或多項：過敏保養 / 葉黃素護眼 / 關節保健 / 體重管理。'),
    ('Step 3', '一鍵產文', '點「開始產文」，系統同時呼叫 AI，為普健總站產 3 篇 + 每間藥局各產 3 篇，每篇角度不同，約需 2 分鐘。'),
    ('Step 4', '審稿 / 批准', '查看所有待審文章（含藥師送來的修改版），逐篇確認或退回修改，附留言說明。'),
    ('Step 5', '全部批准後發布', '批准後即自動上線至各藥局網站，無需手動操作任何技術環境。'),
]

for num, title_s, desc in steps_admin:
    heading3(f'{num}｜{title_s}')
    body(desc)
    spacer()

# ════════════════════════════════════════════════
# 5. 避免 Google 重複內容機制
# ════════════════════════════════════════════════
heading1('五、避免 Google 重複內容的技術設計')
body('100 家藥局同日發布同主題文章，若內容相同會被 Google 判定為「薄內容」並降低排名。本平台透過以下機制確保每篇文章真正不同：')
spacer()

add_table(
    ['分站', '文章角度', '範例標題'],
    [
        ['普健總站', '學術 / 成分解析', '入秋過敏高峰！DP2 外用專利技術的舒緩機制'],
        ['長青台北大安', '在地觀察（藥師視角）', '台北盆地秋季過敏特別嚴重？陳藥師的在地觀察'],
        ['長青新竹東區', '季節 / 氣候切入', '新竹風城乾燥氣候，過敏族這樣保養最有效'],
        ['長青桃園中壢', '家庭故事敘事', '一個媽媽的故事：孩子過敏讓全家不得安寧'],
        ['長青台中西屯', 'Q&A 問答格式', 'Q&A｜過敏藥吃了想睡怎麼辦？台中張藥師解答'],
    ],
    col_widths=[4, 4, 8]
)
spacer()
body('系統在產文時自動注入：藥師姓名、藥局所在縣市區域、主打服務項目、指定角度（系統輪流分配），確保 100 篇文章 100 種視角。')

spacer()

# ════════════════════════════════════════════════
# 6. 藥局加入流程
# ════════════════════════════════════════════════
heading1('六、藥局加入平台完整流程')
body('藥局加入聯盟後，需完成以下一次性設定，約 20-30 分鐘，全程有圖文引導。')
spacer()

onboard_steps = [
    ('Step 1', '普健發出邀請',
     '行銷主管在管理平台點「邀請新藥局」，系統產生專屬邀請連結（7 天有效），透過 Email 或 LINE 傳給藥局。'),
    ('Step 2', '藥局填寫基本資料',
     '藥局點連結後進入引導頁面，填入：藥局名稱、地址、電話、LINE OA 帳號、藥師姓名及執照號碼、藥師大頭照、主打服務項目（多選）。'),
    ('Step 3', '設定藥師登入帳號',
     '填入藥師 Gmail 信箱，之後藥師即以此 Gmail 登入平台。無需另外設定密碼。'),
    ('Step 4', '申請 Vercel 並設定網站\n（平台提供圖文引導，全程約 15 分鐘）',
     '4-1  申請 Vercel 免費帳號\n4-2  在 Vercel 購買藥局專屬域名（如 changqing-daan.tw，約 NT$700/年，刷藥局自己的信用卡，自動完成 DNS 設定）\n4-3  點擊普健提供的「一鍵部署」按鈕，自動建立藥局網站（約 2 分鐘）\n4-4  複製 Vercel Project ID 貼回引導頁面\n4-5  在 Vercel 產生 API Token，複製貼回引導頁面'),
    ('Step 5', '系統自動完成後續',
     '平台驗證連線成功後，自動完成：\n・將藥局資料寫入 Vercel Postgres（出現在聯盟地圖）\n・透過 Vercel API 設定藥局網站環境變數\n・觸發第一次部署（帶入藥師照片、地址、服務項目）\n・通知行銷主管：「XX 藥局已完成設定」\n・給藥師發 Gmail 通知：「您的網站已上線：xxx.tw」'),
]

for num, title_s, desc in onboard_steps:
    heading3(f'{num}｜{title_s}')
    for line in desc.split('\n'):
        if line.strip():
            if line.startswith('・') or line.startswith('4-'):
                bullet(line.lstrip('・').strip(), level=1)
            else:
                body(line)
    spacer()

# ════════════════════════════════════════════════
# 7. 平台功能總覽
# ════════════════════════════════════════════════
heading1('七、平台功能總覽')
spacer()

add_table(
    ['功能模組', '說明', '適用角色'],
    [
        ['今日產文', '選主題 → 一鍵 AI 產文 → 全站草稿', 'admin'],
        ['審稿 / 批准發布', '逐篇確認或退回修改，批准後自動上線', 'admin / super_admin'],
        ['我的文章（藥師）', '查看今日 3 篇草稿、編輯、送審', 'pharmacist'],
        ['藥局管理', '新增 / 編輯藥局資料、藥師帳號、Vercel 設定', 'admin / super_admin'],
        ['邀請藥局', '產生邀請連結，追蹤 onboarding 進度', 'admin / super_admin'],
        ['LINE OA 管理', '設定各藥局 LINE OA 連結，自動注入文章 CTA', 'admin / super_admin'],
        ['見證案例管理', '新增使用者見證（照片 + 短文），同步全站', 'admin / super_admin'],
        ['聯盟地圖', '所有上線藥局顯示於地圖，供顧客查找最近藥局', '前台公開'],
        ['發布記錄', '歷史所有文章，依日期 / 藥局 / 主題篩選', 'admin / super_admin'],
        ['轉化 Dashboard', '各站文章瀏覽量、LINE CTA 點擊數追蹤', 'admin / super_admin'],
    ],
    col_widths=[4, 8, 4]
)

spacer()

# ════════════════════════════════════════════════
# 8. 開發時程
# ════════════════════════════════════════════════
heading1('八、建議開發時程')
spacer()

add_table(
    ['期別', '時程', '內容', '完成標準'],
    [
        ['期一', '第 1-2 週', 'Gmail 登入 + 角色系統\n藥局管理 CRUD\nAI 產文引擎\n基礎審稿流程', '行銷主管可登入、產文、審稿、發布'],
        ['期二', '第 3-4 週', '藥師視角（只看自己文章）\n送審 → 批准流程\n藥局 onboarding 引導頁\n一鍵 Deploy 按鈕', '藥師可完整走完每日作業'],
        ['期三', '第 5-6 週', '見證案例管理\nLINE OA 注入\n聯盟地圖（Google Maps）\n轉化 Dashboard', '完整閉環，行銷主管可看成效數據'],
    ],
    col_widths=[2, 2.5, 7, 5]
)

spacer(2)

# Footer note
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_note = note.add_run('本文件由普健生醫行銷團隊規劃，2026 年 6 月 6 日')
run_note.font.size = Pt(9)
run_note.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
run_note.font.name = '微軟正黑體'

# ── Save ──
out = r'C:\Users\JasonLee\claude_code_projects\CMO\docs\clients\huangweide\普健生醫AI行銷平台規劃書_20260606.docx'
doc.save(out)
print(f'SAVED: {out}')
