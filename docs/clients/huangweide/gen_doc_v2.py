from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

# ── helpers ──
def set_font(run, size=11, bold=False, color=None):
    run.font.name = '微軟正黑體'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    if color: run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_heading('', level=1); p.clear()
    r = p.add_run(text); set_font(r, 17, True, (0x0F,0x76,0x6E))
    return p

def h2(text):
    p = doc.add_heading('', level=2); p.clear()
    r = p.add_run(text); set_font(r, 13, True, (0x05,0x96,0x69))
    return p

def body(text, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text); set_font(r, size)
    return p

def bold_mix(segments, size=11):
    p = doc.add_paragraph()
    for txt, bold in segments:
        r = p.add_run(txt); set_font(r, size, bold)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); set_font(r, 11)
    p.paragraph_format.left_indent = Cm(1 + level*0.8)
    return p

def orange_box(label, content):
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'FFF7ED')
    cell._tc.get_or_add_tcPr().append(shd)
    p = cell.paragraphs[0]
    r1 = p.add_run(label + '　'); set_font(r1, 11, True, (0xC2,0x41,0x0C))
    r2 = p.add_run(content);      set_font(r2, 11, False, (0x9A,0x34,0x12))
    return t

def green_box(content):
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F0FDF4')
    cell._tc.get_or_add_tcPr().append(shd)
    p = cell.paragraphs[0]
    r = p.add_run(content); set_font(r, 11, False, (0x16,0x65,0x34))
    return t

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = h
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'0F766E')
        cell._tc.get_or_add_tcPr().append(shd)
        for r in cell.paragraphs[0].runs: set_font(r,11,True,(0xFF,0xFF,0xFF))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = t.rows[ri+1].cells[ci]; cell.text = str(val)
            if ri%2==0:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F0FDF4')
                cell._tc.get_or_add_tcPr().append(shd)
            for r in cell.paragraphs[0].runs: set_font(r,10)
    if col_widths:
        for i,w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width = Cm(w)
    return t

def sp(n=1):
    for _ in range(n): doc.add_paragraph()

# ══════════════════════════════════════════
# COVER
# ══════════════════════════════════════════
sp()
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('普健生醫\nAI 行銷管理平台\n規格確認書')
set_font(r, 24, True, (0x0F,0x76,0x6E))

sp()
sp2 = doc.add_paragraph(); sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sp2.add_run('藥局聯盟 × AI 產文 × 流量閉環')
set_font(r2, 13, False, (0x64,0x74,0x8B))

sp()
dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = dp.add_run('2026 年 6 月 6 日　　呈：黃維德 藥師')
set_font(r3, 11, False, (0x94,0xA3,0xB8))

doc.add_page_break()

# ══════════════════════════════════════════
# 壹、建置用意（最重要的一頁）
# ══════════════════════════════════════════
h1('壹、為什麼要建置這個平台？')
sp()

body('維德藥師，您好。', 12)
sp()
body(
    '我們計劃為普健生醫建置一個「AI 行銷管理平台」。在談功能之前，我想先說清楚：'
    '我們為什麼要做這件事，以及這件事對普健生醫的意義是什麼。')
sp()

h2('現在遇到的問題')
sp()
bullet('普健生醫的 DP2 外用過敏專利，在市場上幾乎找不到競品，定位極具優勢。')
bullet('但「有好產品」不等於「消費者找得到你」。')
bullet('台灣消費者對保健品的第一個動作，是 Google 搜尋——例如搜尋「過敏 藥局推薦」、「異位性皮膚炎保養」。')
bullet('要在 Google 搜尋上被找到，需要大量的高品質在地文章，持續不斷地更新。')
bullet('若要人工寫作，100 間藥局、每天各 3 篇，每月需要 9,000 篇文章，根本不可能執行。')
sp()

h2('一個您可能已經觀察到的現象')
sp()
body('目前許多公司的行銷部門，已經開始直接使用 AI 工具（如 ChatGPT、Claude）來產出文案和內容。'
     '這本來是好事，但在實際執行上，往往會出現以下情形：')
sp()
add_table(
    ['現象', '對公司的影響'],
    [
        ['行銷人員每天花大量時間下 AI 指令，但主管不清楚他們在做什麼',
         '人力成本增加，工作內容不透明，主管難以評估績效'],
        ['AI 使用費用（Token）持續累積，但看不出與業績的關聯',
         '錢花出去了，不知道帶來多少效果，預算難以控制'],
        ['每個人使用 AI 的方式不同，產出品質參差不齊',
         '品牌形象不一致，文章品質難以把關'],
        ['沒有記錄「哪一篇文章帶來多少流量」，無法優化',
         '行銷工作缺乏可追蹤的成果，無法持續改善'],
    ],
    col_widths=[7.5, 9]
)
sp()
green_box(
    '建置這個平台，就是要解決上述問題。\n\n'
    '平台把 AI 的使用方式「標準化」：行銷人員不再自行操作 AI，'
    '而是在平台上完成固定的工作流程（選主題→產文→審稿→發布）。'
    '每一篇文章、每一次發布，都有完整記錄。\n\n'
    '您（黃維德藥師）在管理後台可以清楚看到：'
    '今天發了幾篇、哪些藥局、哪些主題、文章瀏覽量是多少。'
    '行銷團隊的工作成果，第一次變得「可視化、可追蹤、可管理」。'
)
sp()

h2('這個平台要解決的事')
sp()
green_box(
    '核心概念：讓 AI 做重複性的產文工作，讓行銷人員和藥師專注在確認品質，'
    '讓平台自動把文章推送到每一間藥局的網站——每天只需花 30 分鐘，完成過去需要 10 個人才能做到的事。'
)
sp()

body('具體來說，這個平台做到三件事：')
sp()

add_table(
    ['這個平台做到的事', '帶來的效益'],
    [
        ['AI 每天自動為每間藥局產出 3 篇不同角度的文章',
         '100 間藥局、每月 9,000 篇文章，由 1 位行銷人員即可管理\n大幅降低內容產製成本，規模化不靠人力'],
        ['每篇文章在 Google 被搜尋到後，引導顧客找到離他最近的聯盟藥局',
         '藥局不用自己打廣告，聯盟平台統一導流\n藥局加入動機強：免費獲得客流'],
        ['全部流程在一個網頁管理，不需要任何技術知識',
         '行銷團隊每天 30 分鐘完成全台更新\n藥師只需確認自己的文章，簡單易用'],
    ],
    col_widths=[7, 9.5]
)
sp()

h2('這個平台對普健生醫的戰略意義')
sp()
body('當 100 間藥局的網站都在推廣「DP2 外用過敏專利」，並且每天更新在地化的過敏保健文章，'
     '普健生醫就建立了一個競爭對手幾乎無法複製的護城河：')
sp()
bullet('Google SEO 領先：100 個在地化網站，覆蓋全台各縣市的過敏相關搜尋')
bullet('品牌聲量累積：提到外用過敏，消費者第一個想到的就是普健生醫')
bullet('藥局聯盟黏性：藥局因為獲得客流而留在聯盟，形成長期合作關係')
bullet('擴張速度快：每新增一間藥局，只需填寫資料、點一個按鈕，即可加入平台')
sp()

bold_mix([
    ('一句話總結：', True),
    ('這個平台是普健生醫從「好產品」變成「市場第一品牌」的基礎建設。', False)
], 12)

doc.add_page_break()

# ══════════════════════════════════════════
# 貳、誰會用這個平台
# ══════════════════════════════════════════
h1('貳、誰會用這個平台？')
sp()
body('三種使用者，全部以 Gmail 帳號登入，不需另設密碼。')
sp()

add_table(
    ['使用者', '適用對象', '每天做什麼'],
    [
        ['最高管理員', '黃維德藥師', '查看所有藥局與文章，管理帳號，具備全部權限'],
        ['行銷主管',   '行銷團隊',   '每日一鍵產文、審稿、批准文章上線、管理藥局資料'],
        ['藥師',       '各聯盟藥局藥師', '登入只看到自己藥局今天 3 篇草稿，修改後送出審核'],
    ],
    col_widths=[3.5, 3.5, 9.5]
)
sp()
orange_box('📋 請確認：', '以上三種角色是否符合需求？是否需要新增或調整？')
sp()

# ══════════════════════════════════════════
# 參、每天怎麼運作
# ══════════════════════════════════════════
h1('參、每天是怎麼運作的？')
sp()
h2('行銷主管每日流程')
sp()

add_table(
    ['步驟', '操作', '說明'],
    [
        ['1', '登入平台',   '用 Gmail 登入，看到今日 Dashboard'],
        ['2', '選今日主題', '從四大主題選擇（可多選）：過敏保養 / 葉黃素護眼 / 關節保健 / 體重管理'],
        ['3', '一鍵產文',   '點「開始產文」，AI 幫所有藥局各產 3 篇，每篇角度不同，約 2 分鐘'],
        ['4', '審稿',       '逐篇確認，可修改，或退回給藥師並附留言'],
        ['5', '批准發布',   '批准後，文章即時上到各藥局網站'],
    ],
    col_widths=[1.5, 3.5, 11.5]
)
sp()

h2('藥師每日流程')
sp()

add_table(
    ['步驟', '操作', '說明'],
    [
        ['1', '登入平台',       '用自己的 Gmail 登入，只看到自己藥局的內容'],
        ['2', '查看 3 篇草稿',  'AI 已根據藥師姓名、藥局地區、指定角度自動產出'],
        ['3', '修改內容',       '可直接在編輯器修改，確認符合法規與個人風格'],
        ['4', '送出審核',       '按「送出審核」，行銷主管收到通知'],
        ['5', '收到上線通知',   '主管批准後，文章自動上到藥局網站'],
    ],
    col_widths=[1.5, 3.5, 11.5]
)
sp()
orange_box('📋 請確認：', '藥師送出後需行銷主管批准才上線（藥師無法自行發布），此流程您是否同意？')
sp()

# ══════════════════════════════════════════
# 肆、文章不會重複
# ══════════════════════════════════════════
h1('肆、每間藥局的文章為什麼不會一樣？')
sp()
body('若 100 間藥局的文章內容雷同，Google 會降低所有網站的搜尋排名。'
     '平台在產文時會自動注入每間藥局的專屬資料（藥師姓名、所在縣市、主打服務），'
     '並指定不同的文章角度，確保每篇文章都真正不同。')
sp()

add_table(
    ['分站', '文章角度', '範例標題'],
    [
        ['普健生醫總站',   '學術 / 成分解析', '入秋過敏高峰！DP2 外用專利技術的舒緩機制解析'],
        ['台北大安 長青',  '在地 / 藥師觀察', '台北盆地秋季為何特別容易過敏？陳藥師 10 年觀察'],
        ['新竹東區 長青',  '季節 / 氣候切入', '新竹風城乾燥氣候，過敏族群這樣保養最有效'],
        ['桃園中壢 長青',  '家庭故事敘事',   '一個媽媽的故事：孩子反覆過敏讓全家不安寧'],
        ['台中西屯 長青',  'Q&A 問答格式',   'Q&A｜過敏藥吃了想睡？台中張藥師這樣說'],
    ],
    col_widths=[4, 4, 8.5]
)
sp()
body('每篇文章結尾自動加入「加入 LINE OA 領免費衛教資訊」按鈕，導流到各藥局的 LINE 帳號。')
sp()
orange_box('📋 請確認：', '四大主題（過敏 / 葉黃素 / 關節 / 體重管理）是否為目前主推方向？有無需要調整？')
sp()

# ══════════════════════════════════════════
# 伍、顧客如何找到最近的藥局
# ══════════════════════════════════════════
h1('伍、顧客如何找到最近的聯盟藥局？')
sp()
body('每個藥局的網站底部，都會顯示一張「聯盟藥局地圖」，標示全台所有加入普健聯盟的藥局。')
sp()
body('顧客動線：')
bullet('Google 搜尋「過敏 新竹藥局」→ 落地到對應藥局網站，閱讀文章')
bullet('頁面底部看到「找離你最近的聯盟藥局」地圖')
bullet('點地圖上的藥局 → 直接導航 / 加入 LINE OA 預約諮詢')
sp()
orange_box('📋 請確認：', '顧客進站後，優先引導的行動是？\n□ 加入 LINE OA   □ 電話聯絡   □ 預約現場諮詢   □ 其他：___________')
sp()

# ══════════════════════════════════════════
# 陸、藥局怎麼加入
# ══════════════════════════════════════════
h1('陸、新藥局如何加入聯盟？')
sp()
body('藥局加入時需完成一次性設定，約 20-30 分鐘，全程有圖文引導，不需要技術知識。')
sp()

add_table(
    ['步驟', '誰來做', '內容'],
    [
        ['1  普健發邀請',   '行銷主管', '在平台產生邀請連結，用 LINE 或 Email 傳給藥局'],
        ['2  填寫基本資料', '藥局',     '藥局名稱、地址、電話、藥師姓名、執照、照片、主打服務'],
        ['3  設定登入帳號', '藥局',     '填入藥師 Gmail，之後用此帳號登入平台'],
        ['4  建立網站',     '藥局（平台引導）',
         '① 免費申請 Vercel 帳號\n'
         '② 購買自己的域名，約 NT$700/年（藥局自付）\n'
         '③ 點一下「一鍵建站」，約 2 分鐘完成\n'
         '④ 複製 2 個資料貼回平台（平台有截圖說明哪裡找）'],
        ['5  自動完成',     '系統',
         '驗證成功後：藥局上線、出現在聯盟地圖、通知行銷主管、給藥師發上線通知'],
    ],
    col_widths=[3.5, 3, 10]
)
sp()
orange_box('📋 請確認：', '域名費用（約 NT$700/年）由各藥局自付，是否同意？加盟條件是否需要調整？')
sp()

# ══════════════════════════════════════════
# 柒、開發時程
# ══════════════════════════════════════════
h1('柒、建議開發時程')
sp()

add_table(
    ['期別', '時程', '完成後可以做的事'],
    [
        ['第一期', '第 1–2 週', '行銷主管可登入、產文、審稿、發布，第一批文章上線'],
        ['第二期', '第 3–4 週', '藥師可登入操作，開始邀請第一批藥局試辦'],
        ['第三期', '第 5–6 週', 'LINE OA 注入、聯盟地圖、轉化數據追蹤，完整閉環上線'],
    ],
    col_widths=[2.5, 3, 11]
)
sp()
orange_box('📋 請確認：', '第一期完成後（約 2 週），哪幾間長青藥局適合作為試辦？請提供聯絡人，我們提前準備邀請連結。')
sp()

doc.add_page_break()

# ══════════════════════════════════════════
# 捌、簽核
# ══════════════════════════════════════════
h1('捌、確認簽核')
sp()
body('請您閱讀完畢後，確認以下項目並簽名，我們即可進入開發階段。'
     '如有任何問題，歡迎直接在文件上標注或來電討論。')
sp(2)

chk_t = doc.add_table(rows=5, cols=3); chk_t.style = 'Table Grid'
check_rows = [
    ('確認項目', '確認結果', '備註'),
    ('平台用意與目標',             '□ 同意   □ 需討論', ''),
    ('三種角色設計',               '□ 同意   □ 需調整', ''),
    ('藥師送審→主管批准流程',      '□ 同意   □ 需調整', ''),
    ('域名費用由藥局自付',         '□ 同意   □ 需討論', ''),
]
for ri, (a, b, c) in enumerate(check_rows):
    for ci, val in enumerate([a, b, c]):
        cell = chk_t.rows[ri].cells[ci]; cell.text = val
        for r in cell.paragraphs[0].runs:
            set_font(r, 11, ri==0)
        if ri == 0:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'0F766E')
            cell._tc.get_or_add_tcPr().append(shd)
            for r in cell.paragraphs[0].runs:
                r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
for row in chk_t.rows:
    row.cells[0].width = Cm(5.5)
    row.cells[1].width = Cm(4.5)
    row.cells[2].width = Cm(6.5)

sp(2)
sign = doc.add_paragraph(); sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rs = sign.add_run('黃維德 藥師 簽名：_____________________　　日期：__________')
set_font(rs, 11)

sp(2)
note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = note.add_run('本文件由普健生醫行銷團隊製作　2026 年 6 月 6 日')
set_font(rn, 9, False, (0x94,0xA3,0xB8))

# ── Save ──
out = r'C:\Users\JasonLee\claude_code_projects\CMO\docs\clients\huangweide\普健生醫平台規格確認書_v2_20260606.docx'
doc.save(out)
print(f'SAVED: {out}')
