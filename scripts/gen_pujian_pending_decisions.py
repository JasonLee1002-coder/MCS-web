# -*- coding: utf-8 -*-
"""AI-CMO 普健生醫案待確認事項彙整"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GOLD = RGBColor(0xB8, 0x68, 0x0A)
RED = RGBColor(0x8B, 0x1A, 0x1A)
GREEN = RGBColor(0x1E, 0x7A, 0x3C)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHTBLUE = "EBF3FB"
LIGHTRED = "FDEDEC"
LIGHTGOLD = "FFF3CD"
LIGHTGREEN = "E2EFDA"

def set_font(run, cn="Noto Sans TC", en="Calibri", size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    if color:
        run.font.color.rgb = color

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, "Noto Serif TC", "Cambria", 15, True, NAVY)
    return p

def add_heading2(doc, text, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, "Noto Sans TC", "Calibri", 13, True, color)
    return p

def add_body(doc, text, indent=0, color=None, bold=False, size=12):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(hdr[i], "1F4E79")
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_font(r, size=10)
            if ridx % 2 == 0:
                shade_cell(cells[i], LIGHTBLUE)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_box(doc, title, text, bgcolor, titlecolor):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    shade_cell(cell, bgcolor)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_font(r, size=11.5, bold=True, color=titlecolor)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(text)
    set_font(r2, size=10.5)
    doc.add_paragraph()

doc = Document()
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# ── 封面 ──
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(70)
r = cover.add_run("普健生醫 × 森林藥局聯盟 × 黃維德.com")
set_font(r, "Noto Serif TC", "Cambria", 15, True, NAVY)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(18)
r = title.add_run("待確認事項彙整")
set_font(r, "Noto Serif TC", "Cambria", 24, True, NAVY)
r2 = title.add_run("\n2026-08-03 ～ 08-04 盤點")
set_font(r2, "Noto Serif TC", "Cambria", 16, True, NAVY)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(26)
r = sub.add_run("撰寫者：AI CMO")
set_font(r, size=13, color=GOLD, bold=True)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("版本 v1.0　2026-08-04")
set_font(r, size=11, color=GREY)

doc.add_page_break()

# ── 摘要 ──
add_heading1(doc, "總結：一句話")
add_box(doc, "多數缺口不卡工，但6件事沒人拍板前不能寫進對外文案",
    "本次盤點過敏文章庫（345篇）、AI藥事Agent saleskit、森林藥局聯盟合作藥局說帖、馬偕合作簡報等既有素材後，發現大部分內容工作（VetaCode頁面、引用格式優化、CTA連結修正）現在就能直接做，已排入執行；但有6項涉及商業條款、合作關係真實性、對外用語尺度的事，寫錯就是捏造或洩露未定案資訊，需要黃維德/詩娟/Jason/財務先拍板。",
    LIGHTGOLD, GOLD)

# ── 一、已解決 ──
add_heading1(doc, "一、已解決（僅供留存紀錄，不用再確認）")
add_table(doc, ["項目", "結論", "依據"], [
    ["加盟優惠家數版本之爭（前500/前50/前100/前100家四個版本互相打架）", "定案為「前100家 NT$2,000」", "森林藥局聯盟_合作藥局說帖.pdf（2026.07正式對外發送版）"],
    ["季昭華教授頭銜語序（「前台大動物醫院院長」易被誤讀成「前台大」的動物醫院院長）", "全站統一改為「台大動物醫院前院長」，已上線", "huangweide-seo commit 77460f7"],
], col_widths=[7, 5, 5])
add_body(doc, "建議：08-02馬偕合作簡報內部P12「前500家」與P13「前50家」的不一致，回頭修正對齊「前100家」這個對外實際版本，避免業務同時講出三種說法。", color=GREY, size=10.5)

# ── 二、商業條款類 ──
add_heading1(doc, "二、待確認 — 商業條款類")
add_table(doc, ["項目", "現況/衝突", "需要誰決定"], [
    ["馬偕用語尺度", "簡報內部用「技術主導／醫學背書／實證背書」，公開網頁曝險較高", "黃維德 ＋ 法務"],
    ["研發投入金額", "市場簡報寫「1.5億」，會議中曾提「6億」，兩者未對齊", "財務（富哥）"],
    ["GBP整合（4間代管藥局Google商家檔案）", "Jason已拍板要做，但程式碼完全沒有，尚未排入開發排程", "排優先序"],
], col_widths=[5, 8, 4])
add_body(doc, "這三項在拍板前，網站文案一律維持最保守寫法（不提具體金額、不用強措辭），已同步套用到現有內容。", color=GREY, size=10.5)

# ── 三、合作關係狀態 ──
add_heading1(doc, "三、待確認 — 這個合作是真的還是評估中")
add_box(doc, "AI藥事Agent（台灣生醫大數據科技股份有限公司 × 高通）",
    "這份saleskit在pujian-platform repo裡被發現，內容是藥物交互作用/處方審核的臨床AI工具，瞄準「藥師端多處方交叉審方」——如果普健已經談定要整合進森林藥局聯盟藥師Portal，這是招商的第5項賣點（比GEO/LINE/230萬保戶更硬核，是屈臣氏/康是美拿不到的東西）。但目前完全不確定這是(a)已談定要整合(b)還在評估(c)純參考資料。三選一決定前，網站/招商簡報都不能提「加入聯盟可用AI藥事助手」，否則等於捏造不存在的合作關係。",
    LIGHTRED, RED)

# ── 四、待補資料 ──
add_heading1(doc, "四、待補資料才能落地")
add_box(doc, "台灣人壽230萬保戶＋500大企業專屬導流頁（/taiwanlife）",
    "馬偕簡報提到這個管道，但網站上完全沒有對應著陸頁。這批人流不是自己搜尋過敏才進站的（跟SEO訪客邏輯不同），需要普健那邊（黃維德/詩娟）提供具體活動內容、保戶禮遇門檻、活動期限，才能寫出對應的落地頁文案，不是CMO可以自己編。",
    LIGHTBLUE, NAVY)

# ── 五、待看文件 ──
add_heading1(doc, "五、待Jason看的既有文件")
add_body(doc, "• 「森林藥局聯盟_線上加盟頁待決策事項_20260730.docx」（docs/內）——這份文件尚未讀取消化，下次進來優先看，可能會改變上述幾項的答案。", indent=0.4)

# ── 六、內容方向確認 ──
add_heading1(doc, "六、內容方向確認（jason-seo正在寫的稿）")
add_box(doc, "「生醫保健品要怎麼賣進自助通路？」文章框架",
    "框架整體（避免誤導的五條寫法規則、國際案例佐證）已通過戰略檢視，可以繼續寫。但建議補一段「自助設備是藥局觸角的延伸、不是取代藥局」的定位句——不然森林藥局聯盟正在同步招商的潛在合作藥局，看到「生醫產品要透過自助販賣機賣」可能誤讀成「普健要繞過藥局用機器賣貨」，跟招商說帖的「藥局是主要通路、不會被繞過」打對台，殺傷力比誤導一般消費者更直接。請確認是否同意這個補充方向。",
    LIGHTGOLD, GOLD)

# ── 七、附錄 ──
add_heading1(doc, "七、附錄：不用等確認、已經在做的項目")
add_table(doc, ["項目", "狀態"], [
    ["季昭華教授頭銜語序修正（huangweide-seo）", "✅ 已上線"],
    ["pharmacy-blogs（森林藥局聯盟）新增 /vetacode 頁面", "🔄 執行中"],
    ["huangweide-seo文章補引用來源格式（GEO/E-E-A-T優化）", "🔄 執行中"],
    ["345篇過敏文章排程匯入pharmacy-blogs（在地化內容庫）", "⏳ 待盤點對應工作表欄位後執行"],
], col_widths=[10, 6])

doc.save(r"H:\共用雲端硬碟\2026-銓幻元共用雲端硬碟\2026業務\黃維德-普健生醫\AI-CMO_普健生醫案待確認事項彙整_20260804.docx")
print("saved")
