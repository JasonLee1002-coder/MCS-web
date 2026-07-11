# -*- coding: utf-8 -*-
"""Generate Yo-Kai Express infringement-research report (Word)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GOLD = RGBColor(0xB8, 0x68, 0x0A)
RED = RGBColor(0x8B, 0x1A, 0x1A)
DARKGRAY = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x55, 0x55, 0x55)

SERIF_TC = "Noto Serif TC"
SANS_TC = "Noto Sans TC"
LATIN_TITLE = "Cambria"
LATIN_BODY = "Calibri"


def set_font(run, latin=LATIN_BODY, cjk=SANS_TC, size=12, bold=False, color=None):
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cjk)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    if level == 1:
        set_font(r, LATIN_TITLE, SERIF_TC, 16, True, NAVY)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '1F4E79')
        pBdr.append(left)
        pPr.append(pBdr)
        pPr_ind = p.paragraph_format
        pPr_ind.left_indent = Cm(0.3)
    elif level == 2:
        set_font(r, LATIN_TITLE, SERIF_TC, 13, True, GOLD)
    return p


def add_para(doc, text, size=12, indent=0, color=DARKGRAY, bold=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, LATIN_BODY, SANS_TC, size, bold, color)
    return p


def add_bullet(doc, text, size=11.5):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    set_font(r, LATIN_BODY, SANS_TC, size, False, DARKGRAY)
    return p


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        set_font(r, LATIN_BODY, SANS_TC, 11, True, RGBColor(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(hdr_cells[i], '1F4E79')
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for cidx, val in enumerate(row):
            cells[cidx].text = ''
            p = cells[cidx].paragraphs[0]
            r = p.add_run(str(val))
            set_font(r, LATIN_BODY, SANS_TC, 10.5, False, DARKGRAY)
            if ridx % 2 == 0:
                shade_cell(cells[cidx], 'EBF3FB')
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_source_line(doc, title, url):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run("• " + title + "  ")
    set_font(r1, LATIN_BODY, SANS_TC, 10, False, DARKGRAY)
    r2 = p.add_run(url)
    set_font(r2, LATIN_BODY, SANS_TC, 9.5, False, RGBColor(0x1F, 0x4E, 0x79))


doc = Document()

# 頁面邊界
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# ---------- 封面 ----------
cover_table = doc.add_table(rows=1, cols=1)
cover_table.rows[0].cells[0].width = Cm(17)
cell = cover_table.rows[0].cells[0]
shade_cell(cell, '1F4E79')
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after = Pt(10)
r = p.add_run("龍雲數位整合股份有限公司")
set_font(r, LATIN_TITLE, SERIF_TC, 14, True, RGBColor(0xFF, 0xFF, 0xFF))

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(6)
p2.paragraph_format.space_after = Pt(40)
r2 = p2.add_run("Yo-Kai Express 新聞彙整暨市場擺放點位調查報告")
set_font(r2, LATIN_TITLE, SERIF_TC, 20, True, RGBColor(0xFF, 0xFF, 0xFF))

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("侵權訴訟前置資料彙整 ｜ AI CMO 情報彙整")
set_font(r3, LATIN_BODY, SANS_TC, 13, True, GOLD)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
today = "2026年7月10日（v5：2026年7月11日更新）"
r4 = p4.add_run(f"製作日期：{today}｜密級：內部與委任律師使用｜撰寫者：AI CMO")
set_font(r4, LATIN_BODY, SANS_TC, 10.5, False, GRAY)

doc.add_paragraph()
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run("＊本報告彙整自公開網路新聞、公司官網、法律資料庫（Law360）、商標／專利檢索資料庫（USPTO/Justia）等公開來源，")
set_font(r5, LATIN_BODY, SANS_TC, 9.5, False, GRAY)
p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
r6 = p6.add_run("僅供內部評估侵權訴訟策略參考，正式提告前請律師團隊核實各項原始文件與時效。")
set_font(r6, LATIN_BODY, SANS_TC, 9.5, False, GRAY)

doc.add_page_break()

# ---------- 目錄式摘要 ----------
add_heading(doc, "摘要重點", level=1)
add_para(doc, "本報告針對「Yo-Kai Express, Inc.」（美國矽谷自動化拉麵∕食品販賣機新創公司）進行公開資訊蒐集，並納入顧總（顧祖欣先生，樂活速食坊股份有限公司總經理）提供之自有技術資料，內容分為七大部分：（一）Yo-Kai Express公司背景與重要新聞時間軸；（二）與本案關聯性最高的智慧財產權爭議──Noodle Time Holdings, Ltd. 對 Yo-Kai Express 提起之商業秘密訴訟與 YO-KAI 商標歸屬爭議；（三）針對 Yo-Kai Express 對外宣稱「80件專利、38件已核准」進行真偽查證與6項疑點分析，並專節查證其在台灣（子公司「優豈股份有限公司」）的專利佈局概況；（四）Yo-Kai Express 目前於美國、日本、台灣、南韓等地之實際擺放點位彙整；（五）顧總一方「樂活速食坊／勝十蘭」自有技術研發歷程、AI蒸汽烹調機規格與台灣專利申請號（113108458、113202327）；（六）雙方在台灣市場之時序與技術比較，作為本案侵權主張之核心論證框架。", size=12)

add_heading(doc, "一、公司背景與重要新聞彙總", level=1)

add_heading(doc, "（一）公司概況", level=2)
add_para(doc, "Yo-Kai Express, Inc. 由 Andy Lin（林志鴻，具超過12年半導體產業背景，曾任職日月光）於 2016 年在美國矽谷（加州 Sunnyvale／Hayward 一帶）創立，主打「自動化拉麵∕熟食販賣機」，將半導體產線的精密控制與標準化思維應用於食品加熱：食材以專屬配方預煮後急速冷凍，出餐時以近300°C高溫在45～90秒內完成復熱，並與多位米其林主廚合作開發菜單。公司在北加州、日本、台灣均設有據點，官網並宣稱與多家美國財星500大企業合作。台灣子公司登記名稱為「優豈股份有限公司」，登記地址為台北市松山區南京東路四段2號；創辦人林志鴻本人即為台灣籍工程師。", indent=0.3)

add_heading(doc, "（二）融資與投資人", level=2)
add_para(doc, "累計募資約 1,900萬美元，投資人包括日本一風堂（Ippudo）、韓國不倒翁食品集團 Pulmuone、日本煙草集團（JT，經 Chikaranomoto Holdings 参與）及 Plug and Play 等。2023～2025年間市場消息指出公司規劃於2025年底前於美國那斯達克（Nasdaq）掛牌，惟截至本報告撰寫時（2026年7月）尚未查得掛牌完成之正式新聞。", indent=0.3)

add_heading(doc, "（三）重要新聞時間軸", level=2)
news_rows = [
    ["2016", "Andy Lin 於矽谷創立 Yo-Kai Express，開發自動拉麵販賣機"],
    ["2021.01", "與 SFO（舊金山國際機場）合作，於航廈設置自動拉麵機；同期於安大略國際機場（Ontario, CA）規劃設置4台機器"],
    ["2021.03～04", "【台灣】創辦人林志鴻對外宣布計畫將「無人自助拉麵機」引進台灣市場"],
    ["2021.05～06", "【台灣】首度正式落地：因應COVID-19疫情，將機台設置於衛生福利部桃園醫院、亞東紀念醫院，免費提供醫護人員近1.1萬份餐點，為Yo-Kai Express在台灣的最早部署據點"],
    ["2021.11", "Noodle Time Holdings, Ltd.（香港公司）於加州北區聯邦法院對 Yo-Kai Express 提起商業秘密訴訟（案號 5:21-cv-09263）"],
    ["2022.08", "承審法官 Susan van Keulen 裁定將本案送交仲裁（arbitration）"],
    ["2023.02", "與韓國 Pulmuone 合作，推出「Robot Chef 啜啜盒（Chulchul Box）」，進駐首爾大學醫院、首爾市警察廳等機構"],
    ["2023", "於 CES 2023 展出桌上型迷你機種 YKE Desktop"],
    ["2024.03", "與 SoftBank Robotics 合作，由 SoftBank Robotics 接手日本市場自動烹調機器人「CHEFFY」之經銷業務"],
    ["2024", "於 CES 2024 展出自動珍珠奶茶調製機器人，跨足手搖飲市場"],
    ["2024（確切月份未查得）", "【台灣】台北捷運西門站（1、2號出口）成為北捷系統第一個進駐無人拉麵機＋無人手搖飲機之捷運站，與「一風堂」聯名推出豚骨拉麵、麻辣豚骨拉麵"],
    ["2024～2025間（確切時間未查得）", "【台灣】機台進駐桃園國際機場第1航廈A8管制區，24小時供應，旅客反映常見排隊"],
    ["2025", "傳出擴大進駐明尼亞波里－聖保羅國際機場（MSP）第1、2航廈，共3台機器；市場消息稱規劃年底前於Nasdaq掛牌（尚待證實）"],
    ["2026.01", "執行長 Andy Lin 於 NRF 2026（美國零售業年度大展）接受紐約證交所（NYSE TV）專訪，介紹展出之新品項"],
    ["2026.01.21", "與泰國國民泡麵品牌「MAMA」跨國聯名，於台北捷運西門站B1、3號出口舉辦快閃活動，推出4種口味「現煮MAMA麵」，宣稱獲輝達（NVIDIA）Inception Program支持，導入AI模擬（COSMOS虛擬世界）技術校正各地機台之機械手臂運作"],
]
make_table(doc, ["時間", "事件摘要"], news_rows, col_widths=[3, 13.5])

add_para(doc, "")

add_heading(doc, "二、智慧財產權爭議（訴訟關鍵參考）", level=1)

add_heading(doc, "（一）Noodle Time Holdings, Ltd. v. Yo-Kai Express, Inc. 商業秘密訴訟", level=2)
case_rows = [
    ["案號", "5:21-cv-09263"],
    ["管轄法院", "美國加州北區聯邦地方法院（U.S. District Court, N.D. California）"],
    ["承審法官", "Susan van Keulen"],
    ["原告", "Noodle Time Holdings, Ltd.（香港註冊公司）"],
    ["被告", "Yo-Kai Express, Inc.（原告在美國之代理商／經銷夥伴）"],
    ["起訴日期", "2021年11月30日"],
    ["爭議性質", "商業秘密（trade secrets）侵占爭議，涉及拉麵∕熟食自動販賣機之機台技術與營運機密"],
    ["訴訟代理", "被告方代理律師事務所：Kirkland & Ellis"],
    ["現況", "2022年8月，法官裁定將本案送交仲裁（即便被告一度嘗試撤回仲裁聲請亦未獲准）；本報告撰寫時未查得仲裁最終結果之公開資訊"],
]
make_table(doc, ["項目", "內容"], case_rows, col_widths=[3.5, 13])

add_para(doc, "")
add_heading(doc, "（二）YO-KAI 商標歸屬爭議", level=2)
add_para(doc, "美國商標資料庫（Trademark Elite／USPTO 序號 87346086）顯示，「YO-KAI」商標之註冊權利人為 Noodle Time Holdings Limited（非 Yo-Kai Express, Inc. 本身），商標涵蓋類別包括：冷藏販賣機、販賣機組成之自助服務機台、供應熟食（含麵類與便當）之販賣機。此點顯示 Yo-Kai Express, Inc. 於美國市場實際使用「Yo-Kai」品牌名稱經營自動販賣機事業，但相關商標卻登記於原告 Noodle Time Holdings 名下，與前述商業秘密訴訟的「原廠授權商／代理商反目」情節相互印證。", indent=0.3)

add_para(doc, "")
add_heading(doc, "（三）專利佈局", level=2)
add_para(doc, "依 Justia Patents 檢索，Yo-Kai Express Inc. 名下全球累計申請專利約80件，其中38件已獲核准，主要集中於「食品加熱裝置（Food Warming Device）」「食品模具／自動化烹調流程」等機構設計與製程專利（如 WO2021050686A1）。此專利佈局多數為 Yo-Kai Express 自行申請並持有，與前述「YO-KAI」商標由第三方（Noodle Time Holdings）持有的情況並存，顯示品牌／商標權與部分核心技術專利權可能分屬不同法律主體，此點在評估侵權主張與可能反訴風險時須特別留意。", indent=0.3)

add_para(doc, "")
add_heading(doc, "（四）對本案訴訟策略之初步觀察", level=2)
add_bullet(doc, "Yo-Kai Express 過去已有與原授權方／代理夥伴（Noodle Time Holdings）因商業秘密與品牌歸屬產生訴訟之前例，顯示該公司在商標與技術授權關係上曾發生糾紛，此類先例可作為佐證其商業模式存在「合作方掌握核心資產、雙方事後對簿公堂」風險特徵之參考。")
add_bullet(doc, "若顧總欲主張之侵權態樣與「技術/配方/機台設計遭合作方擅自使用或申請專利」相關，建議進一步調閱 Noodle Time Holdings 一案完整卷宗（可透過 PACER 調閱），比對雙方起訴狀之具體侵權事實描述，評估是否有相似情節可援引。")
add_bullet(doc, "「YO-KAI」商標實際登記人為第三方之事實，意味著若考慮以商標侵權角度提告，需先確認顧總所爭執之品牌／技術與此商標及既有訴訟標的之關聯性與時間先後，避免與既有訴訟主體產生管轄或利害關係衝突。")
add_bullet(doc, "建議正式提告前，委任律師團隊透過 PACER／USPTO TSDR／WIPO Patentscope 等官方資料庫調閱一手文件（起訴狀、仲裁裁決、專利說明書全文），本報告僅整理自公開新聞與二手資料庫摘要，不可直接作為法律文件引用。")

add_para(doc, "")
add_heading(doc, "三、專利真偽與疑點深度分析", level=1)
add_para(doc, "Yo-Kai Express 對外（含官網、募資資料、媒體報導）多次宣稱「全球累計約80件專利、其中38件已獲核准」。本節就公開可查證之專利資料庫（Justia Patents、GreyB 專利分析、ipqwery/WIPO 資料庫）逐一核實，區分「確實存在且已核准」與「僅為申請中、含金量較低」兩類，並整理出6項疑點供律師團隊查證方向參考。", size=12)

add_heading(doc, "（一）查證屬實部分——並非空穴來風", level=2)
add_bullet(doc, "美國設計專利（Design Patent）D985424：申請日2020年11月27日，核准日2023年5月9日，發明人 Lin, Chih Hung（林志鴻），現由 Yo-Kai Express Inc. 持有，可於 USPTO 官網查證，確實存在。")
add_bullet(doc, "美國設計專利 D1003643：申請日同為2020年11月27日，核准日2023年11月7日，發明人同為林志鴻，同樣可查證屬實。")
add_bullet(doc, "多件 WIPO／PCT 國際專利申請案（詳下表）確實已公開，說明書全文可於 WIPO Patentscope／Sumobrain 等資料庫查閱，內容並非捏造。")

pat_rows = [
    ["FOOD PROCESSING DEVICE CAPABLE OF FOOD HEATING AND SELF-CLEANING", "US2024029681", "2024/238785", "2024-05-16", "林志鴻", "PCT國際申請（尚待各國家階段審查）"],
    ["VENDING MACHINE FOR HEATING FOOD AND ITS OPERATING METHOD", "US2024029679", "2024/238784", "2024-05-16", "林志鴻", "PCT國際申請（尚待各國家階段審查）"],
    ["QUICK-FROZEN NOODLE UNFREEZING, COOKING AND VENDING MACHINE", "US2023068734", "2023/205819", "2023-06-20", "Xu, Xiong；林志鴻", "PCT國際申請（尚待各國家階段審查）"],
    ["FOOD WARMING DEVICE", "US2020050139", "2021/050686", "2020-09-10", "林志鴻", "PCT國際申請（尚待各國家階段審查）"],
    ["FOOD MOLD AND METHOD FOR FOOD PROCESSING USING THE SAME", "US2020047002", "2021/041118", "2020-08-19", "林志鴻", "PCT國際申請（尚待各國家階段審查）"],
    ["FOOD HEATING DEVICE AND FOOD HEATING METHOD USING THE SAME", "US2020046977", "2021/034921", "2020-08-19", "林志鴻", "PCT國際申請（尚待各國家階段審查）"],
    ["FOOD VENDING MACHINE AND FOOD VENDING METHOD USING THE SAME", "US2020043028", "2021/016326", "2020-07-22", "林志鴻", "PCT國際申請（尚待各國家階段審查）"],
    ["【設計專利／已核准】二輪把持結構造型", "D985424", "—", "2020-11-27（核准2023-05-09）", "林志鴻", "美國設計專利，已核准，僅保護外觀造型"],
    ["【設計專利／已核准】相關造型變化", "D1003643", "—", "2020-11-27（核准2023-11-07）", "林志鴻", "美國設計專利，已核准，僅保護外觀造型"],
]
make_table(doc, ["專利名稱／類型", "申請案號", "公開號", "申請日期", "發明人", "查證後狀態"], pat_rows, col_widths=[5, 2.8, 2.2, 2.6, 1.8, 3.1])
add_para(doc, "資料來源：ipqwery（WIPO資料庫整理）、Justia Patents、GreyB 專利分析。", size=9.5, color=GRAY)

add_para(doc, "")
add_heading(doc, "（二）六大疑點——建議律師團隊逐項查證", level=2)

add_para(doc, "疑點1：「專利申請」與「已獲核准專利」混為一談", size=12, bold=True, color=RED)
add_para(doc, "上表7件核心功能性技術（食品加熱、解凍、烹調、販賣機構造）目前查證到的均為「PCT國際專利申請案」（WO/US20xx公開號），狀態僅為「In Force」（申請程序仍存續中），並非「已獲核准的專利權」。若對外文宣或募資簡報將這些「申請中」案件包裝成「已取得專利保護」，即屬誇大或誤導，值得比對其原始對外用語與時間點。", indent=0.3)

add_para(doc, "疑點2：核心功能技術於美國本土迄今查無「發明專利」核准紀錄", size=12, bold=True, color=RED)
add_para(doc, "依第三方專利分析機構 GreyB 統計，Yo-Kai Express 在美國直接提出（不含設計專利與PCT案）的「發明專利（utility patent）」申請共4件，核准件數為0件，核准率0%，代理事務所為 Birch Stewart Kolasch & Birch。換言之，目前查證所得，Yo-Kai Express 在美國真正「已核准」的專利權僅限於2件「外觀設計專利」（D985424、D1003643）——這類專利只保護機台或零件的「外觀造型」，並不保護其加熱、解凍、烹調方法等核心功能性技術本身。若對造以「我們有專利」對外主張技術排他權，其法律保護範圍恐遠小於其宣稱之技術核心。", indent=0.3)

add_para(doc, "疑點3：「38件已核准」高度集中於審查門檻較低法域，強度可能被誇大", size=12, bold=True, color=RED)
add_para(doc, "官方／媒體宣稱之38件已核准專利，地域分布集中在台灣、紐西蘭、澳洲、新加坡、歐洲、日本、加拿大、中國等地，美國僅約6件（含2件設計專利）。其中台灣與中國均設有「新型專利」制度，僅做「形式審查」（書面格式齊備即可）、不做「新穎性、進步性」之實質審查，核准速度快、門檻低，但法律效力遠弱於「發明專利」，且極易被第三人提出舉發（invalidation）而撤銷。建議查證這38件「已核准」中，究竟有多少屬於新型專利、多少屬於經實質審查之發明專利，才能判斷其專利含金量與可對抗性。", indent=0.3)

add_para(doc, "疑點4：發明人高度集中一人，與宣稱之研發團隊規模不成比例", size=12, bold=True, color=RED)
add_para(doc, "目前查得之7件核心技術專利申請，發明人幾乎全數僅列創辦人「Lin, Chih Hung」（林志鴻）一人，僅1件與Xu, Xiong共同具名。惟公司對外宣稱擁有工程團隊、機台配備約300個感測器、掌握近30道菜色的「黃金參數」配方，研發規模與專利具名發明人數量明顯不成比例。若日後查證顧總（或其代表之一方）曾實際參與相關技術之共同研發、卻未被列名為共同發明人，此即構成美國專利法上可獨立主張「發明人不實（inventorship defect）」之事由，可作為請求專利無效或不可執行的重要武器，建議列為優先查證方向。", indent=0.3)

add_para(doc, "疑點5：專利申請時間點與 Noodle Time Holdings 商業秘密訴訟時間點高度接近，疑似搶先申請", size=12, bold=True, color=RED)
add_para(doc, "上表核心專利之申請日集中於2020年7～9月間（Food Vending Machine 2020/7/22、Food Mold與Food Heating Device均為2020/8/19、Food Warming Device 2020/9/10），而 Noodle Time Holdings 對 Yo-Kai Express 提起商業秘密訴訟是在其後約一年（2021年11月）。若顧總或對造一方能舉證這些2020年申請的技術內容，實際上源自雙方合作關係中，由對造（原廠／授權方）在更早之前就已經開發、掌握並提供予 Yo-Kai Express 之機密配方或機構設計，則存在「明知他人既有技術／商業秘密，搶先以自己名義申請專利」之重大疑慮，此點與既有 Noodle Time Holdings 一案的訴訟主張高度呼應，建議調閱該案起訴狀，逐條比對雙方陳述之技術移轉時間軸。", indent=0.3)

add_para(doc, "疑點6：品牌（商標）與技術（專利）分屬不同法律主體，權利結構本身自相矛盾", size=12, bold=True, color=RED)
add_para(doc, "如前節所述，「YO-KAI」美國商標登記於 Noodle Time Holdings Limited 名下，而功能性專利卻登記在 Yo-Kai Express Inc.（及創辦人林志鴻個人）名下，形成「品牌歸屬第三方、核心技術專利歸屬自己」的分裂結構。這種安排若非雙方原始合約刻意約定之權利劃分，即可能是 Yo-Kai Express 於代理／合作關係存續期間，將對造提供或雙方共有的技術，片面以自己名義申請專利、卻未同步取得或處理品牌歸屬，形成「兩頭得利」的疑慮——建議作為訴訟主張中，質疑其專利申請正當性與誠信原則（inequitable conduct）的重要佐證脈絡。", indent=0.3)

add_para(doc, "")
add_heading(doc, "（三）台灣專利概況（重點查證）", level=2)
add_para(doc, "Yo-Kai Express 在台灣係以子公司「優豈股份有限公司」註冊經營（登記地址：台北市松山區南京東路四段2號），創辦人林志鴻本人即為台灣籍工程師（曾任職日月光12年）。依第三方專利分析機構 GreyB 的全球專利地域分布統計，「台灣」是 Yo-Kai Express／優豈全球專利佈局中，單一法域申請數量最多的地區，共8件，領先美國（6件，含2件設計專利）、中國（6件）、日本（6件）、歐洲（7件）等其他法域。", indent=0.3)
add_para(doc, "惟本次查證受限於公開網路二手資料庫（Justia／GreyB／ipqwery）之揭露深度，僅能確認「台灣為申請量最大單一法域」此一總量數據，尚無法透過公開網路搜尋直接取得這8件台灣專利／申請案的個別案號、專利類型（發明專利／新型專利／設計專利）、審定核准狀態與請求項內容。原因是台灣智慧財產局「專利公開資訊查詢系統」與「全球專利檢索系統（GPSS）」均為動態網頁（需輸入條件送出查詢），無法透過一般網路搜尋引擎索引或靜態網頁擷取工具取得結果。", indent=0.3)
add_para(doc, "建議下一步（可請律師或專利師執行）：", size=12, indent=0.3)
add_bullet(doc, "直接以申請人「優豈股份有限公司」或發明人「林志鴻」為關鍵字，登入台灣智慧財產局 GPSS 全球專利檢索系統（https://gpss.tipo.gov.tw/）或「專利公開資訊查詢」系統（https://tiponet.tipo.gov.tw/S092_OUT/out）查詢完整清單。")
add_bullet(doc, "查得案號後，務必區分是「發明專利」（實質審查，保護力強）或「新型專利」（僅形式審查，保護力弱、易被舉發撤銷）——這項區分對評估其台灣專利之可對抗性至關重要，與本節前述美國/中國法域的疑慮邏輯一致。")
add_bullet(doc, "若顧總／漢典食品之 AI BAR 產品與 Yo-Kai Express 台灣機台在功能設計上有相似之處，應請專利師針對這8件台灣專利逐案比對請求項範圍（claim scope），確認己方產品是否落入其專利權範圍、或反向確認對方專利是否有得舉發無效之前案（prior art）證據。")

add_para(doc, "")
add_heading(doc, "（四）初步結論", level=2)
add_para(doc, "Yo-Kai Express「擁有專利」本身並非虛假宣稱——確有2件美國設計專利已實際核准、亦有多件國際專利申請案真實存在、可公開查閱。但「專利數量」與「專利含金量」需要拆開檢視：目前查證所得，其核心自動解凍、加熱、烹調技術，在美國本土尚未查到經實質審查、真正核准的發明專利；官方宣稱之「38件已核准」，很可能大量來自審查門檻較低的新型專利或設計專利，法律強度與外界（含潛在投資人、消費者）認知恐有落差。此外，發明人具名集中一人、專利申請時間點與商業秘密訴訟時間點高度接近、品牌與技術權利分屬不同主體等三項疑點，均指向「其專利佈局的正當性與完整性，有進一步查證與挑戰的空間」。", indent=0.3)
add_para(doc, "⚠️ 提醒：以上分析均基於 Justia／GreyB／ipqwery 等第三方專利資料庫之公開摘要整理，並非直接調閱 USPTO Patent Center、WIPO Patentscope、TIPO 全球專利檢索系統之原始檔案全文。正式提告或答辯前，務必請委任律師／專利代理人直接調閱各官方資料庫原始文件（含專利說明書全文、申請歷程檔案 file history、任何舉發或訴願紀錄），逐一核實本報告所列各項疑點是否成立，本報告僅作為研究方向與訴訟策略之初步參考，不構成法律意見。", size=10.5, color=GRAY)

doc.add_page_break()

add_heading(doc, "四、目前擺放點位彙整", level=1)
add_para(doc, "Yo-Kai Express 官網宣稱其機台目前已進駐全球超過150個據點，涵蓋機場、飯店、醫院、企業辦公室、大學校園、會議中心、製造廠與住宅社區等場域類型。以下為公開新聞與旅遊評論網站（Yelp／TripAdvisor／Apple地圖等）可查證之具體據點：", size=12)

loc_rows = [
    ["美國", "舊金山國際機場（SFO）", "第3航廈 E4登機門附近；另於機場旁 Aloft San Francisco Airport 飯店大廳設有一台", "機場／飯店"],
    ["美國", "安大略國際機場（Ontario Intl. Airport, CA）", "第4航廈407登機門、第2航廈212登機門，共規劃4台機器，24小時營運", "機場"],
    ["美國", "明尼亞波里－聖保羅國際機場（MSP）", "第1、2航廈共3台機器，24小時供應熱拉麵", "機場"],
    ["美國", "德州休士頓", "9889 Bellaire Blvd（獨立門市型據點）", "門市／熟食販賣"],
    ["美國", "紐約法拉盛（Flushing, NY）", "136-21 Roosevelt Ave（獨立門市型據點）", "門市"],
    ["美國", "加州太浩湖（Tahoe City, CA）", "1960 Squaw Valley Rd（滑雪度假村商圈）", "度假村／休閒場域"],
    ["日本", "羽田機場（Haneda Airport）", "第2航廈", "機場"],
    ["日本", "東京", "JR東日本東京車站、上野車站新幹線閘門內、首都高速灣岸線芝浦停車區（Shibaura PA）", "車站／高速公路服務區"],
    ["台灣", "桃園國際機場（Taoyuan Intl. Airport）", "第1航廈 A8管制區內，三種口味，24小時供應（每碗等待60～90秒），旅客反映常需排隊", "機場"],
    ["台灣", "台北捷運西門站（Ximen MRT Station）", "1、2號出口設有無人拉麵機＋無人手搖飲機，為台北捷運首座導入據點；與「一風堂」聯名推出「豚骨拉麵」「麻辣豚骨拉麵」；2026/01/21起另於西門站B1、3號出口舉辦與泰國國民泡麵品牌「MAMA」聯名快閃活動，推出4種口味現煮MAMA麵", "捷運站"],
    ["南韓", "首爾大學醫院、首爾市警察廳等企業／公家機構", "以「Robot Chef 啜啜盒（Chulchul Box）」品牌進駐，並規劃擴及飯店、博物館與無人員工餐廳市場", "醫院／公家機關／企業餐飲"],
    ["全球彙總", "150+ 據點（官方宣稱）", "涵蓋北美各城市、日本、台灣、南韓，場域類型含機場、飯店、醫院、企業辦公室、大學校園、會議中心、製造廠、住宅社區", "綜合"],
]
make_table(doc, ["地區", "地點／場域", "細節", "場域類型"], loc_rows, col_widths=[2, 4.5, 7.5, 2.5])

add_para(doc, "")
add_para(doc, "註：以上據點為公開新聞、機場官方社群帳號、旅遊評論平台（Yelp／TripAdvisor／Apple地圖）等管道可查證之時間點快照，實際佈點會隨合約到期／機台汰換而變動，正式作為訴訟損害計算依據前，建議另行實地查證或發函請對方揭露目前完整佈點清單。另查得 Yo-Kai Express 在台灣最早（2021年5～6月COVID疫情期間）曾將機台設置於衛生福利部桃園醫院、亞東紀念醫院，免費提供醫護人員近1.1萬份餐點——此為公益性質之歷史部署，目前查無資料顯示是否仍在運作，故未列入上表「目前」據點，僅供時間軸與市場滲透脈絡參考。", size=10.5, color=GRAY)

doc.add_page_break()

add_heading(doc, "五、顧祖欣／樂活速食坊之自有技術與專利資產（本案權利基礎）", level=1)
add_para(doc, "本節彙整顧總（顧祖欣先生）所屬「樂活速食坊股份有限公司」及其品牌「勝十蘭」的技術研發歷程、專利與市場佈局時間軸。此為評估對 Yo-Kai Express 主張侵權之核心權利基礎，故單獨成節，並於下一節與 Yo-Kai Express 之台灣佈局進行時序比對。", size=12)

add_heading(doc, "（一）公司與品牌背景", level=2)
pp_rows = [
    ["公司名稱", "樂活速食坊股份有限公司（英文：Lohas Fast Food Co., Ltd.）"],
    ["統一編號", "42823196"],
    ["成立時間", "2015年"],
    ["資本額", "新台幣4,500萬元"],
    ["登記地址", "台北市士林區承德路四段182號3樓"],
    ["董事長", "林昌驍（1943年生，曾任大同公司家具部經理，人稱兩岸便利商店「陳列架大王」）"],
    ["總經理", "顧祖欣（Rupert Guh）——即本案委託人「顧總／顧董」"],
    ["核心品牌", "勝十蘭（日式拉麵，Sheng Shih Lan）；另有㗊品軒（中式）、有義'S（義式）、淘泰郎（泰式）、韓拾（韓式）等5大品牌，共用同一套AI蒸汽烹調機平台"],
]
make_table(doc, ["項目", "內容"], pp_rows, col_widths=[3.5, 13])

add_para(doc, "")
add_heading(doc, "（二）技術研發歷程（早於 Yo-Kai Express 進入台灣市場）", level=2)
add_bullet(doc, "2010年：董事長林昌驍開始構思「AI蒸煮機器人」概念。")
add_bullet(doc, "2010～2017年：與大同大學機械工程研究所（所長葉隆吉教授）建教合作，耗時約7年研發第一代機器人，葉隆吉教授稱此為其「有史以來設計過最複雜的機器」，涉及水蒸氣排放、精密溫度控制等技術難題；其中光是為維持「現煮」口感即單獨投入約4年、耗資數千萬元。")
add_bullet(doc, "同期與高雄餐旅大學合作，開發「甜甜圈形」冷凍麵團（具上下左右孔隙特殊設計），可透過蒸氣加熱40秒完成復熱——此為與 Yo-Kai Express「食材急速冷凍、出餐時高溫復熱」概念高度相似的核心技術，惟時間點更早。")
add_bullet(doc, "2017年：研發成功，勝十蘭拉麵於台北南陽街開幕，創下5次翻桌率、單碗麵40秒完成加擺盤1分鐘上桌、日銷700～800碗、月營收近新台幣200萬元之營運實績；同年榮獲經濟部中小企業處「全國服務創新研發首獎」。")
add_bullet(doc, "後續再與台灣科技大學機器人研究所合作約2年，精進至第二代機器人，取得「數項世界性發明專利」。")
add_bullet(doc, "2019年12月（含）以前：勝十蘭於西門町（台北市萬華區中華路一段118號，捷運西門站6號出口步行3分鐘）開設分店，主打「全台首創AI機器人蒸煮」拉麵——**此據點與 Yo-Kai Express 2024年才進駐的台北捷運西門站，位於同一西門商圈，且顧總一方之部署時間至少早了4～5年**。")
add_bullet(doc, "近年：受COVID-19疫情衝擊虧損約新台幣8,000萬元，之後轉型將機台獨立銷售／授權模式，推出第五代機器並完成量產，2024年以「AI蒸氣烹調機」獲頒「2024食創獎－食創之星一星」。")

add_para(doc, "")
add_heading(doc, "（三）AI蒸汽烹調機（SIC001）產品規格", level=2)
spec_rows = [
    ["型號", "SIC001"],
    ["電源規格", "單相 AC230V 50/60Hz，T形插頭"],
    ["額定／最大消耗功率", "2800W ±10%"],
    ["外型尺寸", "主機 318×578×650（mm）"],
    ["重量", "35 ±5 Kg"],
    ["操作面板", "10.4吋 TFT 電容觸控顯示器"],
    ["蒸煮形式", "專用蒸煮容器＋耐高溫發泡紙碗"],
    ["給水方式", "管路直給（需裝設專用過濾器）或專用桶裝水（5.8L）"],
    ["排水方式", "機台前方可拆式水盤"],
    ["核心賣點", "60秒完成烹調；高溫蒸汽滅菌並保存食材營養；中央廚房ISO22000及HACCP認證，全程冷鏈；食材經蒸汽復熱保有原廚師風味；親民售價"],
    ["適用8種商業模式", "一般餐廳／自助或少人供餐(24H)／自販機直接供餐(24H)／一坪變形金剛組合供餐(24H)／1-5坪街邊店(24H+自販機混搭)／貨櫃屋供餐／行動餐車供餐／雲端廚房供餐"],
    ["可烹調5大類別", "湯麵類、乾麵類、水煮類、清蒸類、粥／火鍋類，共涵蓋日式豚骨拉麵、紅燒牛肋條刀削麵、越式鮮牛肉河粉、義大利菌菇燉飯、小籠湯包、麻辣小火鍋等約15道以上菜色"],
]
make_table(doc, ["項目", "內容"], spec_rows, col_widths=[3.5, 13])
add_para(doc, "資料來源：顧總提供之NRA翻譯DM文件與「AI蒸汽烹調機」簡報（2026-07-10補充提供）。", size=9.5, color=GRAY)

add_para(doc, "")
add_heading(doc, "（四）專利與獎項——需優先查證之權利清單", level=2)
own_pat_rows = [
    ["台灣專利申請號 113108458", "AI蒸汽烹調機相關技術（2024年申請，ROC 113年度案號）", "本次由顧總直接提供，尚待赴TIPO查證審定核准狀態、專利類型（發明／新型）與請求項範圍"],
    ["台灣專利申請號 113202327", "AI蒸汽烹調機相關技術（2024年申請，ROC 113年度案號，案號格式屬新型專利常見編碼）", "同上，尚待查證"],
    ["「數項世界性發明專利」（2017年前後，第一、二代機器人）", "與大同大學、台科大合作研發成果，媒體報導稱獲得多項國際專利，惟具體案號未見於公開報導", "建議請顧總提供完整專利清單（含案號、國家、申請/核准日期），並比對是否已逾專利年限（發明專利保護期為申請日起20年，若2010年代初期申請，需注意是否仍在保護期內）"],
    ["2017年 經濟部中小企業處「全國服務創新研發首獎」", "肯定其AI蒸煮機器人商業模式創新，非專利權，但可作為技術原創性與市場肯定之佐證", "可作為訴訟中佐證技術原創時間點的輔助證據"],
    ["2024年「食創獎－食創之星一星」（AI蒸氣烹調機）", "肯定最新一代（SIC001）機型之產品創新", "同上，可作為技術延續發展脈絡的佐證"],
]
make_table(doc, ["權利／榮譽項目", "內容說明", "查證建議"], own_pat_rows, col_widths=[4, 7, 5])
add_para(doc, "⚠️ 重要提醒：上述兩組2024年台灣專利申請號（113108458、113202327）為顧總直接提供之一手資訊，具高度可信度，但本次因 TIPO 專利檢索系統為動態網頁、且無瀏覽器工具可互動查詢，尚未能實際查證其審定核准狀態、專利類型與請求項內容。強烈建議下一步請律師或專利師直接登入台灣智慧財產局 GPSS 系統或專利公開資訊查詢系統，以此二案號查出完整卷宗，這將是後續侵權比對分析（claim mapping）的關鍵素材。", size=10.5, color=RED)

doc.add_page_break()

add_heading(doc, "六、雙方時序與技術比較——本案侵權主張核心論證", level=1)
add_para(doc, "將顧總一方（樂活速食坊／勝十蘭）與 Yo-Kai Express 在台灣市場的發展時序並列比對，可清楚看出顧總一方之技術研發與商業化，明顯早於 Yo-Kai Express 進入台灣市場、且早於其進駐西門商圈：", size=12)

timeline_cmp_rows = [
    ["2010", "顧總一方：董事長林昌驍開始構思AI蒸煮機器人", "—"],
    ["2010～2017", "顧總一方：與大同大學耗時7年研發第一代機器人，含4年專攻「現煮口感」還原", "—"],
    ["2016", "—", "Yo-Kai Express：Andy Lin 於美國矽谷創立公司"],
    ["2017", "顧總一方：研發成功，勝十蘭南陽街店開幕，獲全國服務創新研發首獎", "—"],
    ["2017～2019", "顧總一方：與台科大合作精進第二代機器人，取得數項世界性發明專利", "—"],
    ["2019.12（含）以前", "顧總一方：勝十蘭西門町分店開幕（中華路一段118號，西門站6號出口步行3分鐘），主打「全台首創AI機器人蒸煮」", "—"],
    ["2021.03～06", "—", "Yo-Kai Express：宣布進軍台灣，COVID期間於桃園醫院／亞東醫院提供免費餐點，為其在台最早部署"],
    ["2021.11", "—", "Yo-Kai Express：遭 Noodle Time Holdings 於美國提起商業秘密訴訟"],
    ["2024（確切月份未查得）", "—", "Yo-Kai Express：進駐台北捷運西門站（1、2號出口），與一風堂聯名——**與勝十蘭西門店相距僅約數百公尺內之同一商圈**"],
    ["2024", "顧總一方：以「AI蒸氣烹調機」獲2024食創獎", "—"],
    ["2024（113年度）", "顧總一方：申請台灣專利113108458、113202327", "—"],
    ["2026.01.21", "—", "Yo-Kai Express：於西門站B1/3號出口舉辦泰國MAMA麵聯名快閃活動"],
]
make_table(doc, ["時間", "顧總一方（樂活速食坊／勝十蘭）", "Yo-Kai Express"], timeline_cmp_rows, col_widths=[2.8, 6.6, 6.6])

add_para(doc, "")
add_heading(doc, "（一）本比對表對訴訟策略的意義", level=2)
add_bullet(doc, "**時間優先性（priority）**：顧總一方之「冷凍食材＋蒸氣∕高溫復熱＋自動化機台」技術，研發起始於2010年、商業化於2017年，均早於 Yo-Kai Express 2016年創立、更早於其2021年進入台灣市場。若雙方技術方案有實質相似之處，時間優先性對顧總一方有利。")
add_bullet(doc, "**地緣重疊性**：顧總一方之勝十蘭西門分店至少自2019年底即在西門商圈以「AI機器人蒸煮」為號召營運，而 Yo-Kai Express 遲至2024年才在近乎同一商圈（台北捷運西門站）設置功能相似之自動化熱食機台，兩者消費族群、行銷語彙（「AI」「機器人」「自動」「秒級復熱」）高度重疊，容易使消費者產生來源混淆，此點除可作為專利／技術侵權主張的背景佐證外，亦可一併評估是否有主張「不公平競爭」或「著名表徵/商業外觀（trade dress）」的空間。")
add_bullet(doc, "**專利比對為關鍵下一步**：目前僅查得雙方各自的專利「案號」與「概略技術描述」，尚未取得任一方專利之完整請求項（claims）內容。是否構成侵權，必須由專利師將顧總113108458、113202327（及若能取得的更早期專利）之請求項範圍，與 Yo-Kai Express 之 WO2021050686A1（Food Warming Device）、US2020043028（Food Vending Machine）等專利說明書逐項比對技術特徵（claim charting），才能得出具體侵權結論；本報告僅整理背景事實與時序，不構成侵權與否之法律判斷。")
add_bullet(doc, "**反向風險提醒**：由於雙方多年來在同一產業（自動化熱食/拉麵販賣）並行發展，若正式提告，對造極可能提出反向質疑或反訴（如質疑顧總專利的新穎性、或主張雙方技術路徑其實有本質差異如「蒸氣復熱」vs「微波/水浴加熱」），建議及早請專利師就雙方技術路徑之異同做細部工程分析，避免僅憑「時間較早」與「行業相似」就貿然提告。")

add_para(doc, "")

doc.add_page_break()

add_heading(doc, "七、資料來源清單", level=1)
sources = [
    ("GreyB Insights;Gate - Yo-Kai Express 專利統計分析（含美國核准率0%數據）", "https://insights.greyb.com/yo-kai-express-patents/"),
    ("ipqwery - Yo-Kai Express Inc. WIPO/PCT 專利明細清單", "https://www.ipqwery.com/ipowner/en/owner/ip/4199173-yo-kai-express-inc.html?rgk=IPType&rvk=Patent&rgk=Jurisdiction&rvk=WIPO"),
    ("Justia Patents - D985424／D1003643 設計專利核准資訊", "https://patents.justia.com/assignee/yo-kai-express-inc"),
    ("Yo-Kai Express 官網 - 據點頁面", "https://www.yokaiexpress.com/location"),
    ("Yo-Kai Express 官網", "https://www.yokaiexpress.com/"),
    ("Law360 - Noodle Time Holdings v. Yo-Kai Express 案件頁", "https://www.law360.com/cases/61a7019821fb2a9a7b4e231b"),
    ("Justia Patents - Yo-Kai Express Inc. 專利檢索", "https://patents.justia.com/assignee/yo-kai-express-inc"),
    ("Trademark Elite - YO-KAI 商標（Noodle Time Holdings Limited）", "https://www.trademarkelite.com/trademark/trademark-detail/87346086/YO-KAI"),
    ("WIPO/Sumobrain - Food Warming Device 專利 WO2021050686A1", "https://www.sumobrain.com/patents/wipo/Food-warming-device/WO2021050686A1.html"),
    ("Crunchbase - Yo-Kai Express 公司檔案", "https://www.crunchbase.com/organization/yo-kai-express-inc"),
    ("JSTORIES - Elon Musk approves 報導", "https://jstories.media/article/yo-kai-express"),
    ("經濟部中小及新創企業署 Startup Portal Taiwan 轉載報導", "https://startup.sme.gov.tw/en/NewsContent/708"),
    ("SoftBank Robotics 官方新聞稿（CHEFFY 經銷權移轉）", "https://www.softbankrobotics.com/news/20240313/"),
    ("Restaurant Robotics Technology - SoftBank Robotics/CHEFFY 報導", "https://restaurantroboticstechnology.com/news/softbank-robotics-to-assume-distribution-business-of-yo-kai-express-cheffy/"),
    ("LIVE JAPAN - CHEFFY 介紹", "https://livejapan.com/en/in-tokyo/in-pref-tokyo/in-haneda/article-a0005312/"),
    ("PR Newswire - Pulmuone × Yo-Kai Express 韓國合作新聞稿", "https://www.prnewswire.com/news-releases/pulmuone-launches-gourmet-vending-machines-across-south-korea-in-partnership-with-yo-kai-express-301939908.html"),
    ("The Korea Herald - Pulmuone 熟食販賣機報導", "https://www.koreaherald.com/view.php?ud=20230207000644"),
    ("Startup Terrace - Yo-Kai Express 創辦人林志鴻專訪", "https://www.startupterrace.tw/news/522"),
    ("優果科技（Youguo Tech）官網 - Yo-Kai Express 案例介紹", "https://www.youguotec.com/yo-kai-express/"),
    ("NYSE TV - Andy Lin 談 NRF 2026 展出內容", "https://tv.nyse.com/nyse-tv-highlights/season:4/videos/yo-kai-express-ceo-andy-lin-discusses-what-it-s-showcasing-at-nrf-2026"),
    ("The Spoon - CES 2023 YKE Desktop 報導", "https://thespoon.tech/yo-kai-to-debut-desktop-ramen-robot-for-space-constrained-retail-formats-at-ces-2023/"),
    ("The Spoon - CES 2024 珍珠奶茶機器人報導", "https://thespoon.tech/yo-kai-to-debut-boba-making-robot-at-ces-2024/"),
    ("The Spoon - SFO 機場自動拉麵機報導", "https://thespoon.tech/yo-kai-express-opens-up-its-automated-hot-ramen-machine-at-sfo/"),
    ("APEX - 機場販賣機新聞報導", "https://apex.aero/articles/2021-01-25-yo-kai-express-brings-fresh-airport-vending-machines/"),
    ("Ottomate News - 羽田機場拉麵機報導", "https://www.ottomate.news/p/yo-kai-launches-ramen-vending-machine"),
    ("Taoyuan Airport 官方 X（Twitter）貼文 - 桃園機場據點資訊", "https://x.com/taoyuanairport/status/1775100874576331118"),
    ("Yelp - Ontario 機場門市頁", "https://www.yelp.com/biz/yo-kai-express-ontario-3"),
    ("Yelp - Houston 門市頁", "https://www.yelp.com/biz/yo-kai-express-houston-6"),
    ("Yelp - Flushing 門市頁", "https://www.yelp.com/biz/yo-kai-express-flushing"),
    ("Yelp - Tahoe City 門市頁", "https://www.yelp.com/biz/yo-kai-express-tahoe-city-6"),
    ("MirrorReview - Andy Lin 專訪", "https://www.mirrorreview.com/andy-lin-food-and-beverage-leader/"),
    ("DailyView 網路溫度計 - 西門站機器人現煮拉麵報導", "https://dailyview.tw/popular/detail/27307"),
    ("城市學（遠見）- 西門站無人自助拉麵報導", "https://city.gvm.com.tw/article/116022"),
    ("NOWnews - 桃園機場機台排隊爆滿報導", "https://www.nownews.com/news/6665595"),
    ("壹蘋新聞網 - 林志鴻專訪（西門站/台灣據點僅西門站與桃園機場）", "https://news.nextapple.com/finance/20250327/35BB19A194B69D0AFB705F72CCFBD54A"),
    ("Mega News 報新聞 - Yo-Kai×泰國MAMA麵合作西門站快閃報導", "https://www.contentplatform.info/articles/412066/%E5%BC%B7%E5%BC%B7%E8%81%AF%E6%89%8B%E6%8E%80%E8%B5%B7%E5%8F%B0%E7%81%A3ai-%E5%8D%B3%E9%BB%9E%E5%8D%B3%E7%85%AE%E6%B3%A1%E9%BA%B5%E6%96%B0%E9%A3%9F%E5%B0%9A%EF%BC%81%E3%80%80%E6%B3%B0%E5%9C%8B-mama/"),
    ("Technice - Yo-Kai×MAMA麵、NVIDIA Inception技術合作報導", "https://www.technice.com.tw/experience/enterprise/204482/"),
    ("欣傳媒 XINMEDIA - Yo-Kai×MAMA麵登陸西門町報導", "https://www.xinmedia.com/article/304974"),
    ("Taiwan Tech Arena - Yo-Kai Express 台灣新創檔案（優豈股份有限公司地址）", "https://www.taiwanarena.tech/startups-detail/YoKaiExpress/"),
    ("104人力銀行 - 優豈股份有限公司 公司頁", "https://www.104.com.tw/company/1a2x6bmz10"),
    ("BusinessNext數位時代 - Meet創業小聚：Yo-Kai Express公司成立時間/台灣落地計畫報導", "https://meet.bnext.com.tw/articles/view/47342"),
    ("痞客邦 - 桃園機場第一航廈拉麵自動販賣機部落格報導", "https://flyin24.pixnet.net/blog/post/577905268"),
    ("今周刊 - 樂活速食坊林昌驍／AI蒸煮機器人研發歷程報導", "https://www.businesstoday.com.tw/article/category/154687/post/201805300010/"),
    ("食力foodNEXT - 食・大人物會客室：樂活速食坊AI餐飲革命／顧祖欣專訪", "https://www.foodnext.net/column/columnist/paper/6851080249"),
    ("台灣公司網 twincn - 樂活速食坊股份有限公司登記資料（董監事含顧祖欣）", "https://www.twincn.com/item.aspx?no=42823196"),
    ("痞客邦 - 勝十蘭西門店開幕報導（2019/12/23，全台首創AI機器人蒸煮）", "https://angela820524.pixnet.net/blog/post/47193714"),
    ("樂活速食坊官網 - 旗下五大品牌介紹（勝十蘭/㗊品軒/有義'S/淘泰郎/韓拾）", "https://www.lahasff.com/article.php?lang=tw&tb=4&id=437"),
    ("104人力銀行 - 勝十蘭日式拉麵／樂活速食坊股份有限公司", "https://www.104.com.tw/company/1a2x6bjkvw"),
    ("顧總提供 - NRA翻譯Steam I cooking DM文件（AI Steam Cooking Machine產品DM，內含台灣專利申請號）", "（顧總提供之本機檔案，非公開網路連結）"),
    ("顧總提供 - 「AI蒸汽烹調機」簡報（SIC001規格、五大核心價值、8種商業模式）", "（顧總提供之本機檔案，非公開網路連結）"),
]
for title, url in sources:
    add_source_line(doc, title, url)

out_path = r"C:\Users\JasonLee\claude_code_projects\CMO\docs\AI-CMO_Yo-Kai Express 新聞彙整及專利相關暨市場擺放點位調查報告_v5.docx"
doc.save(out_path)
print("SAVED:", out_path)
