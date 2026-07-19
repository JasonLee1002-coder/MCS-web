# -*- coding: utf-8 -*-
"""AI-CMO 社區智慧櫃競品與場域市場研究報告產生器"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GOLD = RGBColor(0xB8, 0x68, 0x0A)
RED = RGBColor(0x8B, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHTBLUE = "EBF3FB"

def set_font(run, cn="Noto Sans TC", en="Calibri", size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    if color:
        run.font.color.rgb = color

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, "Noto Serif TC", "Cambria", 15, True, NAVY)
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, "Noto Sans TC", "Calibri", 13, True, GOLD)
    return p

def add_body(doc, text, indent=0, color=None, bold=False, size=12):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(hdr[i], "1F4E79")
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_font(r, size=10)
            if ridx % 2 == 0:
                shade_cell(cells[i], LIGHTBLUE)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

doc = Document()
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# ── 封面 ──
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(80)
r = cover.add_run("銓幻元科技 MCS")
set_font(r, "Noto Serif TC", "Cambria", 16, True, NAVY)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(20)
r = title.add_run("社區智慧櫃")
set_font(r, "Noto Serif TC", "Cambria", 26, True, NAVY)
r2 = title.add_run("\n競品與場域市場研究")
set_font(r2, "Noto Serif TC", "Cambria", 22, True, NAVY)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(30)
r = sub.add_run("回覆 CPO MSG-1283EF ｜ 撰寫者：AI CMO")
set_font(r, size=13, color=GOLD, bold=True)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("版本 v1.0　2026-07-19")
set_font(r, size=11, color=GREY)

doc.add_page_break()

# ── 一、研究範疇與方法 ──
add_heading1(doc, "一、研究範疇與方法")
add_body(doc, "本報告回應CPO《智慧自取冰櫃產品規劃深度分析》情境A「社區智慧櫃」的市場驗證需求。CPO定義：住戶線上預訂生鮮/冷凍食材/餐盒，配送方或MCS補貨員放入櫃內，住戶憑會員身份自取回家料理——本質是「生鮮/餐盒最後一哩取貨基礎設施」，非衝動型販賣機。")
add_body(doc, "研究方法：兩組獨立agent分別對「台灣/中國競品」與「台灣目標場域」做網路查證，所有數字均附來源URL，查無可靠資料處誠實標註，不做推測性填補。")
add_body(doc, "⚠ 本報告僅提供market fact base（事實查證），不代替CPO做策略判斷；商業模式選定（SaaS月費/取貨手續費分潤/自營商品毛利三選項）仍需CPO/業務端進一步驗證後定案。", color=RED)

# ── 二、台灣競品掃描 ──
add_heading1(doc, "二、台灣競品掃描")
add_body(doc, "結論：台灣目前查無任何「自動化冷藏/冷凍生鮮智取櫃」案例，是完全空白賽道——但也代表沒有任何第三方驗證數據可借鏡，MCS若推行將是市場首例。", bold=True, color=GOLD)
add_table(doc,
    ["案例", "定位", "收費模式", "規模"],
    [
        ["全聯 PXGo 小時達", "生鮮電商，門市櫃台有人取貨或外送，非智慧櫃", "免運（滿額）", "投入約20億建6座生鮮物流中心，近千店"],
        ["家樂福到店取貨", "門市櫃台有人取貨，非智慧櫃", "滿199免運", "據點數查無資料"],
        ["蝦皮店到店智取門市", "無人商店+電商包裹智取櫃，非生鮮，未見溫控功能", "未載明", "全台僅4間（雙北），對比有人門市800+間"],
        ["中華郵政 i郵箱", "常溫包裹智能櫃", "依規格計價", "全台最大營運商，逾2,400據點"],
        ["掌櫃", "常溫智取櫃", "—", "巔峰逾1,000點，2020年已停業"],
        ["吉達思科技 G-DAS", "社區包裹智慧櫃方案商，官網未列溫控/生鮮功能", "賣斷/月租，未公開報價", "官網僅約12案例照片，客戶數未揭露"],
        ["foodpanda 熊貓超市(pandamart)", "暗倉20分鐘生鮮外送到家（非自取櫃）", "—", "2024年5月全面終止"],
        ["全家「好生凍」冷凍複合店3.0", "冷凍店取，官方稱將試營運於商辦大樓/住宅商業區", "—", "查無最新規模數字"],
    ],
    col_widths=[3.3, 5.5, 3.5, 4.2]
)

# ── 三、中國競品掃描 ──
add_heading1(doc, "三、中國競品掃描")
add_body(doc, "結論：主流社區團購（次日自提+抽成）多為輕資產/有人駐點模式，且大多數已倒閉或全國關停；真正硬體形態的「智能貨櫃」案例規模小但仍持續獲利。", bold=True, color=GOLD)
add_table(doc,
    ["案例", "定位", "收費模式", "規模"],
    [
        ["美團優選", "次日自提+團長抽成，重資產自建倉配", "對團長抽成10-12%", "2019-2024累計虧損逾1,100億人民幣，2025年12月15日全國關停"],
        ["多多買菜", "自提點=既有雜貨店/水果店（非智能櫃），一家獨大", "對自提點抽成10-20%", "2025年GMV近3,000億人民幣，月增速60-80%"],
        ["興盛優選", "區域收縮（湘鄂贛大本營）+私域自提", "團長佣金10-12%（金牌可達12%）", "合作門市逾30萬個，估值100億美元"],
        ["淘菜菜/淘寶買菜", "阿里社區團購", "—", "2025年3月全國停運次日自提業務，具體規模查無資料"],
        ["盒馬鄰里自提店", "社區自提點，硬體形態(智能櫃or有人)查無資料", "—", "巔峰1,500點，已全面關閉"],
        ["錢大妈「菜吧」無人貨櫃 ★最接近案例", "5-8台改造智能冰箱/冷凍櫃，封閉小區生鮮自取，RFID+人臉辨識", "加盟商前期投入約12萬人民幣+銷售額2%品牌管理費", "深圳廣州近百網點，毛利率約15%，單點月淨利可逾1萬人民幣"],
        ["京東到家/物流智能冷鏈自提櫃", "2022年政策鼓勵方向，非已落地具體案例", "—", "查無資料"],
    ],
    col_widths=[4.2, 5.3, 3.7, 3.3]
)
add_body(doc, "查無可靠資料項目：興盛優選對供應商端毛利拆分、淘菜菜停運前具體規模數字、盒馬鄰里硬體形態、永輝/百果園無人貨櫃規模數字、京東系智能冷鏈自提櫃落地規模。", color=GREY, size=10.5)

# ── 四、台灣目標場域盤點 ──
add_heading1(doc, "四、台灣目標場域盤點")
add_body(doc, "依CPO指示優先原則：高密度住宅、性質接近辦公/封閉園區者（如竹科員工宿舍、企業社區），而非一般散戶社區——取貨頻率密度不夠撐不起營運成本。")
add_table(doc,
    ["場域", "具體案例", "估計規模", "生鮮需求佐證"],
    [
        ["竹科周邊社宅", "中雅安居（新竹市北區）", "638戶，2026完工", "foodpanda已覆蓋新竹東區/竹科周邊"],
        ["竹科周邊社宅", "建功安居（新竹市東區）", "743戶，2028完工", "新竹市4處社宅合計2,503戶"],
        ["竹科員工宿舍", "SIPA直營宿舍（松苑/柏苑）", "總戶數查無資料", "—"],
        ["中科周邊社宅", "國安一期好宅（西屯區）", "500戶，周邊即全聯+M平方商場", "媒體點名為中科人首選區段"],
        ["南科周邊社宅 ★最貼近案例", "新市安居（台南新市區）", "670戶，距南科僅500公尺", "家樂福新市民生店官方稱「打造科技人便捷採購新據點」"],
        ["高雄科技園區周邊", "復悅安居（左營/楠梓）", "220戶", "—"],
        ["企業員工宿舍", "台積電新竹員工宿舍", "約300套房（非大型社區規模）", "查無"],
        ["大型社宅基準對照", "林口世大運選手村社宅", "2,500戶（全台最大單一案）", "非科技園區周邊，作密度基準參考"],
    ],
    col_widths=[3.2, 5.0, 3.8, 4.5]
)

add_heading2(doc, "重要澄清/風險訊號")
for t in [
    "「新竹台元宿舍」原假設有誤——台元科技園區是商辦園區（2萬員工），非住宅宿舍，建議從候選清單剔除。",
    "台積電/聯電/日月光在台均未走「大型自建員工社區」模式，多為承租改裝或與學校合作，規模僅數百戶級距。",
    "熊貓超市(pandamart) 2024年5月全台停業是重要負面訊號——曾主打25分鐘生鮮配送，停業前竹科園區內部本身多數不外送，顯示「科技園區生鮮快速配送」商業模式在台灣已失敗過一次。",
    "目前查到的科技園區周邊社宅單案最大僅670-743戶，遠低於選手村2,500戶等級，可能需要組合多案場才夠密度。",
    "查無任何第三方獨立驗證的智取櫃試點案例（僅查到MCS/龍雲數位自家行銷文案）——MCS若推行將是市場首例，無現成失敗/成功案例可借鏡。",
]:
    add_body(doc, "• " + t, indent=0.4, color=RED, size=11)

# ── 五、Market Fact Base 摘要 ──
add_heading1(doc, "五、Market Fact Base 摘要（僅陳述事實，不做策略判斷）")
add_body(doc, "台灣是完全空白賽道（無在地智取櫃競品，但也無驗證數據）；中國最接近的硬體案例（錢大妈菜吧）仍在近百點規模持續運營且有明確15%毛利、月淨利逾萬元人民幣的數字，而規模更大的「次日自提+抽成」平台模式（美團優選、盒馬鄰里）均已在千億虧損/1,500點規模後全面關停——訊號是「小規模封閉場域+硬體自取」比「大規模平台化自提」更可能存活。")
add_body(doc, "台灣候選場域中，新市安居（南科，670戶，家樂福官方認證科技人採購據點）與國安一期好宅（中科，500戶，鄰近全聯商場）是目前查到最具體、佐證最完整的兩個切入點；竹科周邊反而因熊貓超市失敗案例而需要更謹慎評估訂單密度風險。")
add_body(doc, "此為CMO提供的market fact base，收費模式三選項（SaaS月費／取貨手續費分潤／自營商品毛利）與場域切入順序的策略判斷，仍交由CPO/業務端依此事實基礎進一步定案。", bold=True, color=NAVY)

# ── 頁尾 ──
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("— AI CMO 產出．2026-07-19 —")
set_font(r, size=10, color=GREY, bold=True)

out_path = r"C:\Users\JasonLee\claude_code_projects\CMO\docs\AI-CMO_社區智慧櫃競品與場域市場研究_20260719.docx"
doc.save(out_path)
print("Saved:", out_path)
