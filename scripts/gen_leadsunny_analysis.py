# -*- coding: utf-8 -*-
"""AI-CMO 立璨股份有限公司分析報告產生器"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GOLD = RGBColor(0xB8, 0x68, 0x0A)
RED = RGBColor(0x8B, 0x1A, 0x1A)
GREEN = RGBColor(0x1E, 0x7A, 0x3C)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHTBLUE = "EBF3FB"
LIGHTGOLD = "FFF3CD"

def set_font(run, cn="Noto Sans TC", en="Calibri", size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    if color:
        run.font.color.rgb = color

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, "Noto Serif TC", "Cambria", 15, True, NAVY)
    return p

def add_heading2(doc, text, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, "Noto Sans TC", "Calibri", 13, True, color)
    return p

def add_body(doc, text, indent=0, color=None, bold=False, size=12):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(hdr[i], "1F4E79")
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_font(r, size=10)
            if ridx % 2 == 0:
                shade_cell(cells[i], LIGHTBLUE)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_box(doc, title, text, bgcolor, titlecolor):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    shade_cell(cell, bgcolor)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_font(r, size=11.5, bold=True, color=titlecolor)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(text)
    set_font(r2, size=10.5)
    doc.add_paragraph()

doc = Document()
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# ── 封面 ──
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(70)
r = cover.add_run("競品／夥伴研究")
set_font(r, "Noto Serif TC", "Cambria", 15, True, NAVY)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(18)
r = title.add_run("立璨股份有限公司")
set_font(r, "Noto Serif TC", "Cambria", 26, True, NAVY)
r2 = title.add_run("\n公司分析與合作可能性評估")
set_font(r2, "Noto Serif TC", "Cambria", 18, True, NAVY)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(26)
r = sub.add_run("撰寫者：AI CMO")
set_font(r, size=13, color=GOLD, bold=True)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("版本 v1.0　2026-07-25")
set_font(r, size=11, color=GREY)

doc.add_page_break()

# ── 摘要 ──
add_heading1(doc, "摘要")
add_body(doc, "立璨股份有限公司（leadsunny.com.tw）是一家1996年成立、深耕咖啡機與餐飲設備進口代理近30年的公司，2018年引進德國Spengler的AI無人商店設備技術，跨入智慧無人零售領域。與龍雲數位/銓幻元的關係定位為「既競爭又互補」：在辦公室/園區咖啡自助設備這個細分市場上是潛在競爭對手，但在咖啡機硬體供應鏈深度與進口代理經驗上，可能是有價值的零組件/供應合作對象。本報告基於WebFetch/WebSearch真實查證的公開資訊撰寫，未查得雙方現有正式合作關係。", bold=True)

# ── 一、公司基本資料 ──
add_heading1(doc, "一、公司基本資料")
add_table(doc, ["項目", "內容"], [
    ["公司全名", "立璨股份有限公司"],
    ["成立時間", "1996年（近30年）"],
    ["公司地址", "新北市汐止區大同路一段308號，亞太經貿園區B棟"],
    ["聯絡電話", "(02)2643-5000"],
    ["維護專線", "0938-938-982 / 0938-938-991"],
    ["電子信箱", "lesunny@ms15.hinet.net"],
    ["官方網站", "www.leadsunny.com.tw"],
], col_widths=[3.5, 10.5])
add_body(doc, "查證方式：WebFetch官網leadsunny.com.tw + WebSearch多筆結果交叉確認，資訊來源為公司官網公開內容。", color=GREY, size=10)

# ── 二、公司沿革與定位 ──
add_heading1(doc, "二、公司沿革與定位")
add_body(doc, "1996年成立，早期主業是咖啡機與餐飲設備的進口代理，服務台灣餐飲/咖啡市場超過20年。2018年是關鍵轉折點——率先從德國Spengler引進AI科技無人商店自動販賣設備進入台灣市場，官方定位是「運用新AI科技解決人力嚴重缺乏及房租高漲問題」，並自稱「為市場帶來全新的主流」無人商店模式。")
add_body(doc, "官方自我定位關鍵字：「經市場比較驗證的熱銷品牌機型」「操作簡單易上手」「專人維修保養服務」。整體品牌調性偏向「設備供應商／進口代理商」，而非強調自主技術研發。", color=GOLD)

# ── 三、業務範圍與產品線 ──
add_heading1(doc, "三、業務範圍與產品線")
add_table(doc, ["業務範疇", "說明"], [
    ["aicafe AI無人咖啡館", "合作經營及設置推廣，多款型號（605/320/510/96等）"],
    ["AI無人商店自動販賣機", "研發設計與製造，源自德國Spengler技術引進"],
    ["商用咖啡機進口代理", "半自動/全自動咖啡機，代理瑞典品牌、KALERM等國際品牌"],
    ["米其林級廚房設備", "代理ANGELO PO等義大利廚房設備品牌，蒸烤箱等"],
    ["設備租賃服務", "辦公室及業務用設備租賃，含咖啡機、廚房設備"],
    ["物料供應", "咖啡豆、濃縮果汁、即溶飲料、茶葉茶粉、進口鮮奶等耗材"],
], col_widths=[4, 10])

add_heading2(doc, "經營模式")
add_table(doc, ["方案類型", "內容"], [
    ["自購優惠方案", "直接購買機器享折扣"],
    ["租賃方案", "一年租借、三年分期租送、辦公室咖啡租借"],
    ["合作經營模式", "合夥經營、免租設置、業務派遣"],
    ["共享咖啡與無人店舖", "aicafe無人咖啡館設置推廣"],
], col_widths=[4, 10])

# ── 四、國際品牌合作網絡 ──
add_heading1(doc, "四、國際品牌合作網絡")
add_body(doc, "立璨的核心優勢是進口代理網絡的深度與廣度，官網明確提及以下合作品牌：")
for t in ["德國 Spengler — AI自動販賣設備技術來源(2018年引進)",
          "瑞典品牌咖啡機 — EB-61、EX系列、ONE等機型",
          "義大利 ANGELO PO — 米其林級蒸烤箱設備",
          "KALERM — K96全自動咖啡機"]:
    add_body(doc, "• " + t, indent=0.4)
add_body(doc, "官網未具體列舉台灣本地通路合作案例，本報告不做未經查證的具體案例引用。", color=RED, size=10.5)

# ── 五、與龍雲數位/銓幻元的關聯分析 ──
add_heading1(doc, "五、與龍雲數位/銓幻元的關聯分析")

add_heading2(doc, "本質差異：進口代理商 vs 自主研發商", NAVY)
add_table(doc, ["維度", "立璨股份", "龍雲數位/銓幻元"], [
    ["核心模式", "進口代理／設備轉售／租賃", "軟體＋硬體＋韌體全自主研發"],
    ["技術來源", "德國Spengler等國外技術引進", "自有IVM/OmniCore平台、GraBox自主開發"],
    ["優勢領域", "咖啡機／餐飲設備供應鏈深度（近30年）", "IoT數據平台、深度客製化、量身訂做能力"],
    ["客製彈性", "受限於原廠規格，代理商角色", "軟硬韌三層可依場域客製"],
    ["市場歷史", "1996年成立，餐飲設備老牌", "2013年成立，IoT智慧零售新創"],
], col_widths=[3, 5.5, 5.5])

add_box(doc, "⚠️ 值得注意：立璨正是「代理轉售模式」的真實案例",
    "本公司07-24發布的SEO文章〈智慧販賣機廠商怎麼選？自主研發 vs 代理轉售的坪效差異〉，論述的「代理商拿到的機器是黑盒子，改不了底層韌體邏輯」，立璨正是這個模式的具體市場案例——他們的AI無人商店技術是2018年從德國Spengler整批引進，而非自主開發。這對我方內容行銷是一個可以引用的真實市場對照組（但撰寫時仍需比照全站規範，不做針對特定公司的負面斷言，只做客觀模式對照，且不可捏造具體業績數字）。",
    LIGHTGOLD, GOLD)

add_heading2(doc, "潛在競爭關係", RED)
add_body(doc, "在「辦公室／園區咖啡自助設備」這個細分市場上，立璨的aicafe系列與我方GraBox智取櫃＋咖啡機組合方案，客群高度重疊（辦公室、廠區、學校）。立璨近30年的餐飲設備銷售關係網絡，可能已經跟不少企業客戶有既有往來，這是需要留意的競爭變數。")

add_heading2(doc, "潛在合作/互補可能性", GREEN)
for t in [
    "**咖啡機硬體供應合作**：立璨在商用咖啡機進口代理有近30年經驗，代理多個國際精品咖啡機品牌（KALERM、瑞典品牌等）。我方「辦公室茶水間坪效優化」等場域方案若需要專業咖啡機硬體，可評估向立璨採購/代理，而非自行開發咖啡沖煮硬體——把資源集中在軟體平台與韌體整合的核心優勢上，硬體向專業供應商採購。",
    "**通路互補/異業轉介**：立璨客群偏餐飲/咖啡背景，我方客群偏IoT/企業數位轉型背景，兩者接觸的業主圈子可能不完全重疊，有機會做客戶轉介（客戶若需要深度咖啡品項規劃找立璨，需要深度數據平台/客製化找我方）。",
    "**耗材/物料供應鏈**：立璨已建立咖啡豆、濃縮果汁、進口鮮奶等物料供應體系，我方若擴大咖啡類自助設備佈局，物料採購可評估與立璨接洽，不用重新建立供應鏈關係。",
]:
    add_body(doc, "• " + t, indent=0.4)

# ── 六、建議 ──
add_heading1(doc, "六、建議")
for t in [
    "① 這份分析目前完全基於公開網路資訊，尚未查證雙方是否已有既存往來關係，建議先由COO/業務端確認過去是否已接觸過立璨",
    "② 若考慮咖啡機硬體供應合作，建議先小規模接觸探詢（例如單一場域方案的咖啡機採購洽詢），確認報價與供貨穩定度後再評估深化合作",
    "③ 內容行銷上可以把立璨當作「代理轉售模式」的市場對照案例佐證，但撰寫時要客觀陳述經營模式差異，不做未經查證的負面評價或誇大己方優勢的具體數字宣稱",
    "④ 若要進一步了解立璨的合作經營/加盟具體條件，官網「合作經營」頁面(product-444.html)本次WebFetch未能取得完整內容，建議直接電洽(02)2643-5000或請業務窗口接觸了解實際條件",
]:
    add_body(doc, t, indent=0.4)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("— AI CMO 產出．2026-07-25 —")
set_font(r, size=10, color=GREY, bold=True)

out_path = r"C:\Users\JasonLee\claude_code_projects\CMO\docs\AI-CMO_立璨股份有限公司分析與合作可能性評估_20260725.docx"
doc.save(out_path)
print("Saved:", out_path)
