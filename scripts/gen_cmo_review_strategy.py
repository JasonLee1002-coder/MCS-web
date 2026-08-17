# -*- coding: utf-8 -*-
"""AI-CMO 成績檢討與下一輪策略報告"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GOLD = RGBColor(0xB8, 0x68, 0x0A)
RED = RGBColor(0x8B, 0x1A, 0x1A)
GREEN = RGBColor(0x1E, 0x7A, 0x3C)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHTBLUE = "EBF3FB"
LIGHTRED = "FDEDEC"
LIGHTGOLD = "FFF3CD"

def set_font(run, cn="Noto Sans TC", en="Calibri", size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    if color:
        run.font.color.rgb = color

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, "Noto Serif TC", "Cambria", 15, True, NAVY)
    return p

def add_heading2(doc, text, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, "Noto Sans TC", "Calibri", 13, True, color)
    return p

def add_body(doc, text, indent=0, color=None, bold=False, size=12):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(hdr[i], "1F4E79")
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_font(r, size=10)
            if ridx % 2 == 0:
                shade_cell(cells[i], LIGHTBLUE)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_box(doc, title, text, bgcolor, titlecolor):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    shade_cell(cell, bgcolor)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_font(r, size=11.5, bold=True, color=titlecolor)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(text)
    set_font(r2, size=10.5)
    doc.add_paragraph()

doc = Document()
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# ── 封面 ──
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(70)
r = cover.add_run("CMO 成績檢討")
set_font(r, "Noto Serif TC", "Cambria", 15, True, NAVY)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(18)
r = title.add_run("三週戰果檢討與下一輪策略")
set_font(r, "Noto Serif TC", "Cambria", 24, True, NAVY)
r2 = title.add_run("\n2026/07/06 — 2026/07/27")
set_font(r2, "Noto Serif TC", "Cambria", 16, True, NAVY)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(26)
r = sub.add_run("撰寫者：AI CMO")
set_font(r, size=13, color=GOLD, bold=True)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("版本 v1.0　2026-07-27")
set_font(r, size=11, color=GREY)

doc.add_page_break()

# ── 摘要 ──
add_heading1(doc, "總結：一句話")
add_box(doc, "內容引擎全速運轉，但轉換漏斗100%卡死",
    "近3週在5個網站產出超過280篇文章/模板，成功把陌生人變成39筆具體客戶詢問（Lead）。但這39筆詢問的處理狀態全部卡在「新詢問」，0筆進到「已聯絡」以後的任何階段——行銷端的努力目前是有去無回。下一輪最高優先級不是「再生產更多內容」，而是「把現有39筆詢問接起來」。",
    LIGHTRED, RED)

# ── 一、內容產出戰果 ──
add_heading1(doc, "一、內容產出戰果（07/06–07/27）")
add_table(doc, ["網站", "新增內容量", "備註"], [
    ["李奇申.com (Jason-SEO)", "17篇新文章", "含旗艦文「坪效提升王」定位＋餐飲AI自助化12篇補缺口批次"],
    ["transtep.com", "18篇新文章", "含「自主研發vs代理轉售」等5篇＋B2B/GEO/場域補缺口13篇批次"],
    ["mcstation.ai（已上線）", "5篇新頁面", "辦公茶水間/坪效計算法/診所/美容業/一坪店 等場域文"],
    ["mcstation.ai（已寫但未上線）", "17篇待發布", "⚠️卡在等「Chat Brain」上線才接手，目前是閒置產能"],
    ["森林藥局聯盟（pharmacy-blogs）", "157個新模板", "×4間加盟藥局＝約628個實際頁面；07/26起每日自動化5篇/天上線"],
    ["黃維德.com（huangweide-seo）", "82篇新文章", "含既有站台SEO持續擴充"],
], col_widths=[4.5, 4, 6.5])
add_body(doc, "三週內5個網站合計新增內容量超過280個單位（文章/模板），是本輪最主要的產出項目。", color=GREY, size=10.5)

# ── 二、關鍵警訊 ──
add_heading1(doc, "二、關鍵警訊：轉換漏斗100%卡死")
add_body(doc, "查Notion CRM（龍雲數位—客戶詢問CRM）全部39筆真實客戶詢問記錄，依來源網站與狀態交叉統計：")
add_table(doc, ["來源網站", "狀態＝新詢問（筆數）", "已聯絡／報價／成交（筆數）"], [
    ["transtep.com", "20", "0"],
    ["李奇申.com", "13", "0"],
    ["mcstation.ai", "4", "0"],
    ["普健生醫", "1", "0"],
    ["森林藥局聯盟", "1", "0"],
    ["黃維德.com", "1", "0"],
], col_widths=[5, 6, 6])
add_body(doc, "全站39筆詢問，「已聯絡」以後的狀態是0筆。這不是單一網站的問題，是所有網站共通的結構性缺口。", bold=True, color=RED)

add_heading2(doc, "根因鏈（已知、部分已通報未修復）", RED)
add_table(doc, ["問題", "發現日期", "狀態"], [
    ["Monique從未收到CRM新單LINE通知（MONIQUE_LINE_USER_ID正式環境從未設定）", "07-21", "❌ 未修復，07-26二次驗證仍未修"],
    ["hasValidContact()防呆過鬆，「test@」「not sure yet」等垃圾聯絡方式被判定有效", "07-19", "❓ 已通報CTO，修復狀態未確認"],
    ["無lead派單/SLA自動釋放機制（CTO已設計但未部署）", "07-21前", "❌ 未部署"],
], col_widths=[8, 3, 6])
add_body(doc, "行銷端（CMO）的AI顧問/SEO/內容在正常把訪客轉成有效詢問；業務承接端（CRM通知→真人跟進）完全沒有運作。這代表過去三週投入的內容產出，商業轉換價值目前是0——除非現在補上這段。", color=RED, bold=True)

# ── 三、次要警訊 ──
add_heading1(doc, "三、次要警訊")
add_heading2(doc, "① mcstation.ai 有17篇文章寫好卻沒上線")
add_body(doc, "這批內容卡在等一個叫「Chat Brain」的架構上線才接手發布，目前形同閒置產能——已經花的產出成本沒有變成任何SEO/流量價值。")
add_heading2(doc, "② 待審閱清單（DELIVERABLES_CHECKLIST.md）累積速度超過Jason review速度")
add_body(doc, "目前28筆未勾選項目中12筆是CMO的，且全部集中在07-24～07-26這3天內產生。清單設計初衷是「產出後等你確認品質」，但堆積速度已經超過確認速度，清單正在失去「待辦」的意義，變成純紀錄。")
add_heading2(doc, "③ 專案owner標記與實際執行不一致")
add_body(doc, "ACTIVE_PROJECTS.md把Jason-SEO標記為CEO-PERSONAL，但實際內容產出/SEO/AI顧問維護一直是CMO在做，登記冊需要跟現實同步。")

# ── 四、下一輪策略 ──
add_heading1(doc, "四、下一輪策略（依優先序）")

add_heading2(doc, "P0（最高槓桿）修復轉換漏斗——先接住現有39筆詢問再談新增", RED)
for t in ["找CTO/COO一起排查並補上MONIQUE_LINE_USER_ID，確認Monique真的能收到LINE通知",
          "跟進hasValidContact()防呆修復狀態，確認垃圾聯絡方式不會再混進CRM",
          "推動CTO已設計但未部署的lead派單/SLA自動釋放機制真正上線",
          "這一項的投資報酬率遠高於「再寫10篇文章」——現有39筆詢問已經是花成本換來的，接不住就是純損失"]:
    add_body(doc, "• " + t, indent=0.4)

add_heading2(doc, "P1 決定mcstation.ai那17篇閒置內容的去留", GOLD)
add_body(doc, "• 建議不等「Chat Brain」，比照現有5篇上線頁面沿用的靜態路由機制先發布，或明確排出上線時程，不要繼續閒置", indent=0.4)

add_heading2(doc, "P2 安排一次批次審閱，清空待審閱清單", GOLD)
add_body(doc, "• 找一段時間集中過目12筆CMO待審閱項目，避免清單持續累積失去追蹤意義", indent=0.4)

add_heading2(doc, "P3 延續已驗證有效的內容策略", GREEN)
for t in ["「坪效提升王」公司願景定位＋「自助設備＝行銷活動」框架已驗證對AI顧問回答品質有實質提升，可以複製到目前還沒套用這個角度的既有頁面",
          "持續每站的場域補缺口內容策略（B2B/GEO/特定產業垂直），近三週證明這個方法論可以規模化執行"]:
    add_body(doc, "• " + t, indent=0.4)

add_heading2(doc, "P4 專案登記冊校正", GREY)
add_body(doc, "• 修正ACTIVE_PROJECTS.md，把Jason-SEO的owner標記同步成實際執行狀況（CMO）", indent=0.4)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("— AI CMO 產出．2026-07-27 —")
set_font(r, size=10, color=GREY, bold=True)

out_path = r"C:\Users\JasonLee\claude_code_projects\CMO\docs\AI-CMO_三週戰果檢討與下一輪策略_20260727.docx"
doc.save(out_path)
print("Saved:", out_path)
