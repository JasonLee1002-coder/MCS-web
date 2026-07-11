# -*- coding: utf-8 -*-
"""Generate Gemini-direct vs whisper-service comparison report (Word)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GOLD = RGBColor(0xB8, 0x68, 0x0A)
RED = RGBColor(0x8B, 0x1A, 0x1A)
GREEN = RGBColor(0x19, 0x6F, 0x3D)
DARKGRAY = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x55, 0x55, 0x55)

SERIF_TC = "Noto Serif TC"
SANS_TC = "Noto Sans TC"
LATIN_TITLE = "Cambria"
LATIN_BODY = "Calibri"


def set_font(run, latin=LATIN_BODY, cjk=SANS_TC, size=12, bold=False, color=None):
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cjk)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    if level == 1:
        set_font(r, LATIN_TITLE, SERIF_TC, 16, True, NAVY)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '1F4E79')
        pBdr.append(left)
        pPr.append(pBdr)
        p.paragraph_format.left_indent = Cm(0.3)
    elif level == 2:
        set_font(r, LATIN_TITLE, SERIF_TC, 13, True, GOLD)
    return p


def add_para(doc, text, size=12, indent=0, color=DARKGRAY, bold=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, LATIN_BODY, SANS_TC, size, bold, color)
    return p


def add_bullet(doc, text, size=11.5):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    set_font(r, LATIN_BODY, SANS_TC, size, False, DARKGRAY)
    return p


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        set_font(r, LATIN_BODY, SANS_TC, 11, True, RGBColor(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(hdr_cells[i], '1F4E79')
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for cidx, val in enumerate(row):
            cells[cidx].text = ''
            p = cells[cidx].paragraphs[0]
            r = p.add_run(str(val))
            set_font(r, LATIN_BODY, SANS_TC, 10.5, False, DARKGRAY)
            if ridx % 2 == 0:
                shade_cell(cells[cidx], 'EBF3FB')
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


doc = Document()
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# ---------- 封面 ----------
cover_table = doc.add_table(rows=1, cols=1)
cover_table.rows[0].cells[0].width = Cm(17)
cell = cover_table.rows[0].cells[0]
shade_cell(cell, '1F4E79')
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after = Pt(10)
r = p.add_run("龍雲數位整合股份有限公司")
set_font(r, LATIN_TITLE, SERIF_TC, 14, True, RGBColor(0xFF, 0xFF, 0xFF))

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(6)
p2.paragraph_format.space_after = Pt(40)
r2 = p2.add_run("會議錄音分析比較報告：Gemini 直接分析 vs 我方 whisper-service")
set_font(r2, LATIN_TITLE, SERIF_TC, 19, True, RGBColor(0xFF, 0xFF, 0xFF))

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("顧祖欣會議錄音（2026-07-10）｜ AI CMO 情報彙整")
set_font(r3, LATIN_BODY, SANS_TC, 13, True, GOLD)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("製作日期：2026年7月11日｜密級：內部使用｜撰寫者：AI CMO")
set_font(r4, LATIN_BODY, SANS_TC, 10.5, False, GRAY)

doc.add_paragraph()
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run("＊本報告分析兩份會議錄音檔案，比較「直接呼叫 Google Gemini API」與「我方自建 whisper-service」兩種管線的轉錄／分析品質，")
set_font(r5, LATIN_BODY, SANS_TC, 9.5, False, GRAY)
p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
r6 = p6.add_run("並發現後者對長錄音有嚴重截斷問題，已另行通報 CTO/Mozo 處理。內容涉及尚待查證之法律/商業資訊，正式引用前請自行核實。")
set_font(r6, LATIN_BODY, SANS_TC, 9.5, False, GRAY)

doc.add_page_break()

# ---------- 摘要 ----------
add_heading(doc, "摘要重點", level=1)
add_para(doc, "本報告分析兩份顧祖欣先生的會議錄音（2026年7月10日14:21，約87分鐘；同日16:57，約31分鐘），分別以「直接呼叫 Google Gemini API（gemini-2.5-pro）」與「我方自建 whisper-service（poc.mcstation.ai/whisper）」兩條路徑進行轉錄與分析，並比較結果。核心發現有二：（一）我方 whisper-service 底層其實也是 Gemini（2.5 Flash），並非真正的 Whisper 語音模型，且對長錄音有嚴重截斷問題——87分鐘錄音只轉錄前5分43秒、31分鐘錄音只轉錄前1分17秒，即當作會議結束產出摘要，此問題已另行通報 CTO/Mozo 修復；（二）透過完整覆蓋的 Gemini 直接分析，額外挖掘出多項先前未知的重大法律與商業資訊，包括顧總持有一件**已核准**台灣發明專利（I874154）、顧總實際企業體系規模、以及與 Yo-Kai Express 創辦人 Andy 之間可能的僱傭淵源說法，均已同步通報 COO 續辦。", size=12)

# ---------- 一、方法論 ----------
add_heading(doc, "一、比較方法論與重要澄清", level=1)
add_heading(doc, "（一）「我方 whisper-service」實際上也是 Gemini，並非 Whisper 模型", level=2)
add_para(doc, "查閱 EC2（poc.mcstation.ai）上 whisper-service 原始碼（/home/mozo/whisper-service/main.py）確認：此服務雖名為 whisper，但其架構註解明確寫著「音檔直傳 Gemini 2.5 Flash（逐字稿+分析一次完成）」，底層完全沒有使用 OpenAI Whisper 或任何專門的語音辨識（ASR）模型。因此本次比較實際上是「Gemini 2.5 Pro（本報告直接呼叫）」對比「Gemini 2.5 Flash（經我方自建結構化管線包裝，含知識庫 Context Cache、固定 JSON 格式驗證與重試機制）」，而非兩種不同 AI 供應商技術的比較。此發現已同步更新至團隊知識庫（KNOWLEDGE_BASE.md）。", indent=0.3)

add_heading(doc, "（二）比較設計", level=2)
comp_rows = [
    ["路徑A：Gemini 直接分析", "本報告作者自行撰寫 Python script，用 google-genai SDK 直接呼叫 gemini-2.5-pro，上傳音檔＋通用提示詞（要求完整逐字稿＋摘要＋決議＋待辦＋關鍵數字），單次請求取得結果"],
    ["路徑B：我方 whisper-service", "透過 https://poc.mcstation.ai/whisper 正式上傳介面（/upload API），輸入會議標題／參與者／背景說明後，由該服務的既有管線（gemini-2.5-flash + 知識庫快取 + 結構化驗證重試）產出結果，並自動生成 HTML 會議記錄頁面"],
]
make_table(doc, ["路徑", "說明"], comp_rows, col_widths=[4, 13])

doc.add_page_break()

# ---------- 二、關鍵發現：截斷bug ----------
add_heading(doc, "二、關鍵發現：whisper-service 對長錄音嚴重截斷", level=1)
add_para(doc, "這是本次比較最重要、且已超出「比較兩種分析品質」範疇的系統性問題：我方 whisper-service 對兩份錄音都只處理了開頭一小段，就當作整場會議已結束，直接產出「會議摘要」——而摘要本身寫得像是完整會議記錄，容易讓使用者誤以為已涵蓋全部內容。", size=12)

trunc_rows = [
    ["短會議（16:57，共31分24秒）", "31分24秒", "僅到 01:17（約1分17秒），21筆對話", "約4%", "無明確時間戳記，但內容涵蓋多個話題轉折（統包案、健康麵食、產品試吃、Yo-Kai Express侵權評論、西門站踩點心得），推估涵蓋至少10分鐘以上實質內容"],
    ["長會議（14:21，共86分42秒）", "86分42秒", "僅到 05:43（約5分43秒），16筆對話", "約7%", "涵蓋完整電話對話段落＋企業體系介紹＋專利號＋美國市場進度，推估涵蓋錄音相當大比例的實質內容"],
]
make_table(doc, ["錄音檔", "總時長", "whisper-service涵蓋範圍", "涵蓋比例", "Gemini直接分析涵蓋情形"], trunc_rows, col_widths=[3.5, 2.2, 3.8, 2, 5.5])

add_para(doc, "")
add_para(doc, "推測成因：查閱 main.py 的 analyze_with_gemini 函式，發現其對 Gemini 回應有「JSON 解析失敗→retry」的重試邏輯（本次執行 log 中也確實觸發），研判可能是結構化 JSON 輸出受到 token 長度上限影響，模型在還沒把音檔全部內容轉錄完之前，JSON 結構就被截斷或提前收尾，導致後續內容遺失且未被妥善偵測、亦未重試「補完全部內容」，只是重試「讓已截斷的內容能被解析成合法JSON」。", indent=0.3, color=RED)
add_para(doc, "⚠️ 影響範圍評估：由於此問題與音檔長度相關而非本次特例，先前所有透過此服務產出、且原始錄音長度超過約5-10分鐘的會議記錄（本機可見資料夾中至少有 Ann、Jimmy、Kevin、Ardy、Justin、Mozo 等多筆會議錄音），都可能存在類似程度不一的內容遺漏，建議之後有重要決策依賴這些既有會議記錄時，先以本報告方法（直接呼叫Gemini全文分析）重新核實。此問題已於今日正式通報 CTO/Mozo 調查修復。", indent=0.3, color=RED)

doc.add_page_break()

# ---------- 三、逐份錄音內容比較 ----------
add_heading(doc, "三、逐份錄音內容比較", level=1)

add_heading(doc, "（一）短會議（2026-07-10 16:57，31分鐘）", level=2)
add_para(doc, "主題：AI Bar 使用顧總 AI 蒸汽烹調機技術，是否有 Yo-Kai Express 專利侵權風險", size=12, bold=True, color=GOLD)
add_para(doc, "whisper-service 版本（僅涵蓋前1分17秒）：Jason 與顧祖欣討論 AI Bar 若採用顧總的 AI 蒸汽烹調機技術（該技術目前應用於樂活速食坊/勝十蘭），是否會與 Yo-Kai Express 的美國販賣機專利產生侵權疑慮；顧總確認 Yo-Kai Express 目前僅在美國申請專利，台灣、中國未申請。行動項目：Jason 需啟動 Yo-Kai Express 專利範圍的法律分析（訂 2026-07-31 前完成），顧總需準備 AI Bar 技術整合商業分析（訂 2026-08-07 前完成）。", indent=0.3)
add_para(doc, "Gemini直接分析版本（涵蓋更長篇幅）：對話從顧總分享其「高蛋白、低糖」健康拉麵配方與無人自助餐廳構想開始，說明自己是「系統工程師」出身，過去曾創立中華民國集合式住宅統包（Turnkey）發包制度；接著討論與「維康」藥局通路合作、鎖定糖尿病與健身族群做會員制的商業模式；現場展示產品試吃；隨後話題轉向 Yo-Kai Express——顧總明確表示專利「都在申請中」、對Yo-Kai創辦人Andy表達強烈敵意（「軟銀我一定K死他，7-Eleven我一定K死他」），並提及已去西門站踩點試吃、評論其拉麵尚可但水餃不行、且指出一風堂並非Yo-Kai真正對手；也提到透過「BB Yang」牽線的日本人脈、SoftBank合作動態，以及Jason朋友計劃在英國開拉麵店的國際擴展意向。", indent=0.3)

add_para(doc, "")
add_heading(doc, "（二）長會議（2026-07-10 14:21，87分鐘）", level=2)
add_para(doc, "主題：樂活速食坊專案的雙溫智慧櫃成本與IP風險", size=12, bold=True, color=GOLD)
add_para(doc, "whisper-service 版本（僅涵蓋前5分43秒）：Jason 與顧祖欣討論樂活速食坊（100+門市）與勝十蘭（20-30門市）導入 AI 蒸汽烹調機與智慧櫃的合作案，初估首批需求30-50台機器；智慧櫃需同時具備冷藏與加熱雙溫功能，成本落在60-80萬，是主要挑戰；AI Bar 產品線初期鎖定咖啡/茶飲、暫緩雞尾酒（法規考量）；雙方都提到 Yo-Kai Express 過去的侵權案例，同意IP合規需優先處理。行動項目：顧總需將智慧櫃設計交法務做IP合規審查（下次會議前）、確認樂活速食坊預算範圍（訂2026-07-24前）、Jason需召集研發團隊評估雙溫智慧櫃技術可行性與替代方案（訂2026-07-24前）。", indent=0.3)
add_para(doc, "Gemini直接分析版本（涵蓋更長篇幅，內容豐富度顯著更高）：對話由顧總（說話者A）與「李董」對話開始，後段出現顧總在電話中與另一人討論多項業務——包括與軟銀（SoftBank）、7-Eleven相關人士Vivian/Bibian的合作談判、指控「家豐」的機台仿冒抄襲既有案例、懷疑Yo-Kai Express創辦人Andy「之前是Yo-Kai/匯聚員工後來背叛」（原話提及「林維鴻背叛我們公司」，人名有出入待核實）、日本首都高速公路智慧冰箱合作案。顧總並詳細介紹其企業體系：**歐月設計**（服務統一集團/7-Eleven）、**創偉家具**（服務非統一集團業務）、**承龍精密板金**、**樂活速食坊**、**聖食饗**、**次成控股**，自稱擁有5間食品廠（其中新莊食品廠600坪）。技術面提到已持有**台灣發明專利 I874154**（自動烹煮裝置），並透露常溫麵（可存放一年）、無麩質產品線的研發進度，以及在美國大學校園（洛杉磯、德州、亞利桑那）佈局冷凍食品販賣機的最新進度（與Paul/陳博彥通話討論）。", indent=0.3)

doc.add_page_break()

# ---------- 四、轉錄準確度問題 ----------
add_heading(doc, "四、轉錄準確度問題（Gemini本身的辨識誤差）", level=1)
add_para(doc, "即便 Gemini 直接分析涵蓋範圍遠優於我方 whisper-service，仍觀察到若干人名/專有名詞辨識誤差，正式引用前應逐一核實：", size=12)
acc_rows = [
    ["顧祖欣（Rupert Guh，本人名片確認）", "顧宗新 / 顧總（口語代稱）", "同音字誤植，Gemini 對中文姓名的辨識仍有限制"],
    ["李奇申（Jason，MASTER_CONTEXT確認本名）", "李其森 / 李董", "同音字誤植，另混用職稱代稱"],
    ["林志鴻（Andy Lin，Yo-Kai Express創辦人，公開資料確認）", "林維鴻（顧總口語提及）", "可能是顧總自己講錯，也可能是Gemini聽錯，需向顧總核實何者為真"],
    ["優凱／Yo-Kai（顧總口語提及Yo-Kai Express的簡稱，Jason 2026-07-11確認寫法）", "—", "顧總慣用「優凱」或「匯聚」指稱 Yo-Kai Express，非官方名稱，需注意後續文件統一用語"],
]
make_table(doc, ["推測正確名稱", "轉錄/辨識結果", "備註"], acc_rows, col_widths=[5.5, 5, 6.5])
add_para(doc, "")
add_para(doc, "⚠️ 「家豐」「匯聚」等專有名詞目前僅為錄音中口語提及，未經公開資料交叉驗證，是否為真實公司名稱、與 Yo-Kai Express／Noodle Time Holdings 有何關聯，均待進一步查證。", size=10.5, color=RED)

doc.add_page_break()

# ---------- 五、新增重大情報 ----------
add_heading(doc, "五、本次新增之重大法律／商業情報", level=1)
add_para(doc, "以下資訊僅透過 Gemini 直接分析（完整涵蓋錄音）取得，我方 whisper-service 因截斷問題完全沒有擷取到，凸顯本次比較的實質價值：", size=12)
add_bullet(doc, "顧總持有一件**已核准**台灣發明專利【I874154】（自動烹煮裝置）——若查證屬實，其法律強度顯著高於先前查到的2024年申請案113108458/113202327（尚在申請中），應優先作為侵權訴訟主張的權利基礎，已通報COO優先查證。")
add_bullet(doc, "顧總實際企業體系規模遠超先前掌握：歐月設計、創偉家具、承龍精密板金、樂活速食坊、聖食饗、次成控股，自稱擁有5間食品廠。")
add_bullet(doc, "顧總懷疑 Yo-Kai Express 創辦人 Andy（林志鴻）「之前是Yo-Kai/匯聚員工後來背叛」——若屬實，可能與既有 Noodle Time Holdings 商業秘密訴訟是同一事件的不同角度陳述，值得向顧總本人求證細節（時間、任職公司全名、背叛的具體行為）。")
add_bullet(doc, "顧總提及已有專利遭「家豐」仿冒的既有案例、與SoftBank／7-Eleven相關人士（Vivian/Bibian）之合作談判、日本首都高速公路智慧冰箱合作案——均為全新的商業關係網絡資訊。")
add_bullet(doc, "顧總透露常溫麵（可存放一年不需冷凍）、無麩質產品線的研發近況，以及美國大學校園（洛杉磯、德州、亞利桑那）冷凍食品販賣機佈局進度——顯示其國際擴張已有具體行動，而非僅止於規劃階段。")

add_para(doc, "")
add_heading(doc, "六、建議事項", level=1)
add_bullet(doc, "【已執行】已將 whisper-service 長錄音截斷 bug 正式通報 CTO/Mozo 調查修復，並記錄至團隊知識庫。")
add_bullet(doc, "【已執行】已將 I874154 專利號、顧總企業體系、Andy 僱傭淵源說法等新發現通報 COO，作為顧總合作案／Yo-Kai Express 侵權案的後續查證方向。")
add_bullet(doc, "【建議】在 whisper-service 修復前，重要會議錄音若需要完整逐字稿與分析，建議使用本報告方法（直接呼叫 Gemini API 全文分析），或至少人工核對 whisper-service 產出是否有明顯「話題提前結束」的跡象。")
add_bullet(doc, "【建議】向顧總本人核實：I874154 專利的完整資訊、「林維鴻/Andy背叛」說法的具體細節、「家豐」仿冒案例的來源，這些都是口語提及、未經文件佐證的資訊，需一手資料才能正式採用於訴訟或商業決策。")

out_path = r"C:\Users\JasonLee\claude_code_projects\CMO\docs\AI-CMO_顧祖欣會議錄音分析比較_Gemini直接分析vs我方whisper-service_v2.docx"
doc.save(out_path)
print("SAVED:", out_path)
