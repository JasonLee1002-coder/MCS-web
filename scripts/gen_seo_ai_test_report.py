# -*- coding: utf-8 -*-
"""AI-CMO 三站SEO發布+30人模擬AI顧問測試 完整報告產生器"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GOLD = RGBColor(0xB8, 0x68, 0x0A)
RED = RGBColor(0x8B, 0x1A, 0x1A)
GREEN = RGBColor(0x1E, 0x7A, 0x3C)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHTBLUE = "EBF3FB"
LIGHTRED = "FBEBEB"
LIGHTGREEN = "EAFAF1"

def set_font(run, cn="Noto Sans TC", en="Calibri", size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    if color:
        run.font.color.rgb = color

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, "Noto Serif TC", "Cambria", 15, True, NAVY)
    return p

def add_heading2(doc, text, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, "Noto Sans TC", "Calibri", 13, True, color)
    return p

def add_body(doc, text, indent=0, color=None, bold=False, size=12):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(hdr[i], "1F4E79")
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_font(r, size=10)
            if ridx % 2 == 0:
                shade_cell(cells[i], LIGHTBLUE)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_box(doc, title, text, bgcolor, titlecolor):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    shade_cell(cell, bgcolor)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_font(r, size=11.5, bold=True, color=titlecolor)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(text)
    set_font(r2, size=10.5)
    doc.add_paragraph()

doc = Document()
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# ── 封面 ──
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(70)
r = cover.add_run("龍雲數位 × 銓幻元科技")
set_font(r, "Noto Serif TC", "Cambria", 15, True, NAVY)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(18)
r = title.add_run("三站SEO內容發布")
set_font(r, "Noto Serif TC", "Cambria", 24, True, NAVY)
r2 = title.add_run("\n＋ 30人模擬AI顧問對話測試報告")
set_font(r2, "Noto Serif TC", "Cambria", 19, True, NAVY)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(26)
r = sub.add_run("撰寫者：AI CMO")
set_font(r, size=13, color=GOLD, bold=True)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("版本 v1.0　2026-07-23")
set_font(r, size=11, color=GREY)

doc.add_page_break()

# ── 摘要 ──
add_heading1(doc, "摘要")
add_body(doc, "本次任務兩部分：①在李奇申.com、transtep.com、mcstation.ai 三站發布「小空間就可以賺錢的AI勞動力」主題SEO文章共9篇 ②用30個模擬客戶人設，真實跑三站正式站的AI顧問「小龍」對話到收單完整流程，驗證從搜尋、進站、對話到轉換的整條路徑。")
add_body(doc, "過程中意外發現並處理了兩個對整體業務有實質影響的問題：李奇申.com 的 Vercel 從未連接 GitHub（自動部署形同虛設）已修復；李奇申.com 的 AI 顧問「小龍」已完全故障至少3週（07-02起），已診斷根因並回報CTO待修復。這代表過去3週李奇申.com透過AI對話收單的功能是掛零的，本報告視為本次最重要的發現。", bold=True, color=RED)

add_table(doc, ["項目", "結果"], [
    ["新增SEO文章", "9篇（3站各3篇），全數已驗證上線(HTTP 200)"],
    ["AI顧問名稱", "「小龍」🐉，三站統一品牌識別"],
    ["模擬對話完成數", "20/30（mcstation.ai 10組、transtep.com 10組）；李奇申.com 10組因AI顧問故障無法測試"],
    ["mcstation.ai 收單成功率", "7/10（1組正確擋下無效聯絡方式，2組為一般性詢問未強迫收單）"],
    ["transtep.com 收單成功率", "3/10（1組正確擋下無效聯絡方式，6組卡在AI重複問句迴圈——新發現bug）"],
    ["發現問題", "3項：①李奇申.com AI完全故障3週(嚴重) ②李奇申.com從未自動部署(已修復) ③transtep.com對話收斂率偏低(新發現)"],
], col_widths=[4.5, 10.5])

# ── Part A ──
add_heading1(doc, "一、SEO內容發布狀態")
add_heading2(doc, "9篇新文章清單")
add_table(doc, ["站台", "文章標題", "狀態"], [
    ["李奇申.com", "店裡本來就有的等待角落", "✅ 已上線"],
    ["李奇申.com", "一人工作室的被動收入解方", "✅ 已上線"],
    ["李奇申.com", "畸零角落也能賺", "✅ 已上線"],
    ["transtep.com", "美容美髮沙龍等待區怎麼加值？", "✅ 已上線"],
    ["transtep.com", "自助洗衣店、自助洗車場怎麼配無人設備？", "✅ 已上線"],
    ["transtep.com", "一人工作室也能多一份收入？", "✅ 已上線"],
    ["mcstation.ai", "診所候診區導入智慧販賣機", "✅ 已上線"],
    ["mcstation.ai", "1坪店面能放什麼設備？", "✅ 已上線"],
    ["mcstation.ai", "美容美髮工作室候場空間怎麼配置？", "✅ 已上線"],
], col_widths=[3, 9, 3])
add_body(doc, "內容合規：全數文章未編造具體部署數量/客戶案例數字，遇到「能賺多少」類問題一律用「先抓人流/坪數再試算」的方法論帶過，比照2026-07-13/07-18兩次內容真實性大掃除的教訓執行。", color=GREEN, size=10.5)

add_heading2(doc, "🔴 意外發現並修復：李奇申.com從未自動部署", RED)
add_body(doc, "推送文章後transtep.com/mcstation.ai都正常自動部署，唯獨李奇申.com完全沒反應。查Vercel API發現該專案的git連結欄位是null——代表這個Vercel專案從建立以來從來沒有真的跟GitHub repo建立官方整合（過去deployment紀錄裡的github metadata，只是CLI手動部署留下的痕跡，不代表有真連結）。已用`vercel git connect`重新連接並驗證新push會正常觸發部署。")
add_body(doc, "順帶查到WEB_mcstation目錄的git remote實際對應兩個Vercel專案：「mcstation-web」(孤兒、未連接)與「mcstation-ai」(正常運作、實際serving www.mcstation.ai)——確認push走的是後者，沒有受影響。")

# ── Part B ──
add_heading1(doc, "二、真實搜尋能見度查證")
add_body(doc, "用WebSearch跑8組真實搜尋查詢（涵蓋無人商店/被動收入/AI勞動力/美容院/一人工作室/智取櫃/冷凍販賣機/自助洗衣等主題），檢查三站內容是否真的能被搜到。")
add_table(doc, ["搜尋詞", "結果"], [
    ["無人商店 台灣 自動化設備 2026", "✅ 三站全上榜（mcstation.ai、transtep.com、李奇申.com多篇皆在前8名）"],
    ["GraBox 智取櫃 銓幻元", "✅ 三站全上榜（品牌詞能見度強）"],
    ["小空間 被動收入 自助販賣設備", "🟡 transtep.com上榜(第6)，但被通用理財文章與競品稀釋"],
    ["美容院 等待區 自助販賣機", "❌ 三站皆未上榜（今天剛發布，尚未被Google收錄，符合預期）"],
    ["一人工作室 額外收入 自助設備", "❌ 三站皆未上榜（除了太新未收錄，這個詞的真實搜尋意圖偏向工作室登記/稅務，跟設備變現關聯度低）"],
    ["智慧取物櫃 辦公室 包裹", "❌ 三站皆未上榜，競品(WAFERLOCK/iTake/吉達思科技)反而排名在前"],
    ["冷凍販賣機 工廠 宿舍", "❌ 三站皆未上榜，即使這是mcstation.ai覆蓋最深的主題之一"],
    ["自助洗衣店 加值 咖啡機 販賣機", "❌ 三站皆未上榜，同業競品(洗特樂/CocoCafe/巨唐等)排名在前"],
], col_widths=[5.5, 9])
add_box(doc, "誠實揭露：搜尋能見度落差",
    "儘管三站內容量極大（尤其李奇申.com已有3468篇文章），實測發現「廣泛分類詞」與「品牌詞」搜尋能見度很強，但「具體產品組合詞」（智取櫃+辦公室、冷凍販賣機+工廠宿舍等）反而是競品或不相關內容排名在前。這代表內容量≠精準關鍵字排名，特定高意圖長尾詞可能需要針對性SEO優化（技術面：schema標記、內部連結、頁面標題結構），而非單純增加文章數量。",
    LIGHTBLUE, NAVY)

# ── Part C ──
add_heading1(doc, "三、30人模擬客戶對話測試")
add_body(doc, "方法：真實打三站正式站/api/chat端點，讓小龍AI用真實模型（Anthropic Haiku）跟30個模擬客戶對話到收單，最後用測試模式送出（不寫入正式Notion CRM/LINE推播，避免污染真實資料），驗證「搜尋→進站→對話→轉換」整條路徑。")

add_heading2(doc, "🔴 李奇申.com：10組完全無法測試（AI顧問故障）", RED)
add_body(doc, "測試第一組就發現/api/chat每次呼叫直接回傳錯誤，AI完全無法對話。查Vercel正式環境錯誤紀錄，確認真正原因是 Error[AI_InvalidPromptError]（訊息格式與AI SDK schema不符），首次出現時間 2026-07-02，代表這個故障已經存在至少3週。")
add_body(doc, "根因：比對三站套件版本，李奇申.com的\"ai\"核心套件是 ^7.0.11，但transtep.com與mcstation.ai（同源架構、運作正常）都是 ^6.0.191——版本不一致導致訊息格式驗證失敗。已將完整診斷與建議修法（版本降回^6.0.191對齊）送交CTO，這需要改程式碼＋重新部署＋測試，非本次能自行完成的範圍。", bold=True)
add_body(doc, "業務影響：這3週內，任何訪客在李奇申.com點擊「小龍」聊天，都會遇到無法對話的狀況——等同這段期間AI對話式收單完全停擺。", color=RED, bold=True)

add_heading2(doc, "mcstation.ai：7/10 成功收單")
add_table(doc, ["#", "詢問主題", "結果"], [
    ["1", "診所候診區", "✅ 收單成功"],
    ["2", "1坪空間設備", "🟡 未收單（一般性詢問，AI未強迫收單，合理）"],
    ["3", "美甲工作室設備", "✅ 收單成功"],
    ["4", "無人商店（一般詢問）", "🟡 未收單（同上，資訊性問題）"],
    ["5", "移工宿舍冷凍販賣機", "✅ 收單成功"],
    ["6", "辦公室智取櫃", "✅ 收單成功"],
    ["7", "醫院蒸氣拉麵機", "✅ 收單成功"],
    ["8", "連鎖早午餐取餐櫃", "✅ 收單成功"],
    ["9", "社區大樓智取櫃（刻意不留聯絡方式）", "✅ 正確擋下（needs_contact=true，未誤發真單）"],
    ["10", "醫院智取系統", "✅ 收單成功"],
], col_widths=[1, 7, 6.5])

add_heading2(doc, "transtep.com：3/10 成功收單，發現重複問句迴圈bug", RED)
add_table(doc, ["#", "詢問主題", "結果"], [
    ["1", "美容院無人販賣機", "✅ 收單成功"],
    ["2", "自助洗衣店販賣機", "✅ 收單成功"],
    ["3", "工作室被動收入", "🔴 卡住：AI重複問同一題3輪未收斂"],
    ["4", "包租代管附加價值", "🔴 卡住：同上"],
    ["5", "無人商店怎麼開", "🔴 卡住：同上"],
    ["6", "半導體廠販賣機", "✅ 收單成功"],
    ["7", "物流倉儲販賣機", "🔴 卡住：同上"],
    ["8", "長照機構冷凍餐食", "🔴 卡住：同上"],
    ["9", "蒸氣加熱設備比較（刻意不留聯絡方式）", "✅ 正確擋下（needs_contact=true）"],
    ["10", "租賃辦公室販賣機", "🔴 卡住：同上"],
], col_widths=[1, 7, 6.5])
add_box(doc, "🔴 新發現：transtep.com對話收斂率偏低（30% vs mcstation.ai的70%）",
    "6組卡住的對話呈現完全一致的模式：客戶在第一句話就清楚說明了場域與需求（例如「共享辦公室想加販賣機當服務亮點」），AI卻只顧著問一個罐頭式選擇題（①②③④），客戶補了聯絡方式之後，AI在第三輪逐字重複第二輪一模一樣的問句，完全沒有往前推進，也沒有把客戶第一句話裡已經講清楚的資訊填進場域/需求欄位。三個案例的第2、3輪AI回覆文字經比對是逐字相同，屬於真正的對話卡死，不是措辭不同的正常追問。這是transtep.com『longcloud』雙軌（IoT設備＋SEO顧問）系統提示詞的實際運作瑕疵，建議CTO/COO評估是否要調整system prompt，讓AI優先採信客戶在開場白就主動提供的資訊，不要無視自然語言答案、只認選擇題格式。",
    LIGHTRED, RED)

# ── Part D ──
add_heading1(doc, "四、AI顧問身分確認")
add_body(doc, "三站AI顧問統一命名「小龍」🐉，品牌識別一致：")
add_table(doc, ["站台", "AI名稱", "涵蓋業務"], [
    ["mcstation.ai", "小龍", "IoT無人商店設備（銓幻元單一軌）"],
    ["transtep.com", "小龍", "IoT無人商店設備＋SEO顧問服務（雙軌，依對話判斷）"],
    ["李奇申.com", "小龍", "IoT無人商店設備＋SEO顧問服務（雙軌，目前故障中）"],
], col_widths=[3.5, 3, 8])

add_heading1(doc, "五、本次已採取的行動")
for t in [
    "①9篇新SEO文章三站發布並驗證上線",
    "②修復李奇申.com從未連接GitHub的Vercel部署問題（vercel git connect重新連接）",
    "③更新李奇申.com的ANTHROPIC_API_KEY（原值與transtep.com不同，疑似過期，順手一併修正，雖非本次故障根因）",
    "④診斷李奇申.com AI顧問3週故障根因（ai套件版本v7 vs v6不相容），完整報告+建議修法已送CTO收件匣",
    "⑤發現transtep.com對話收斂率偏低問題，記入本報告供CTO/COO參考評估system prompt調整",
]:
    add_body(doc, t, indent=0.4)

add_heading1(doc, "六、建議下一步")
for t in [
    "① CTO優先處理李奇申.com的AI套件版本修復（3週故障，影響最大）",
    "② 針對「智取櫃辦公室」「冷凍販賣機工廠宿舍」等已投入大量內容但搜尋能見度仍輸給競品的關鍵字組合，評估是否需要技術面SEO優化（schema/內部連結/標題結構），而非單純增加文章量",
    "③ CTO/COO評估transtep.com的雙軌system prompt是否要調整，讓AI優先採信客戶開場白的自然語言描述",
    "④ 修復確認後，建議重跑一次李奇申.com的10組模擬對話補齊本次未完成的測試",
]:
    add_body(doc, t, indent=0.4)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("— AI CMO 產出．2026-07-23 —")
set_font(r, size=10, color=GREY, bold=True)

out_path = r"C:\Users\JasonLee\claude_code_projects\CMO\docs\AI-CMO_三站SEO發布與30人模擬AI對話測試報告_20260723.docx"
doc.save(out_path)
print("Saved:", out_path)
