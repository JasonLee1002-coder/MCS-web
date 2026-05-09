"""
產生說帖函文 DOCX：欣殿萬飲智慧外帶取物站
致高雄市政府文化局
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

# ── 路徑 ──────────────────────────────────────────────
ASSET_DIR = "G:/我的雲端硬碟/公司經營戰略 & 科專規劃/2026科專-提升商業服務業營運效能強化韌性計畫計畫/高雄駁二特區-欣殿萬飲"
IMG_MAP  = os.path.join(ASSET_DIR, "messageImage_1777554962475.jpg")   # 鳥瞰定位圖
IMG_CAB1 = os.path.join(ASSET_DIR, "智能櫃_004.jpg")                    # 智慧櫃實景
IMG_CAB2 = os.path.join(ASSET_DIR, "智能櫃_001.jpg")                    # 備用圖
IMG_MWD  = "G:/我的雲端硬碟/龍雲數位/2022-新公司/銓幻元/業務/麥味登/麥味登-01.jpg"  # 麥味登實景
OUT_PATH = os.path.join(ASSET_DIR, "智慧外帶取物站說帖_致高雄市政府文化局_20260509.docx")

doc = Document()

# ── 頁面設定 ──────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2)

# ── 預設字型（標楷體）─────────────────────────────────
from docx.oxml.ns import qn
style = doc.styles['Normal']
style.font.name = '標楷體'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

def set_font(run, size=12, bold=False, color=None):
    run.font.name = '標楷體'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p

def body(text, size=12, indent=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(indent)
    r = p.add_run(text)
    set_font(r, size=size)
    return p

def add_image(path, width=Cm(14), caption=None):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cp.runs:
            set_font(r, size=10, color=(100,100,100))

def divider():
    p = doc.add_paragraph('─' * 40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        set_font(r, size=9, color=(180,180,180))

# ════════════════════════════════════════════════════
#  封面區
# ════════════════════════════════════════════════════
doc.add_paragraph()
heading('龍雲科技股份有限公司', size=13, align=WD_ALIGN_PARAGRAPH.RIGHT)
heading('Transtep Technology Co., Ltd.', size=10, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
doc.add_paragraph()
doc.add_paragraph()

heading('說　　帖', size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
heading('欣殿萬飲 × 龍雲科技', size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
heading('智慧外帶取物站導入計畫', size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
heading('暨 駁二藝術特區 智慧城市示範點申請說明', size=13, align=WD_ALIGN_PARAGRAPH.CENTER, color=(60,60,60))
doc.add_paragraph()

body('致：高雄市政府文化局', size=13)
body('發文單位：龍雲科技股份有限公司（統編：90965685）', size=12)
body('發文日期：中華民國 114 年 5 月 9 日', size=12)
body('承辦聯絡：李奇申  lcs@transtep.com  ／ 台北市大同區長安西路78巷4弄10號1樓', size=12)

doc.add_page_break()

# ════════════════════════════════════════════════════
#  主旨
# ════════════════════════════════════════════════════
heading('壹、主旨', size=14)
divider()
body(
    '為響應高雄市政府推動「5G 應用」、「智慧城市」及「智慧雨林產業創生」等重點政策，'
    '本公司擬於駁二藝術特區內之欣殿萬飲（大勇路100號）導入 IoT 智慧外帶取物站，'
    '打造全臺首座結合歷史倉庫建築美學與現代科技的智慧取餐示範場域，'
    '敬請 鈞局惠予支持並協助協調相關許可事宜，請　查照。',
    indent=1
)
doc.add_paragraph()

# ════════════════════════════════════════════════════
#  場域說明 + 鳥瞰圖
# ════════════════════════════════════════════════════
heading('貳、場域位置說明', size=14)
divider()
body(
    '欣殿萬飲位於高雄駁二藝術特區核心地帶，緊鄰大港橋旋轉開合秀景點，'
    '周邊匯聚高雄流行音樂中心、福容大飯店，每逢假日人流量龐大，'
    '為駁二最具能見度之商業據點之一。',
    indent=1
)
doc.add_paragraph()

add_image(IMG_MAP, width=Cm(14),
    caption='圖一：欣殿萬飲位置鳥瞰圖（黃色標示處）——緊鄰大港橋、高雄流行音樂中心')

doc.add_paragraph()

# ════════════════════════════════════════════════════
#  設備說明 + 智慧櫃圖
# ════════════════════════════════════════════════════
heading('參、智慧外帶取物站設備說明', size=14)
divider()

body('（一）設備名稱：GraBox 智慧取物站（IoT 多格溫控智慧櫃）', size=12)
body('（二）安裝位置：店面外牆（嵌入現有牆體開口，不佔人行道面積）', size=12)
body('（三）設備規格：', size=12)

specs = [
    ('外觀尺寸', '寬 80cm × 深 55cm × 高 180cm（嵌入牆面，不外凸）'),
    ('格口數量', '12～16 格，含常溫格／保溫格／冷藏格'),
    ('操作介面', '觸控螢幕 + QR Code 掃碼取餐，免接觸'),
    ('連網方式', '5G / Wi-Fi 雙模，搭配 OmniCore 雲端後台即時管控'),
    ('LED 看板', '嵌入式全彩 LED，呈現品牌形象與取餐資訊（非廣告燈箱）'),
    ('電力需求', '單相 110V，10A，獨立迴路，無需額外施工'),
]

for k, v in specs:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(f'{k}：')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(v)
    set_font(r2, size=12)

doc.add_paragraph()
add_image(IMG_CAB1, width=Cm(13),
    caption='圖二：智慧外帶取物站實景示意圖——嵌入駁二磚牆，兩位遊客正在操作取餐')

doc.add_paragraph()

# ════════════════════════════════════════════════════
#  政策呼應
# ════════════════════════════════════════════════════
heading('肆、與高雄市政府重點政策之呼應', size=14)
divider()

policies = [
    (
        '一、5G 智慧應用推廣',
        '本設備全程透過 5G/Wi-Fi 連網，實現遠端庫存監控、即時溫度警報、消費數據分析，'
        '是 5G 應用落地於零售服務業的具體實踐，可作為高雄市 5G 場域實證案例。'
    ),
    (
        '二、智慧雨林產業創生計畫',
        '龍雲科技為本土 IoT 新創，本次導入正是產業創生計畫所鼓勵的「在地科技解決方案商業化」'
        '之典型案例，有助於驗證臺灣自研智慧設備的市場可行性，並帶動相關產業鏈發展。'
    ),
    (
        '三、駁二智慧城市形象提升',
        '駁二藝術特區兼具歷史感與潮流氛圍，智慧外帶取物站以工業美學外觀嵌入百年磚牆，'
        '呈現「科技與文化共融」的視覺衝擊，能有效強化駁二的智慧城市識別，'
        '吸引媒體報導與社群擴散。'
    ),
    (
        '四、遊客體驗升級',
        '大港橋旋轉開合秀散場後，遊客於戶外等待外帶餐飲時，可透過 App 預訂、'
        '抵達後 QR Code 直接取餐，減少排隊時間，提升旅遊滿意度，'
        '為駁二增添「科技驚喜感」的差異化記憶點。'
    ),
    (
        '五、示範典範推廣潛力',
        '本場域一旦成功落地，可作為高雄市標竿示範案例，供其他文創園區、'
        '觀光景點複製導入，形成高雄特有的「智慧文創商業模式」，'
        '有利於市府對外招商與城市行銷。'
    ),
]

for title, content in policies:
    heading(title, size=13, color=(20, 80, 150))
    body(content, indent=1)
    doc.add_paragraph()

# ════════════════════════════════════════════════════
#  麥味登標竿案例
# ════════════════════════════════════════════════════
heading('伍、標竿參考：全臺最大上市連鎖早餐品牌「麥味登」已率先導入', size=14)
divider()
body(
    '麥味登（My Warm Day，股票代號 2723，全臺逾 2,000 間直營加盟門市）'
    '為臺灣規模最大之上市連鎖餐飲業者。',
    indent=1
)
body(
    '2025 年起，麥味登已正式與龍雲科技簽訂「智能取餐櫃 POC 專案合約」，'
    '在多間門市外牆導入「My Express TAKE OUT」智慧取物站，'
    '消費者可透過 App 預訂、QR Code 掃碼自助取餐，大幅縮短排隊時間。',
    indent=1
)
body(
    '▶ 若連全臺最大連鎖餐飲品牌都已在一般街道門市外牆導入此設備，'
    '高雄駁二藝術特區率先採用，不僅是跟隨趨勢，更是走在全臺文創園區的最前端。',
    indent=1
)
doc.add_paragraph()
add_image(IMG_MWD, width=Cm(13),
    caption='圖三：麥味登門市「My Express TAKE OUT」智慧外帶取物站實景——外牆嵌入式，不佔騎樓空間')
doc.add_paragraph()

# ════════════════════════════════════════════════════
#  安全與景觀說明（針對局處疑慮）
# ════════════════════════════════════════════════════
heading('陸、針對主管機關關切事項之說明', size=14)
divider()
body(
    '本公司充分尊重 鈞局對駁二園區景觀維護之用心，就外牆智慧取物站之安全性與美觀性，'
    '提出以下具體保證：',
    indent=1
)

concerns = [
    ('不佔人行道', '設備採嵌入式安裝，完全收納於現有牆面開口內，外凸深度為零，行人通行空間不受任何影響。'),
    ('外觀融合歷史建築', '機體外框採工業鐵件材質，配色呼應紅磚色調，LED 看板以低亮度運行，不產生光害，視覺上與駁二老倉庫風格協調。'),
    ('無噪音、無廢氣', '設備為純電力靜音運作，壓縮機噪音低於 45dB（低於一般對話音量），不影響周遭環境品質。'),
    ('可逆性施工', '若日後政策調整，設備可完整移除，牆面恢復原狀，不留永久結構性損壞。'),
    ('符合消防法規', '所有電力配置均依消防安全規範施作，並取得合格電氣承裝業者施工憑證。'),
]

for k, v in concerns:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(f'{k}：')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(v)
    set_font(r2, size=12)

doc.add_paragraph()

# ════════════════════════════════════════════════════
#  請求
# ════════════════════════════════════════════════════
heading('柒、懇請事項', size=14)
divider()
body(
    '敬請 高雄市政府文化局惠予支持，協助協調下列事項：',
    indent=1
)
requests = [
    '協助確認外牆智慧取物站裝設之相關許可程序（建管、景觀審查等）',
    '若需補件，請惠予書面通知，本公司將全力配合',
    '建議將本計畫納入「駁二智慧城市示範點」推廣案例',
]
for r in requests:
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(r)
    set_font(run, size=12)

doc.add_paragraph()

# ════════════════════════════════════════════════════
#  結語
# ════════════════════════════════════════════════════
heading('捌、結語', size=14)
divider()
body(
    '高雄駁二藝術特區是臺灣最具代表性的文創再生場域之一。'
    '龍雲科技希望藉由此次合作，以科技為橋梁，為駁二注入新一層的智慧生命力，'
    '讓每一位造訪的遊客都能感受到「高雄，走在智慧城市的前端」。'
    '懇請　鈞局惠予指導與支持，共同打造具全臺示範性的智慧文創場域。',
    indent=1
)
doc.add_paragraph()
body('此致', indent=1)
body('高雄市政府文化局', indent=3, size=13)
doc.add_paragraph()
body('龍雲科技股份有限公司', align=WD_ALIGN_PARAGRAPH.RIGHT)
body('董事長　李奇申　敬上', align=WD_ALIGN_PARAGRAPH.RIGHT)
body('中華民國 114 年 5 月 9 日', align=WD_ALIGN_PARAGRAPH.RIGHT)

doc.add_paragraph()
divider()
body('【附件一】設備規格說明書（另附）', size=11)
body('【附件二】欣殿萬飲現場平面配置圖（另附）', size=11)
body('【附件三】GraBox 智慧取物站安全認證文件（另附）', size=11)
body('【附件四】麥味登×龍雲科技智能取餐櫃 POC 合約（另附，供參）', size=11)

# ════════════════════════════════════════════════════
#  儲存
# ════════════════════════════════════════════════════
doc.save(OUT_PATH)
print(f"✅ 已儲存：{OUT_PATH}")
