# -*- coding: utf-8 -*-
"""
產出《分租 PO 文投放與操作說明》Word 文件。
用法：python scripts/gen_rental_fb_playbook.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GOLD = RGBColor(0xB8, 0x68, 0x0A)
RED = RGBColor(0x8B, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x22, 0x22, 0x22)

OUT = (r"G:\我的雲端硬碟\2025_銓幻元_MCS相關資料\臨時紀錄雜事\FB分租PO_20260817"
       r"\AI-CMO_分租PO投放與操作說明_20260817.docx")

CN_H = "Noto Serif TC"
CN_B = "Noto Sans TC"


def setfont(run, name=CN_B, size=12, bold=False, color=BLACK):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), name)
    r.set(qn('w:ascii'), "Calibri")
    r.set(qn('w:hAnsi'), "Calibri")


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def para(doc, text, size=12, bold=False, color=BLACK, indent=0, space_after=6,
         name=CN_B, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if align:
        p.alignment = align
    setfont(p.add_run(text), name, size, bold, color)
    return p


def h1(doc, text):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '8'); left.set(qn('w:color'), '1F4E79')
    bdr.append(left); pPr.append(bdr)
    p.paragraph_format.left_indent = Cm(0.25)
    setfont(p.add_run(text), CN_H, 14.5, True, NAVY)


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    setfont(p.add_run(text), CN_H, 13, True, GOLD)


def bullets(doc, items, indent=0.6, size=11.5):
    for it in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.3
        setfont(p.add_run("· " + it), CN_B, size)


def callout(doc, title, body, fill="FFF3CD", color=GOLD):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, fill)
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    setfont(p.add_run(title), CN_H, 12, True, color)
    p2 = c.add_paragraph()
    p2.paragraph_format.line_spacing = 1.3
    setfont(p2.add_run(body), CN_B, 11, False, BLACK)
    doc.add_paragraph()


def table(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(header):
        c = t.rows[0].cells[i]
        shade(c, "1F4E79")
        c.text = ""
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        setfont(p.add_run(htext), CN_B, 10.5, True, RGBColor(0xFF, 0xFF, 0xFF))
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1:
                shade(cells[i], "EBF3FB")
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.25
            setfont(p.add_run(val), CN_B, 10, False, BLACK)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    doc.add_paragraph()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)

    # ---------- 封面 ----------
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0); shade(c, "1F4E79"); c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(16)
    setfont(p.add_run("銓幻元科技"), CN_B, 11, True, RGBColor(0xC8, 0xDC, 0xF0))
    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(18)
    setfont(p2.add_run("分租 PO 文\n投放與操作說明"), CN_H, 24, True,
            RGBColor(0xFF, 0xFF, 0xFF))
    doc.add_paragraph()
    para(doc, "大坪林站一樓空間　×　南京三民站開放式商辦", 13, True, NAVY, name=CN_H)
    para(doc, "撰寫：AI CMO　｜　2026-08-17", 10.5, False, GREY)
    para(doc, "本文件為對內文件。所有對外文案、圖片與看屋頁另存於同一資料夾。",
         10.5, False, GREY, space_after=2)

    callout(doc, "發布前唯一一句話",
            "任何印在封面圖卡或寫進文案的「專屬」「可用」「可進駐」，都必須能由主約、"
            "平面圖或實測證明。證明不了，就先別發。\n"
            "（這句是跨模型稽核三輪之後留下的結論，也是這次實際踩到的坑——"
            "「專屬衛浴」是從平面圖推論出來的，並不存在，已全數移除。）",
            "FDEDEC", RED)

    # ---------- 一 ----------
    h1(doc, "一、結論先講")
    para(doc, "大坪林這篇補完四件事就可以發；南京三民這篇現在不要發。", 12, True)
    bullets(doc, [
        "大坪林：補押金金額、泱璇的聯絡方式（並先取得她同意公開）、"
        "前庭車位到底給不給分租方固定使用、側院能不能讓分租方專用——補完就能 PO。",
        "南京三民：租金、樓層、坪數、押金、管理費、可進駐日、車位、工商登記 八項全空。"
        "陌生人看到沒辦法判斷要不要聯絡，發出去只會收到一堆「多少錢」然後沒下文。"
        "先當成內部預覽頁，數字補齊再公開。",
    ])

    # ---------- 二 ----------
    h1(doc, "二、要 PO 到哪裡")
    para(doc, "兩案都是商用空間，不要投住宅類通路——住宅客會問能不能住、"
              "能不能養寵物，問了也不會成交。", 11.5, False, GREY)

    h2(doc, "【一定要投】")
    table(doc,
          ["通路", "為什麼", "注意"],
          [["你本人的 FB 公開貼文＋限時動態",
            "「一起用空間」這種事，熟人背書比制式出租有效。來的多半是設計、行銷、"
            "電商、接案團隊本人，或朋友轉介。",
            "貼文設「公開」才轉得出去。第一句就寫明你是承租人、房東同意分租。"],
           ["591 商用（辦公／住辦）",
            "搜尋意圖最強，會來直接比較交通、坪數、價格的人。591 沒有全面禁二房東，"
            "站上本來就有「承租分租」「共享辦公」類物件。",
            "資料要真實填：價格、坪數、樓層、經辦身分，不可自稱屋主。"
            "坪數是必填欄位，所以上 591 就不能維持「不寫坪數」。"
            "591 規則禁止放非政府的外部網址，看屋頁連結只能用在 FB／LINE／LinkedIn。"
            "付費刊登，買刊期前先把資料補齊。"],
           ["地區型與同溫層 FB 社團",
            "大坪林投新店在地、接案／工作室類；南京三民投台北創業、"
            "中小企業、辦公室分租類。",
            "社團名稱我沒有寫死，因為社團會關、會改規則、會突然只剩廣告。"
            "用第三節的檢查表自己搜一次再投。"]],
          [3.6, 6.4, 6.4])

    h2(doc, "【值得投】")
    table(doc,
          ["通路", "為什麼", "注意"],
          [["好房網快租",
            "基本刊登免費，刊登規範明文包含二房東，可以補到 591 以外的搜尋者。",
            "可能要求物件證明文件，先把主約與房東同意備好。"],
           ["樂屋網 商用辦公",
            "台北商用辦公出租頻道確實存在，多一個曝光面。",
            "承租人身分怎麼填，刊登前先問平台客服。"],
           ["LinkedIn 個人貼文",
            "B2B 決策者密度比一般 FB 租屋社團高，適合工程、顧問、業務、外商分點。",
            "南京三民那案要先有座位數、月總成本、能不能登記這三項才值得投。"],
           ["既有 LINE 工作群、客戶群",
            "轉介品質最高，成交最快。",
            "只發已有關係的群，不要陌生大量私訊。"],
           ["會計師／記帳士／公司設立顧問（南京三民專用）",
            "這些人最常接觸正在設公司、搬辦公室的客戶。",
            "工商登記還沒確認前，不要拿「可設址」當賣點。"]],
          [3.6, 6.4, 6.4])

    h2(doc, "【不要投】")
    table(doc,
          ["通路", "原因"],
          [["FB Marketplace",
            "商辦不是商品也不是住宅，分類會錯放；來的訊息多半是「還在嗎」和詐騙式私訊。"],
           ["學生租屋、套房、找室友社團",
            "36,000 的團隊總價和商用需求對不上，多數社團也禁商用與二房東。"],
           ["現階段買 FB 廣告加強推廣",
            "範圍、用途、價格都還沒完全定案，付費只會放大不精準的詢問。"
            "先跑自然觸及和 591，累積 10–15 個有效詢問再判斷。"]],
          [4.5, 11.9])

    # ---------- 三 ----------
    h1(doc, "三、投社團之前的四個檢查")
    para(doc, "社團名單我刻意不寫死。名稱、活躍度、規則都會變，寫死了三個月後就是錯的。"
              "改成一張你自己能維護的表：", 11.5)
    bullets(doc, [
        "最近 30 天還有沒有自然貼文？（只剩廣告的社團投了沒用）",
        "允不允許商用租賃、允不允許二房東？（規則寫在社團公告或置頂）",
        "要不要審文？審多久？",
        "有沒有限制發文頻率或禁止外部連結？",
    ])
    para(doc, "建議欄位：搜尋詞｜社團網址｜最近自然貼文日期｜准不准商用／二房東｜"
              "要不要審文｜你實際發文日。一次投 2–3 個就好，不要同一分鐘洗版。",
         11.5, False, GREY)

    # ---------- 四 ----------
    h1(doc, "四、怎麼 PO（素材組合與節奏）")

    h2(doc, "【貼文本身】")
    bullets(doc, [
        "FB 不吃 Markdown，文案檔案裡是純文字，複製貼上就好，不要再加符號。",
        "前兩行是唯一會被完整看到的部分，其他會被「查看更多」收起來——"
        "所以第一行放捷運距離和價格，第二行放「我是承租人、房東同意分租」。",
        "圖片一次上傳，順序照資料夾裡的檔名編號。第一張是封面圖卡，"
        "第二張是標色平面圖——先讓人搞懂範圍，再看空間。",
        "影片不要跟照片混在同一篇。隔一到兩天單獨發，或發成 Reels。",
        "看屋頁的連結放第一則留言，不要放本文。"
        "（註：Meta 從沒公告「本文帶外連一定降觸及」這條規則，"
        "這是實務上多數人的觀察，你可以兩種都試一次再決定。）",
    ])

    h2(doc, "【發文時段】")
    para(doc, "個人帳號沒有後台數據可以參考，先測週二、週三的 12:10–13:10 和 "
              "18:30–20:00 這兩個時段，比較哪個回應多。", 11.5)

    h2(doc, "【一篇不夠，排三篇】")
    table(doc,
          ["", "第一篇", "第二篇（3–4 天後）", "第三篇（清空後）"],
          [["大坪林",
            "講清楚專用與共用範圍，主圖用平面圖",
            "講成本：2–3 人怎麼分、每人多少，附雙捷運",
            "「已清空、可進駐」，全部換新照片重拍"],
           ["南京三民",
            "（數字補齊才發）完整價格與坪數",
            "可以放幾席、每席成本、怎麼隔",
            "管理處、客貨梯、空調時段、登記與車位的最終答案"]],
          [2.4, 4.6, 4.7, 4.7])

    # ---------- 五 ----------
    h1(doc, "五、看屋頁怎麼用")
    para(doc, "看屋頁是一個網址，手機打開就能看，裡面有兩案的專用／共用對照、"
              "費用表、平面圖、全部照片，還有各 30 秒左右的空間影片。", 11.5)
    bullets(doc, [
        "用途：FB 貼文的第一則留言、LINE 轉傳、私訊回覆時直接丟連結。",
        "不要放進 591（平台規則禁外部網址）。",
        "頁面上還沒確認的欄位一律顯示成橘色「待補」，不是留白也不是含糊帶過。"
        "這是刻意的——少量待補是誠實，會加分；但南京三民那案八個欄位都待補，"
        "所以頁面上已經標明「條件確認中，尚未開放洽詢」。",
        "數字補齊之後告訴我，我更新同一個網址，你發出去的連結不用換。",
    ])

    # ---------- 六 ----------
    h1(doc, "六、還沒解決的事")
    table(doc,
          ["項目", "現況", "誰處理"],
          [["泱璇的電話／LINE",
            "文案裡是【待填】。也要先取得她同意公開在網路上。", "你"],
           ["大坪林押金金額", "文案與看屋頁都顯示「待補」。", "你"],
           ["前庭車位歸屬",
            "只有一個平面車位。文案現在寫「怎麼用一起談」，"
            "如果你要自己固定用，看屋時要先講清楚，不要讓對方誤會。", "你"],
           ["側院能不能專用",
            "平面圖的粉紅線切進側院，但我沒跟你確認過。"
            "現在文案寫「側院（範圍依平面圖）」，平面圖也加註「範圍提案，非最終約定」。",
            "你確認後我改"],
           ["大坪林工商登記",
            "要看主約與建物法定用途。文案寫成「先問我」，沒有承諾。", "你查主約"],
           ["南京三民八項數字",
            "租金、樓層、坪數、押金、管理費、可進駐日、車位、工商登記。"
            "另外「是不是整層」也要確認——同棟一層可能不只一戶，只憑影片不能寫整層。",
            "你補給我"],
           ["南京三民你會不會自留自用",
            "看屋頁原本寫「兩處我自己也在裡面用」，已改成不提；確認後再決定要不要寫回去。",
            "你確認"],
           ["大坪林租金 52,000 vs 68,000",
            "591 刊登寫 52,000，你說 68,000。不影響貼文內容，但影響你分租後的實際成本。",
            "你確認"]],
          [3.9, 8.6, 3.9])

    # ---------- 七 ----------
    h1(doc, "七、定價這件事，先讓你知道")
    para(doc, "查到兩筆同區刊登價，放在這裡不是要你降價，是讓你知道對方會拿什麼來比：",
         11.5)
    bullets(doc, [
        "大坪林站附近：37 坪一樓辦公室，開價 37,000／月，可登記。"
        "（591 商用 rent/21189692）",
        "大坪林站附近共享辦公：約 30 坪，推薦 4–6 個工位，18,000 起含水電與 Wi-Fi、"
        "19,000 起含登記。（591 商用 rent/21130943）",
        "南京三民站：五段 208 號 12F、50 坪、48,000／月、距站 131 公尺、含平面車位。"
        "注意這 48,000 含車位，不能直接除以 50 坪算單價。（591 商用 rent/21730579）",
    ])
    callout(doc, "這代表什麼",
            "以上都是刊登開價、不是成交價，也不完全同質。但它們會成為對方心裡的比較基準。\n"
            "大坪林分租 36,000、又不寫坪數的情況下，如果沒有固定車位、工商登記、"
            "24 小時進出這類具體條件，價格會很難撐。\n"
            "解法有兩條，選一條：把實際存在的條件補上去讓 36,000 有內容，"
            "或是把價格往下修。不存在的條件不能寫——那是前面那句警告的意思。",
            "FFF3CD", GOLD)

    # ---------- 附錄 ----------
    h1(doc, "附錄、檔案在哪裡")
    para(doc, r"G:\我的雲端硬碟\2025_銓幻元_MCS相關資料\臨時紀錄雜事\FB分租PO_20260817\\",
         10.5, True, NAVY)
    bullets(doc, [
        "01_大坪林站\\FB貼文_大坪林_主版本.txt —— 主版本、精簡版、第一則留言、"
        "PO 前確認清單（後兩段不要貼出去）",
        "01_大坪林站\\圖片\\ —— 12 張，檔名編號就是上傳順序",
        "02_南京東路五段\\FB貼文_南京東路五段_主版本.txt —— 同上，但標【　】的要先補",
        "02_南京東路五段\\圖片\\ —— 11 張",
        "看屋頁.html —— 本機檔；線上版網址見對話",
    ], size=11)

    doc.save(OUT)
    print("✓", OUT)


if __name__ == "__main__":
    build()
