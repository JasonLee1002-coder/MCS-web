"""
森林藥局聯盟 導流斷點追查與修復報告 → Word
收件人：黃維德 執行長（森林藥局聯盟 / 普健生物科技）
撰寫：AI CMO　日期：2026-08-11

依 docx-beautiful 規範：Noto Serif TC（標題）/ Noto Sans TC（正文），
主色 #1F4E79、輔色 #B8680A、警示 #8B1A1A。禁用標楷體/新細明體/微軟正黑體。
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GOLD = RGBColor(0xB8, 0x68, 0x0A)
RED = RGBColor(0x8B, 0x1A, 0x1A)
GREEN = RGBColor(0x1D, 0x6B, 0x4F)
GREY = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x22, 0x22, 0x22)

SERIF = "Noto Serif TC"
SANS = "Noto Sans TC"


def set_font(run, name=SANS, size=12, bold=False, color=BLACK):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hex_color)
    tcPr.append(el)


def para(doc, text, size=12, bold=False, color=BLACK, indent=0, space_after=6, name=SANS):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    set_font(p.add_run(text), name=name, size=size, bold=bold, color=color)
    return p


def h1(doc, text):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(text), name=SERIF, size=15, bold=True, color=NAVY)
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '1F4E79')
    bdr.append(bottom)
    pPr.append(bdr)
    return p


def h2(doc, text, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(text), name=SANS, size=13, bold=True, color=color)
    return p


def bullet(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(text), size=size)
    return p


def build_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        shade(c, '1F4E79')
        set_font(c.paragraphs[0].add_run(htxt), size=11, bold=True,
                 color=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            if ri % 2 == 0:
                shade(cells[ci], 'EBF3FB')
            set_font(cells[ci].paragraphs[0].add_run(val), size=11)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


doc = Document()
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Cm(2.4)
sec.top_margin = sec.bottom_margin = Cm(2.2)

# ── 封面 ──────────────────────────────────────────
cover = doc.add_table(rows=1, cols=1)
cover.rows[0].cells[0].text = ''
shade(cover.rows[0].cells[0], '1F4E79')
cp = cover.rows[0].cells[0].paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(cp.add_run('森林藥局聯盟 × 普健生物科技'), name=SERIF, size=13, bold=True,
         color=RGBColor(0xD6, 0xE4, 0xF7))
cp2 = cover.rows[0].cells[0].add_paragraph()
cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(cp2.add_run('搜尋流量導流斷點：追查結果與修復計畫'), name=SERIF, size=19,
         bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
cp3 = cover.rows[0].cells[0].add_paragraph()
cp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(cp3.add_run('顧客搜尋到我們之後，走不走得到藥局？'), size=11,
         color=RGBColor(0xD6, 0xE4, 0xF7))

doc.add_paragraph()
for label, val in [('呈閱', '黃維德 執行長（森林藥局聯盟 / 普健生物科技）'),
                   ('撰寫', 'AI CMO'),
                   ('日期', '2026 年 8 月 11 日'),
                   ('版本', 'v1')]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(f'{label}：'), size=11, bold=True, color=NAVY)
    set_font(p.add_run(val), size=11, color=GREY)

# ── 一、結論摘要 ──────────────────────────────────
h1(doc, '一、結論摘要')
para(doc, '我們在 Google 搜尋上已經開始有排名，但顧客讀完文章之後，'
          '目前「走不到」實體藥局——這條路中間是斷的。', size=12, bold=True)
para(doc, '本次逐頁追查後確認三個斷點，其中兩個屬於我方系統問題，'
          '已排入本週修復；一個屬於通路涵蓋範圍的現實限制，需要與貴方一起決定因應方式。',
     space_after=10)

build_table(doc,
            ['斷點', '狀態', '負責'],
            [['文章內沒有任何前往藥局的入口（47 篇）', '我方系統問題，本週修復', 'AI CMO'],
             ['藥局資料缺 Google 地圖連結，手機無法一鍵導航', '我方系統問題，本週修復', 'AI CMO'],
             ['僅 4 家門市，多數縣市讀者「附近沒有店」', '通路現實，需共同決策', '雙方'],
             ['藥局資料缺營業時間', '需貴方提供資料', '普健生醫']],
            widths=[9.0, 4.6, 2.6])

# ── 二、追查過程 ──────────────────────────────────
h1(doc, '二、追查過程與證據')
para(doc, '以下結論皆以實際檢查程式碼與內容檔案取得，非推測。', size=11, color=GREY)

h2(doc, '斷點一｜文章讀完沒有出口（最嚴重）', RED)
para(doc, '檢查本週復原上線的 47 篇藥局衛教文章，結果為：', indent=0.6)
bullet(doc, '內文含「森林藥局聯盟」導流區塊：0 篇')
bullet(doc, '內文含任何連往森林藥局聯盟網站的連結：0 篇')
bullet(doc, '兩者皆無、讀者無路可走：47 篇（100%）')
para(doc, '原因已查明：這批文章先前依「品牌聚焦」決策被暫時下架，執行下架時'
          '判定導流策略已無意義，因而一併移除了文章內的導流連結。本週依新策略'
          '將文章復原上線時，內容回來了，但導流管線沒有一併接回。',
     indent=0.6, space_after=4)
para(doc, '影響：顧客搜尋「抗生素什麼時候吃」「兒童退燒藥」這類問題找到我們、'
          '讀完專業內容、對藥師產生信任——然後關掉網頁，我們沒有給他下一步。',
     indent=0.6, bold=True, color=RED)

h2(doc, '斷點二｜藥局資料缺地圖連結', RED)
para(doc, '目前四家合作藥局的資料欄位有：名稱、藥師、電話、地址、專長、'
          'LINE 官方帳號、藥師照片。缺少 Google 地圖連結與營業時間。', indent=0.6)
para(doc, '影響：手機使用者看到的是一段地址文字，必須自行複製、切換到地圖 App、'
          '貼上搜尋。在「現在不舒服、想立刻找到最近藥局」的情境下，'
          '每多一個步驟就流失一批人。', indent=0.6)
para(doc, '另補充：營業時間的缺漏特別關鍵，因為我們目前排名表現最好的關鍵字群'
          '正是「24 小時藥局」「半夜買藥」這一類——讀者一定會問幾點關門，'
          '而我們現在答不出來。', indent=0.6, color=GOLD, bold=True)

h2(doc, '斷點三｜通路涵蓋範圍（非系統問題）', RED)
para(doc, '目前四家合作門市分布：', indent=0.6)
build_table(doc,
            ['藥局', '藥師', '縣市 / 區'],
            [['皇慶藥局', '黃維德', '新北市 板橋'],
             ['皇誼藥局', '（見網站）', '新北市 永和'],
             ['向宇藥局', '（見網站）', '台北市 士林'],
             ['國良藥局', '（見網站）', '台中市 西區']],
            widths=[5.0, 4.0, 7.2])
para(doc, '也就是說，只有雙北與台中西區一帶的讀者有機會走進門市。'
          '我們的內容是全台可搜尋的，但門市不是全台都有——這個落差不是靠網站'
          '能解決的，需要在文案策略上正面處理。', indent=0.6)

# ── 三、修復方案 ──────────────────────────────────
h1(doc, '三、修復方案（我方本週執行，不需貴方等待）')

h2(doc, '① 導流區塊補回 47 篇，並改為「雙軌」寫法', GREEN)
para(doc, '不採用單純「找附近藥局」的寫法。理由是：對台南、高雄、花蓮的讀者而言，'
          '「附近」並不存在，這種文案會讓八成讀者撞牆，反而傷害信任感。改為兩條路並行：',
     indent=0.6)
bullet(doc, '雙北／台中讀者 → 直接給門市資訊（地址、電話、一鍵導航、該店藥師 LINE）')
bullet(doc, '其他縣市讀者 → 給線上諮詢入口，由聯盟藥師線上回覆保健問題')
para(doc, '這樣不論讀者身在何處，都有一個成立的下一步，不會走進死路。'
          '同時符合既有原則：導流要低調、自然，不做生硬的廣告插入。', indent=0.6)

h2(doc, '② 藥局資料補上 Google 地圖連結', GREEN)
para(doc, '依現有地址產生標準 Google Maps 導航連結，手機點擊即可直接開啟導航，'
          '不需複製貼上。此項不需貴方提供額外資料，我方可直接完成。', indent=0.6)

# ── 四、需要貴方提供 ──────────────────────────────
h1(doc, '四、需要普健生醫提供或決策的事項')
para(doc, '以下六項會直接影響推廣成效與對外說法的正確性，煩請協助確認。', size=11, color=GREY)

build_table(doc,
            ['項目', '說明', '急迫性'],
            [['各店營業時間',
              '目前資料缺此欄位。「24 小時藥局」「半夜買藥」是我們排名最好的關鍵字群，'
              '讀者必問，現在答不出來。', '高'],
             ['開辦費服務項目完整清單',
              '目前依已確認資料列出 8 項。若實際服務內容更多，請補齊，'
              '避免對外說明不完整。', '高'],
             ['開辦費退費規則',
              '舊版網頁的按比例退費說法是綁在已廢除的年費制上，現暫改為「以合約為準」。'
              '一次性開辦費的退費規則需要正式定義。', '高'],
             ['馬偕授權商品完整名單',
              '官方文件寫「全系列擴增至 11 支」，但明確列名者為 8 支。'
              '我方不自行臆測補足，需要完整名稱才能對外說明 11 支。', '中'],
             ['壽險合作夥伴是否可具名',
              '因內部曾註記未經對方同意不宜對外具名，網頁現以「大型壽險」不具名描述。'
              '若已取得具名同意，可還原具名。', '中'],
             ['招商簡報內部數字矛盾',
              '「普健生醫 × 馬偕」簡報第 12 頁寫「前 500 家」、第 13 頁寫「前 50 家」。'
              '已依指示以 50 家為準，但簡報本身尚未更新，業務攜舊版出門恐口徑不一。', '高']],
            widths=[4.4, 9.4, 2.0])

# ── 五、後續 ──────────────────────────────────────
h1(doc, '五、後續追蹤')
para(doc, '導流修復完成後，我方將以 Google Search Console 實際數據追蹤成效，'
          '重點觀察兩項指標：文章頁到藥局頁的點擊轉換、以及線上諮詢的進線量。'
          '目前站上已累積 295 篇通過合規檢查的衛教文章，'
          '內容量能不是瓶頸——把讀完文章的人接住，才是。', space_after=4)
para(doc, '另需說明：本站內容在發布前皆會通過自動化檢查閘門，'
          '比對普健生醫官方禁字清單，違規用詞在上線前即被攔下，'
          '不會出現在正式頁面上。', color=GREY, size=11)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(p.add_run('AI CMO　2026.08.11'), size=10, color=GREY)

OUT = (r'H:\共用雲端硬碟\2026-銓幻元共用雲端硬碟\2026業務\黃維德-普健生醫'
       r'\AI-CMO_森林藥局聯盟導流斷點追查與修復報告_20260811.docx')
doc.save(OUT)
print('SAVED:', OUT)
