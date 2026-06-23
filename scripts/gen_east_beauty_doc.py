"""Generate 東方美二代接班人策略 Word document"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ===== HELPERS =====
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def run(para, text, bold=False, size=11, color='2C3E50', italic=False):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return r

def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string('1A5276')
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p

def divider(doc):
    p = doc.add_paragraph()
    r = p.add_run('─' * 70)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string('AAAAAA')

def make_table_header(table, headers, bg='1F3864'):
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, h, bold=True, size=11, color='FFFFFF')

def fill_row(table, row_idx, vals, bg='FFFFFF', bold_cols=None):
    bold_cols = bold_cols or []
    for j, val in enumerate(vals):
        cell = table.cell(row_idx, j)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run(p, val, bold=(j in bold_cols), size=10)

# ===== COVER =====
cover_tbl = doc.add_table(rows=1, cols=1)
cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cc = cover_tbl.cell(0, 0)
set_cell_bg(cc, '1F3864')

p = cc.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(28)
run(p, 'OMNICORE AI 顧問團隊 × 東方美集團', size=10, color='BDC3C7')

p2 = cc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p2, '東方美集團　二代接班人策略藍圖', bold=True, size=24, color='FFFFFF')

p3 = cc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p3, 'AI 時代的傳承、轉型與國際擴張', size=14, color='85C1E9')

p4 = cc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(18)
run(p4, '2026 年 6 月 23 日', size=11, color='AEB6BF')

p5 = cc.add_paragraph()
p5.paragraph_format.space_after = Pt(28)

doc.add_paragraph()

# ===== PREFACE =====
pt = doc.add_table(rows=1, cols=1)
pt.alignment = WD_TABLE_ALIGNMENT.CENTER
pc = pt.cell(0, 0)
set_cell_bg(pc, 'EBF5FB')

pp = pc.paragraphs[0]
pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
pp.paragraph_format.space_before = Pt(12)
run(pp, '📌 給東方美二代接班人的一句話', bold=True, size=13, color='1A5276')

pp2 = pc.add_paragraph()
pp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(pp2, '您繼承的，不只是一家餐廳——而是三十年物流網絡、供應鏈掌控力、與遍布全台的場域佈局。', size=12, color='2C3E50')

pp3 = pc.add_paragraph()
pp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
pp3.paragraph_format.space_after = Pt(12)
run(pp3, '在 AI 的翅膀下，這三項資產將讓東方美飛向國際。', bold=True, size=12, color='1A5276')

doc.add_paragraph()

# ===== SECTION 1: HERITAGE =====
h1(doc, '一、你繼承的護城河：三十年積累的競爭優勢')

p = doc.add_paragraph()
run(p, '許多企業追求數位轉型，卻忽略了真正的競爭壁壘不在科技，而在難以複製的實體基礎。東方美已握有三把金鑰，而你，是第一個能用 AI 解鎖它們的人。', size=11)

doc.add_paragraph()

tbl = doc.add_table(rows=4, cols=3)
tbl.style = 'Table Grid'
make_table_header(tbl, ['核心優勢', '具體內容', '對 AI 擴張的戰略價值'])
rows = [
    ('🏭 南區物流中心', '完整冷鏈倉儲、車隊調度、原物料管理', '新品項快速上架；AI自助吧供貨主幹道；日本加盟統一配送基礎'),
    ('🔗 完整供應鏈', '從原料採購到末端配送全面掌握', '咖啡、拉麵、義大利燉飯等自助吧品項無縫納入；降低各店備貨成本 40%+'),
    ('📍 多元場域佈局', '辦公、商圈、校園等多類型據點', '早餐高峰→全日自助雙模式切換；場域資源模組化複製到海外'),
]
for i, row in enumerate(rows):
    fill_row(tbl, i+1, row, 'EBF5FB' if i%2==0 else 'FFFFFF', bold_cols=[0])

doc.add_paragraph()
p = doc.add_paragraph()
run(p, '⚡ 關鍵認知：', bold=True, size=11, color='E74C3C')
run(p, '任何競爭對手若想從零建立這三項優勢，至少需要 10 年與數億資本投入。你已站在起跑點的終點線。', size=11)

doc.add_paragraph()
divider(doc)

# ===== SECTION 2: LOCATION STRATEGY =====
h1(doc, '二、選址策略：二代接班人的展店藍圖')

p = doc.add_paragraph()
run(p, '餐飲選址不是找一個「便宜的店面」——而是找一個', size=11)
run(p, '能讓 AI 自助 BAR 發揮最大效益的場域', bold=True, size=11, color='1A5276')
run(p, '。以下是二代接班展店的核心選址框架：', size=11)

doc.add_paragraph()

ltbl = doc.add_table(rows=5, cols=3)
ltbl.style = 'Table Grid'
make_table_header(ltbl, ['選址維度', '評估標準', '東方美 AI 模型加分項'], '154360')
lrows = [
    ('人流屬性', '早餐客群密度（通勤、辦公室、學校）', 'QR 點餐效率最高；自助 BAR 下午客源穩定'),
    ('場域類型', '辦公大樓 B1~1F / 商圈轉角 / 校園周邊', '三段式運營（早餐→自助→夜間）皆適用'),
    ('競爭態勢', '同商圈無台式早餐連鎖品牌', 'AI 早餐店建立搜尋壟斷優勢'),
    ('租金結構', '月租佔預估月營收 ≤ 15%', '自助 BAR 全日營收拉高天花板，租金合理性提升'),
]
for i, row in enumerate(lrows):
    fill_row(ltbl, i+1, row, 'EAF2FF' if i%2==0 else 'FFFFFF', bold_cols=[0])

doc.add_paragraph()
tt = doc.add_table(rows=1, cols=1)
tt.alignment = WD_TABLE_ALIGNMENT.CENTER
tc = tt.cell(0, 0)
set_cell_bg(tc, 'FEF9E7')
tp = tc.paragraphs[0]
tp.paragraph_format.space_before = Pt(8)
run(tp, '💡 選址心法：', bold=True, size=11, color='B7770D')
run(tp, '一家新店用 AI 自助 BAR 模式，等於同時在一個地點開了「早餐店＋咖啡廳＋簡餐站」三間店，租金效益是傳統模式的 3 倍。', size=10, color='7D6608')
tp2 = tc.add_paragraph()
tp2.paragraph_format.space_after = Pt(8)

doc.add_paragraph()
divider(doc)

# ===== SECTION 3: AI DIGITAL HUB =====
h1(doc, '三、AI 數位中央台：讓三十年變成智慧資產')

p = doc.add_paragraph()
run(p, '數位化的本質不是「換一套系統」，而是讓三十年的經驗與流程，變成可以被 AI 學習、加速、複製的智慧資產——讓你一個人管十間店，比老一代管一間店還輕鬆。', size=11)

doc.add_paragraph()

atbl = doc.add_table(rows=5, cols=2)
atbl.style = 'Table Grid'
make_table_header(atbl, ['建設項目', '二代接班人獲得的能力'], '1A5276')
arows = [
    ('🏭 物流數位化\n南區物流中心 AI 化', '庫存即時追蹤 + 自動補貨觸發\nAI 配送路線最佳化 → 降低 30% 物流成本'),
    ('🔗 供應鏈整合\n原物料採購預測', '供應商資料整合 + 採購預測模型\n庫存共享平台 → 各店零缺貨'),
    ('📲 智慧門市\n各店 AI 代理人', 'QR 點餐 + 智慧取餐櫃 + 客戶行為數據\n你坐在家就能看每間店的即時報告'),
    ('🌐 品牌數位門面\n官網 + APP（規劃中）', '品牌形象、加盟招募、線上預訂整合\n未來加盟主直接從 APP 下訂貨'),
]
for i, row in enumerate(arows):
    fill_row(atbl, i+1, row, 'EBF5FB' if i%2==0 else 'FFFFFF')

doc.add_paragraph()
divider(doc)

# ===== SECTION 4: AI SELF-SERVICE BAR =====
h1(doc, '四、AI 自助服務 BAR：讓場域全日產值最大化')

p = doc.add_paragraph()
run(p, '東方美的場域，早餐過後就閒置了。二代接班人最大的機會，就是把', size=11)
run(p, '同一個空間變成三倍收入的機器', bold=True, size=11, color='E74C3C')
run(p, '，人力成本幾乎不變。', size=11)

doc.add_paragraph()

btbl = doc.add_table(rows=4, cols=3)
btbl.style = 'Table Grid'
make_table_header(btbl, ['時段', '運營模式', '關鍵設計'], '1F618D')
brows = [
    ('☀️ 早餐高峰\n06:00 ~ 10:00', '傳統餐廳模式\n保留東方美品牌溫度', 'QR 點餐 + 智慧取餐\n快速翻桌率優化'),
    ('🤖 離峰自助\n10:00 ~ 22:00', 'AI 無人自助 BAR\n咖啡、拉麵、義大利燉飯', '東方美供應鏈供貨\n設備自動加熱出餐'),
    ('🌙 夜間輕運營\n22:00 ~ 隔日', '極低成本維持運作', '清潔人員定期巡邏\nCCTV 全程監控'),
]
bar_bgs = ['FEF9E7', 'E8F8F5', 'EAF2FF']
for i, row in enumerate(brows):
    fill_row(btbl, i+1, row, bar_bgs[i])

doc.add_paragraph()
p = doc.add_paragraph()
run(p, '🏪 7-11 等級模組化設計：', bold=True, size=11, color='1A5276')
for item in [
    '模組標準化：廚房設備、動線、補貨流程完全標準化，加盟主拿到 SOP 即可複製',
    '全場 CCTV：AI 異常偵測，遠端管理多間門市，大幅降低巡店成本',
    '智慧廚房設備：AI 咖啡機、自動加熱、智慧補貨——無需廚藝背景',
    '輕裝潢高格調：早餐結束直接切換自助吧模式，整套可複製推廣',
]:
    pi = doc.add_paragraph(style='List Bullet')
    run(pi, item, size=10)

doc.add_paragraph()
divider(doc)

# ===== SECTION 5: JAPAN =====
h1(doc, '五、日本市場：台式早餐的藍海商機')

qt = doc.add_table(rows=1, cols=1)
qt.alignment = WD_TABLE_ALIGNMENT.CENTER
qc = qt.cell(0, 0)
set_cell_bg(qc, '1A252F')
qp = qc.paragraphs[0]
qp.alignment = WD_ALIGN_PARAGRAPH.CENTER
qp.paragraph_format.space_before = Pt(12)
run(qp, '「在日本，每天真的去找早餐，找不到。」', bold=True, size=14, color='F9E79F', italic=True)
qp2 = qc.add_paragraph()
qp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(qp2, '「我也是 APP 一直找早餐。」', size=12, color='AEB6BF', italic=True)
qp3 = qc.add_paragraph()
qp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
qp3.paragraph_format.space_after = Pt(12)
run(qp3, '— 2026.06.23 東方美策略會議逐字記錄', size=10, color='7F8C8D')

doc.add_paragraph()
p = doc.add_paragraph()
run(p, '台灣「美而美」早餐文化（CP值高、效率快、選擇多）在日本幾乎不存在。這是一個巨大的市場空缺，也是東方美最懂、最能填補的商機。', size=11)

doc.add_paragraph()

jtbl = doc.add_table(rows=5, cols=3)
jtbl.style = 'Table Grid'
make_table_header(jtbl, ['階段', '里程碑', '執行內容'], '922B21')
jrows = [
    ('Phase 1', '勁耘日本樣板店', '東京台灣人聚集區開設 AI 早餐 BAR，以台灣客為核心測試模型'),
    ('Phase 2', '成功數據包裝', '整理翻桌率、人事成本、顧客滿意度，製作日本加盟招募手冊'),
    ('Phase 3', '加盟網絡展開', '招募日本在地加盟主，東南亞員工培訓體系，東方美負責供貨'),
    ('Phase 4', '平台訂閱帝國', '所有加盟店統一 Google 廣告，品牌搜尋壟斷，平台收訂閱費形成持續現金流'),
]
for i, row in enumerate(jrows):
    fill_row(jtbl, i+1, row, 'FDEDEC' if i%2==0 else 'FFFFFF', bold_cols=[0, 1])

doc.add_paragraph()
p = doc.add_paragraph()
run(p, '🌏 人力革命優勢：', bold=True, size=11, color='1A5276')
run(p, '日本面臨嚴重人力短缺，AI 智慧廚房設備操作無需廚藝背景，可大量招募東南亞移工，降低人力成本 40~60%，培訓週期縮短至 3~5 天。', size=11)

doc.add_paragraph()
divider(doc)

# ===== SECTION 6: BUDGET =====
h1(doc, '六、OmniCore 合作方案：市場 1/10 的費用，100% 的效果')

btbl2 = doc.add_table(rows=6, cols=3)
btbl2.style = 'Table Grid'
make_table_header(btbl2, ['服務項目', '市場行情', 'OmniCore 報價'])
bdata = [
    ('企業官網設計建置', '$80,000 ~ $200,000', ''),
    ('APP 原型設計', '$100,000 ~ $300,000', '整合套餐'),
    ('AI 代理人部署（各店）', '$30,000 ~ $80,000 / 店', ''),
    ('供應鏈數位化顧問', '$50,000 ~ $150,000', ''),
    ('市場行情總計', '$280,000 ~ $790,000+', '$50,000（節省逾 80%）'),
]
for i, row in enumerate(bdata):
    is_total = (i == 4)
    bg = 'D5F5E3' if is_total else ('EBF5FB' if i%2==0 else 'FFFFFF')
    fill_row(btbl2, i+1, row, bg, bold_cols=[0, 2] if is_total else [])

doc.add_paragraph()
divider(doc)

# ===== SECTION 7: ACTION =====
h1(doc, '七、接班人的立即行動清單')

act_tbl = doc.add_table(rows=5, cols=3)
act_tbl.style = 'Table Grid'
make_table_header(act_tbl, ['優先級', '行動項目', '說明'], '1B2631')
act_rows = [
    ('🔴 立即', '顧總授權確認', '確認 AI 自助吧導入東方美供應鏈的合作授權'),
    ('🔴 立即', '五萬元合約簽署', 'AI 數位中央台整合套餐正式啟動'),
    ('🟡 近期', '官網內容深化訪談', '安排 2~3 次現場訪談，完善品牌定位與功能模組'),
    ('🟡 近期', '展店選址啟動', '依選址框架篩選 3~5 個候選店面，AI 評估人流數據'),
]
for i, row in enumerate(act_rows):
    bg = 'FDEDEC' if '🔴' in row[0] else 'FEF9E7'
    fill_row(act_tbl, i+1, row, bg, bold_cols=[0, 1])

doc.add_paragraph()

# ===== CLOSING =====
ct = doc.add_table(rows=1, cols=1)
ct.alignment = WD_TABLE_ALIGNMENT.CENTER
cc2 = ct.cell(0, 0)
set_cell_bg(cc2, '1A252F')

cps = [
    ('🚀 給接班人的最後一句話', True, 14, 'F9E79F', 16),
    ('東方美花了三十年建立的物流、供應鏈、場域優勢，', False, 12, 'FFFFFF', 6),
    ('正是 AI 時代最難被複製的核心資產。', True, 13, '85C1E9', 6),
    ('你的使命，不是守住父執輩的成果——', False, 12, 'AEB6BF', 6),
    ('而是讓東方美的三十年，在 AI 的加持下，', False, 12, 'AEB6BF', 6),
    ('成為下一個三十年國際擴張的最強基礎。', True, 13, 'F9E79F', 6),
    ('台灣出發，日本落地，世界為家。', True, 15, 'FFFFFF', 10),
    ('本文件依 2026-06-23 Jason 李奇申 × 佳恩 東方美策略會議內容生成  |  OmniCore AI 顧問團隊', False, 9, '717D7E', 16),
]
for i, (text, bold, size, color, space_after) in enumerate(cps):
    if i == 0:
        cp = cc2.paragraphs[0]
        cp.paragraph_format.space_before = Pt(16)
    else:
        cp = cc2.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(space_after)
    run(cp, text, bold=bold, size=size, color=color)

output_path = r'C:\Users\JasonLee\claude_code_projects\CMO\docs\東方美二代接班人策略_20260623.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print('SUCCESS:', output_path)
