# -*- coding: utf-8 -*-
"""AI-CMO 蔡慧玲律師網域申請手把手教學"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GOLD = RGBColor(0xB8, 0x68, 0x0A)
RED = RGBColor(0x8B, 0x1A, 0x1A)
GREEN = RGBColor(0x1E, 0x7A, 0x3C)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHTBLUE = "EBF3FB"
LIGHTGOLD = "FFF3CD"
LIGHTGREEN = "E2EFDA"

def set_font(run, cn="Noto Sans TC", en="Calibri", size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    if color:
        run.font.color.rgb = color

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, "Noto Serif TC", "Cambria", 16, True, NAVY)
    return p

def add_heading2(doc, text, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, "Noto Sans TC", "Calibri", 13, True, color)
    return p

def add_body(doc, text, indent=0, color=None, bold=False, size=12):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p

def add_step(doc, num, title, detail_lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"步驟 {num}｜{title}")
    set_font(r, size=13, bold=True, color=NAVY)
    for line in detail_lines:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.6)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run("• " + line)
        set_font(r2, size=11.5)

def add_box(doc, title, text, bgcolor, titlecolor):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    shade_cell(cell, bgcolor)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_font(r, size=11.5, bold=True, color=titlecolor)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(text)
    set_font(r2, size=10.5)
    doc.add_paragraph()

doc = Document()
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# 封面
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run("手把手教學")
set_font(r, "Noto Serif TC", "Cambria", 15, True, NAVY)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(18)
r = title.add_run("蔡慧玲律師 個人網站\nwww.蔡慧玲.com 申請與上線教學")
set_font(r, "Noto Serif TC", "Cambria", 22, True, NAVY)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(26)
r = sub.add_run("撰寫者：AI CMO")
set_font(r, size=13, color=GOLD, bold=True)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("版本 v1.0　2026-07-28")
set_font(r, size=11, color=GREY)

doc.add_page_break()

# 總覽
add_heading1(doc, "在開始之前：整件事怎麼分工")
add_box(doc, "這份教學把工作分成兩塊",
    "第一塊（Part 1）是「買網域名稱」，這是律師您本人需要親自做的事，因為要用您自己的信用卡付款、您自己的帳號擁有這個網域，別人不能代替您買。步驟很簡單，跟在Momo購物網買東西差不多。\n\n第二塊（Part 2）是「架設網站、把網域接上網站」，這部分技術性比較高（涉及一種叫DNS的網路設定），建議直接交給龍雲數位技術團隊處理，您只需要在買好網域後，把帳號權限開放給我們，或是把登入資訊安全地提供給我們，接下來全部由我們完成，完成後您只要打開瀏覽器輸入網址確認網站有出現就好。",
    LIGHTBLUE, NAVY)

# Part 1
add_heading1(doc, "Part 1｜買網域名稱（律師本人操作，約10分鐘）")

add_step(doc, 1, "打開瀏覽器，前往 Vercel 網域購買頁面", [
    "在瀏覽器網址列輸入完整網址：https://vercel.com/domains",
    "（Vercel 是一家專門提供網站服務的國際公司，很多正式企業網站都是用它架設，安全可靠）",
])

add_step(doc, 2, "註冊/登入 Vercel 帳號", [
    "第一次使用會要求註冊，畫面右上角會有「Sign Up」（註冊）按鈕",
    "建議選「Continue with Email」用您的電子郵件註冊，輸入信箱後，Vercel會寄一封確認信到您信箱，點開信裡的連結完成註冊",
    "如果您已經有 Google 帳號，也可以直接選「Continue with Google」用Google帳號快速登入，比較省事",
])

add_step(doc, 3, "搜尋「蔡慧玲.com」這個網域", [
    "登入後回到 https://vercel.com/domains 這個頁面",
    "畫面中間會有一個搜尋框，輸入「蔡慧玲.com」（中文網域，直接打中文即可，不用轉拼音）",
    "按 Enter 或搜尋按鈕，系統會顯示這個網域可不可以購買、價格多少",
])

add_step(doc, 4, "確認可購買後，點選購買並結帳", [
    "如果顯示「Available」（可購買），點選旁邊的「Buy」或「Purchase」按鈕",
    "系統會請您輸入信用卡資訊付款（網域通常是一年一次的年費，中文.com網域一年大約新台幣1,000~2,000元左右，實際金額以網頁顯示為準）",
    "付款完成後，這個網域就正式屬於您了，帳號登入資訊是您自己的信箱+密碼，請自行妥善保管",
])

add_box(doc, "⚠️ 如果搜尋結果顯示「已被其他人註冊」",
    "代表這個網域名稱已經有人登記走了。這種情況請截圖給AI CMO或龍雲技術團隊，我們會協助確認狀況、討論備案（例如改用其他網域字尾，或查詢是否能協商轉讓）。",
    LIGHTGOLD, GOLD)

doc.add_paragraph()

# Part 2
add_heading1(doc, "Part 2｜網站架設與上線（交給龍雲技術團隊）")
add_body(doc, "網域買好之後，接下來這部分屬於技術設定，律師您不需要自己動手，請把下面兩項資訊提供給龍雲數位技術團隊即可：")

add_table_data = None
add_heading2(doc, "您需要提供給我們的東西")
add_body(doc, "① 剛剛註冊 Vercel 帳號用的電子郵件信箱\n② 是否已經有現成的網站內容（例如律師事務所簡介、個人經歷、聯絡方式、想放的照片等），如果還沒有也沒關係，我們可以協助規劃一個簡單專業的個人形象網站頁面內容，屆時會先給您過目確認再正式上線", indent=0.4)

add_heading2(doc, "接下來我們會做的事（不需要律師操心）")
for t in [
    "在您的 Vercel 帳號底下架設網站程式",
    "把 www.蔡慧玲.com 這個網址正確「接上」剛架好的網站（這步叫DNS設定，是比較技術性的部分）",
    "測試網站在電腦、手機上都能正常開啟",
    "上線後把最終網址回報給您，您只要打開瀏覽器輸入 https://www.蔡慧玲.com 看到網站正常顯示，就代表完成了",
]:
    add_body(doc, "• " + t, indent=0.4)

doc.add_paragraph()

# 名詞小百科
add_heading1(doc, "名詞小百科（看不懂的專有名詞查這裡）")
add_table = doc.add_table(rows=1, cols=2)
add_table.style = 'Table Grid'
add_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = add_table.rows[0].cells
for i, h in enumerate(["名詞", "白話解釋"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    r = p.add_run(h)
    set_font(r, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    shade_cell(hdr[i], "1F4E79")
rows_data = [
    ("網域 (Domain)", "就是網址，例如 www.蔡慧玲.com，是網站的「門牌地址」，別人輸入這個地址就能找到您的網站"),
    ("Vercel", "一家提供網站託管服務的公司，負責讓您的網站24小時都能被看到，很多知名企業都用它"),
    ("DNS 設定", "把「門牌地址」（網域）跟實際的「房子」（網站程式）連在一起的技術設定，比較技術性，交給我們處理即可"),
    ("上線 (Deploy)", "把做好的網站正式公開到網路上，讓所有人都能透過網址看到"),
]
for ridx, (a, b) in enumerate(rows_data):
    cells = add_table.add_row().cells
    for i, val in enumerate([a, b]):
        cells[i].text = ""
        p = cells[i].paragraphs[0]
        r = p.add_run(val)
        set_font(r, size=10.5)
        if ridx % 2 == 0:
            shade_cell(cells[i], LIGHTBLUE)

doc.add_paragraph()
add_heading1(doc, "卡關怎麼辦？")
add_body(doc, "任何一個步驟卡住、看不懂、或畫面跟教學不一樣，都不用自己摸索，直接把當下畫面截圖傳給 Jason 或龍雲數位技術團隊，我們會直接協助處理，不會讓您自己一個人卡關。")

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("— AI CMO 產出．2026-07-28 —")
set_font(r, size=10, color=GREY, bold=True)

out_path = r"C:\Users\JasonLee\claude_code_projects\CMO\docs\AI-CMO_蔡慧玲律師網域申請手把手教學_20260728.docx"
doc.save(out_path)
print("Saved:", out_path)
